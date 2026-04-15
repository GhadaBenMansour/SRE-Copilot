from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── helpers ────────────────────────────────────────────────────
def shd(cell, fill):
    p = cell.paragraphs[0]._p.get_or_add_pPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), fill)
    p.append(s)

def para_shd(p, fill):
    pr = p._p.get_or_add_pPr()
    s = OxmlElement('w:shd')
    s.set(qn('w:val'), 'clear'); s.set(qn('w:color'), 'auto'); s.set(qn('w:fill'), fill)
    pr.append(s)

def add_title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

def add_subtitle(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x60,0x60,0x60)

def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F,0x49,0x7D)

def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E,0x74,0xB5)

def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x23,0x85,0x56)

def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.size = Pt(11)

def add_bullet(doc, text, bold=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    if bold:
        r1 = p.add_run(bold); r1.bold = True; r1.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(5)
    para_shd(p, 'F2F2F2')
    r = p.add_run(text); r.font.name = 'Courier New'; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1E,0x1E,0x1E)

def add_note(doc, text, fill='FFF3CD', color=RGBColor(0x85,0x64,0x04)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    para_shd(p, fill)
    r = p.add_run(text); r.font.size = Pt(10); r.font.color.rgb = color

def add_sep(doc):
    p = doc.add_paragraph()
    r = p.add_run('─'*95); r.font.size = Pt(7)
    r.font.color.rgb = RGBColor(0xCC,0xCC,0xCC)
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)

def make_table(doc, headers, rows, col_fill='1F497D'):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for run in c.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            run.font.size = Pt(10)
        shd(c, col_fill)
    for row_data in rows:
        row = t.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
            for p in row[i].paragraphs:
                for r in p.runs: r.font.size = Pt(10)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph(); doc.add_paragraph()
add_title(doc, "SRE Copilot")
add_subtitle(doc, "Référence complète — Commandes & Fichiers du projet")
doc.add_paragraph()
add_subtitle(doc, "Stage PFE · PwC DigiTech / Cloud Native Operations")
add_subtitle(doc, "Remédiation GitOps pilotée par un agent IA orchestré avec n8n")
doc.add_paragraph()
add_subtitle(doc, "Avril 2026")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ══════════════════════════════════════════════════════════════
add_h1(doc, "Table des matières")
toc = [
    ("1.", "État actuel du projet et prochaines étapes"),
    ("2.", "Commandes de mise en place du cluster"),
    ("3.", "Commandes de déploiement des composants"),
    ("4.", "Commandes de correction et debugging"),
    ("5.", "Commandes de test et validation"),
    ("6.", "Commandes de maintenance quotidienne"),
    ("7.", "Description de chaque fichier du projet"),
]
for num, title in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{num}  "); r1.bold = True; r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1F,0x49,0x7D)
    p.add_run(title).font.size = Pt(11)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. ÉTAT ACTUEL
# ══════════════════════════════════════════════════════════════
add_h1(doc, "1. État actuel du projet et prochaines étapes")
add_sep(doc)

add_h2(doc, "1.1  Ce qui fonctionne (au 14 Avril 2026)")
ok = [
    ("Cluster KinD sre-cluster",      "1 control-plane + 2 workers — Running"),
    ("Prometheus + Alertmanager",      "Stack de monitoring — alertes actives"),
    ("Grafana",                        "Dashboards K8s — http://localhost:3000"),
    ("ArgoCD",                         "Déploiement GitOps — http://localhost:8080"),
    ("n8n",                            "Workflow actif et exécuté avec succès — http://localhost:5678"),
    ("SRE Agent (FastAPI)",            "Pod Running, /health OK, Mistral joignable"),
    ("Ollama + Mistral 7B",            "Modèle disponible sur Windows — port 11434"),
    ("crash-app (app de test)",        "CrashLoopBackOff actif — 300+ restarts — namespace demo"),
    ("Workflow n8n bout en bout",      "Tous les noeuds verts — flux complet validé"),
]
make_table(doc, ["Composant", "Statut"], ok, '23855A')

add_h2(doc, "1.2  Prochaines étapes")
next_steps = [
    ("Étape A", "Valider le flux automatique Alertmanager → n8n",
     "Vérifier que les alertes réelles de Prometheus déclenchent automatiquement le workflow sans action manuelle. "
     "Ouvrir http://localhost:9093, confirmer que PodRestartingTooMuch est en Firing, et observer les Executions n8n."),
    ("Étape B", "Activer AUTO_REMEDIATE pour tester la remédiation automatique",
     "Dans sre-agent-deploy.yaml, mettre AUTO_REMEDIATE=true dans le ConfigMap. "
     "Rebuild + redeploy. L'agent appliquera kubectl scale deployment/crash-app --replicas=0 -n demo automatiquement."),
    ("Étape C", "Tester la remédiation GitOps avec ArgoCD",
     "Configurer git_tools.py avec un token GitHub. Créer un dépôt gitops. "
     "L'agent créera une PR avec le manifest corrigé. ArgoCD synchronisera automatiquement."),
    ("Étape D", "Intégrer une vraie alerte applicative",
     "Déployer une vraie application (nginx, flask) avec une règle Prometheus personnalisée. "
     "Observer le flux complet de bout en bout sans intervention manuelle."),
    ("Étape E", "Documentation finale PFE",
     "Rédiger le rapport de stage : architecture, résultats, métriques (temps de détection, MTTR), "
     "captures d'écran du workflow n8n, comparaison avant/après."),
]
for num, title, desc in next_steps:
    add_h3(doc, f"{num} — {title}")
    add_body(doc, desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 2. COMMANDES MISE EN PLACE CLUSTER
# ══════════════════════════════════════════════════════════════
add_h1(doc, "2. Commandes de mise en place du cluster")
add_sep(doc)

add_h2(doc, "2.1  Installation et création du cluster KinD")
cmds_setup = [
    ("kind create cluster --config kind-cluster.yaml --name sre-cluster",
     "Crée le cluster Kubernetes local avec 1 control-plane et 2 workers selon la config YAML."),
    ("kind get clusters",
     "Liste tous les clusters KinD actifs sur la machine."),
    ("kubectl cluster-info --context kind-sre-cluster",
     "Affiche les URLs du cluster et confirme que kubectl pointe vers le bon contexte."),
    ("kubectl get nodes",
     "Affiche tous les noeuds du cluster avec leur statut (Ready/NotReady)."),
]
for cmd, desc in cmds_setup:
    add_h3(doc, "")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("Commande : "); r.bold = True; r.font.size = Pt(11)
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "2.2  Installation de la stack Prometheus/Grafana (kube-prometheus-stack)")
cmds_prom = [
    ("helm repo add prometheus-community https://prometheus-community.github.io/helm-charts",
     "Ajoute le dépôt Helm officiel Prometheus Community."),
    ("helm repo update",
     "Met à jour le cache local des charts Helm."),
    ("helm install monitoring prometheus-community/kube-prometheus-stack -n monitoring --create-namespace",
     "Installe Prometheus, Alertmanager, Grafana et les exporters dans le namespace monitoring."),
    ("kubectl get pods -n monitoring",
     "Vérifie que tous les pods de la stack monitoring sont en état Running."),
    ("kubectl apply -f pod-restart-alert.yaml",
     "Crée la règle PrometheusRule qui déclenche l'alerte PodRestartingTooMuch (>5 restarts en 10 min)."),
    ("kubectl apply -f alertmanager-config.yaml",
     "Applique la config Alertmanager qui route les alertes critical/warning vers le webhook n8n."),
]
for cmd, desc in cmds_prom:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "2.3  Installation d'ArgoCD")
cmds_argocd = [
    ("kubectl create namespace argocd",
     "Crée le namespace dédié à ArgoCD."),
    ("kubectl apply -n argocd -f https://raw.githubusercontent.com/argoproj/argo-cd/stable/manifests/install.yaml",
     "Installe ArgoCD (tous les composants : server, repo-server, application-controller, etc.)."),
    ("kubectl get pods -n argocd",
     "Vérifie que tous les pods ArgoCD sont Running."),
    ("kubectl -n argocd get secret argocd-initial-admin-secret -o jsonpath='{.data.password}' | base64 -d",
     "Récupère le mot de passe initial de l'admin ArgoCD (valide seulement au premier démarrage)."),
]
for cmd, desc in cmds_argocd:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "2.4  Déploiement de n8n et de l'app de test")
cmds_n8n = [
    ("kubectl apply -f n8n-deploy.yaml",
     "Déploie n8n (orchestrateur de workflow) avec un Service NodePort 30567 dans le namespace n8n."),
    ("kubectl apply -f demo-app.yaml",
     "Déploie crash-app (busybox exit 1) dans le namespace demo — application de test volontairement défectueuse."),
    ("kubectl get pods -n n8n",
     "Vérifie que le pod n8n est Running."),
    ("kubectl get pods -n demo",
     "Vérifie que crash-app est bien en CrashLoopBackOff (comportement attendu)."),
]
for cmd, desc in cmds_n8n:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 3. COMMANDES DÉPLOIEMENT COMPOSANTS
# ══════════════════════════════════════════════════════════════
add_h1(doc, "3. Commandes de déploiement des composants")
add_sep(doc)

add_h2(doc, "3.1  Build et déploiement du SRE Agent")
cmds_agent = [
    ("docker build -t sre-agent:latest .",
     "Construit l'image Docker du SRE Agent depuis sre-agent/Dockerfile. "
     "À exécuter depuis le dossier sre-agent/."),
    ("kind load docker-image sre-agent:latest --name sre-cluster",
     "Charge l'image locale dans les noeuds KinD (évite de passer par Docker Hub). "
     "Obligatoire car imagePullPolicy=Never dans le Deployment."),
    ("kubectl apply -f sre-agent-deploy.yaml",
     "Crée le namespace sre-agent, ServiceAccount, ClusterRole, RBAC, ConfigMap, Secret, PVC, Deployment et Service."),
    ("kubectl create secret generic sre-agent-kubeconfig --from-file=config=$HOME/.kube/config -n sre-agent --dry-run=client -o yaml | kubectl apply -f -",
     "Crée/met à jour le secret kubeconfig dans le cluster (utilisé en fallback si le SA token échoue)."),
    ("kubectl rollout status deployment/sre-agent -n sre-agent --timeout=120s",
     "Attend que le pod sre-agent soit Ready avant de continuer."),
    ("kubectl rollout restart deployment/sre-agent -n sre-agent",
     "Redémarre le pod sre-agent (utile après un rebuild d'image pour charger la nouvelle version)."),
]
for cmd, desc in cmds_agent:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "3.2  Script de déploiement automatisé")
add_code(doc, ".\\scripts\\deploy-sre-agent.ps1")
add_body(doc, "Utilité : Script PowerShell qui enchaîne automatiquement les 6 étapes : docker build → kind load → kubectl apply → kubeconfig secret → rollout status → port-forward. Équivalent à exécuter manuellement toutes les commandes ci-dessus.")

add_h2(doc, "3.3  Réinitialisation du mot de passe ArgoCD")
cmds_argocd_pw = [
    ("pip install bcrypt",
     "Installe le module Python bcrypt nécessaire pour hasher le nouveau mot de passe."),
    ('python -c "import bcrypt; h=bcrypt.hashpw(b\'admin123\', bcrypt.gensalt(10)); print(h.decode())"',
     "Génère un hash bcrypt du mot de passe admin123 (format attendu par ArgoCD)."),
    ('kubectl -n argocd patch secret argocd-secret -p \'{"data":{"admin.password":"<HASH_B64>","admin.passwordMtime":"<DATE_B64>"}}\'',
     "Injecte le nouveau hash dans le secret ArgoCD. Remplacer <HASH_B64> par le hash encodé en base64."),
    ("kubectl -n argocd rollout restart deployment argocd-server",
     "Redémarre argocd-server pour qu'il prenne en compte le nouveau mot de passe."),
]
for cmd, desc in cmds_argocd_pw:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 4. COMMANDES DEBUGGING
# ══════════════════════════════════════════════════════════════
add_h1(doc, "4. Commandes de correction et debugging")
add_sep(doc)

add_h2(doc, "4.1  Correction du conflit de dépendances Python (requirements.txt)")
add_body(doc, "Problème rencontré : langchain-core==0.3.15 incompatible avec langchain-ollama==0.2.1 (exige >=0.3.20).")
add_code(doc, "pip install langchain==0.3.7 langchain-core==0.3.27 langchain-ollama==0.2.1 langchain-community==0.3.7 --dry-run")
add_body(doc, "Utilité : Vérifie la compatibilité des versions sans installer (dry-run). Permet de tester les combinaisons de versions avant de modifier requirements.txt.")

add_h2(doc, "4.2  Vérification de connectivité réseau depuis n8n vers le SRE Agent")
cmds_net = [
    ("kubectl exec -n n8n <pod-n8n> -- wget -q -O- --timeout=10 http://sre-agent-service.sre-agent.svc.cluster.local:8000/health",
     "Teste la connectivité DNS et HTTP depuis le pod n8n vers le service SRE Agent via le DNS interne du cluster. Permet de confirmer que le service est joignable avant de configurer le workflow."),
    ("kubectl exec -n n8n <pod-n8n> -- wget -q -O- --timeout=10 http://10.96.61.178:8000/health",
     "Teste la connectivité via l'adresse ClusterIP directement (contourne le DNS). Utile pour isoler les problèmes DNS vs réseau."),
]
for cmd, desc in cmds_net:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "4.3  Vérification que kubectl fonctionne depuis le pod SRE Agent")
add_code(doc, "kubectl exec -n sre-agent <pod-sre-agent> -- sh -c 'TOKEN=$(cat /var/run/secrets/kubernetes.io/serviceaccount/token) && kubectl --server=https://kubernetes.default.svc --token=$TOKEN --certificate-authority=/var/run/secrets/kubernetes.io/serviceaccount/ca.crt get pods -n demo'")
add_body(doc, "Utilité : Vérifie que le ServiceAccount token in-cluster permet bien au pod d'exécuter des commandes kubectl. C'est le mécanisme utilisé par les tools LangChain pour investiguer le cluster. Si cette commande échoue, l'agent ne peut pas investiguer.")

add_h2(doc, "4.4  Correction du port Grafana (port déjà occupé)")
cmds_pf = [
    ("netstat -ano | findstr :3000",
     "Identifie quel processus occupe le port 3000 (PID affiché en dernière colonne)."),
    ("Stop-Process -Id <PID> -Force",
     "Tue le processus qui bloque le port (PowerShell). Remplacer <PID> par le PID trouvé."),
    ("Get-Process kubectl | Stop-Process -Force",
     "Tue TOUS les processus kubectl actifs (port-forwards inclus). Nettoie avant un redémarrage complet."),
]
for cmd, desc in cmds_pf:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "4.5  Correction de n8n (Exec format error — mauvaise architecture image)")
add_code(doc, "# Dans n8n-deploy.yaml : remplacer n8nio/n8n:latest par n8nio/n8n:1.82.3\nkubectl apply -f n8n-deploy.yaml\nkubectl rollout restart deployment/n8n -n n8n\nkubectl rollout status deployment/n8n -n n8n --timeout=120s")
add_body(doc, "Utilité : L'image n8nio/n8n:latest avait été tirée en ARM64 sur Docker Desktop Windows, causant 'Exec format error'. La version 1.82.3 forcée garantit l'architecture amd64 compatible avec KinD.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 5. COMMANDES TEST ET VALIDATION
# ══════════════════════════════════════════════════════════════
add_h1(doc, "5. Commandes de test et validation")
add_sep(doc)

add_h2(doc, "5.1  Vérification de l'état des services")
cmds_check = [
    ("curl http://localhost:8001/health",
     "Vérifie que l'API FastAPI du SRE Agent répond. Retourne {status: ok, service: sre-copilot}."),
    ("curl http://localhost:8001/status",
     "Vérifie que le SRE Agent peut joindre Ollama et que le modèle Mistral est disponible. Retourne {ollama: reachable, model_available: true}."),
    ("curl http://localhost:11434/api/tags",
     "Interroge Ollama directement pour lister les modèles disponibles. Confirme que mistral:latest est chargé."),
    ("kubectl get pods -A",
     "Vue globale de tous les pods du cluster. Identifie rapidement les pods en erreur."),
    ("kubectl get pods -A --field-selector=status.phase!=Running",
     "Filtre uniquement les pods qui ne sont PAS Running. Utile pour le diagnostic rapide."),
]
for cmd, desc in cmds_check:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "5.2  Test direct du SRE Agent (sans n8n)")
add_code(doc, "Invoke-RestMethod -Uri 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 600")
add_body(doc, "Utilité : Déclenche directement l'analyse d'une alerte CrashLoopBackOff fictive sans passer par n8n. L'agent LangChain + Mistral s'exécute, investigate le cluster avec kubectl, et retourne un JSON structuré avec classification, root_cause et runbook. Timeout 600s car Mistral 7B prend 3-5 minutes en local.")

add_h2(doc, "5.3  Test du workflow complet via n8n")
add_code(doc,
    "Invoke-RestMethod -Uri 'http://localhost:5678/webhook/sre-alert' `\n"
    "  -Method POST `\n"
    "  -ContentType 'application/json' `\n"
    "  -TimeoutSec 600 `\n"
    "  -Body '{\n"
    '    "receiver":"n8n-webhook",\n'
    '    "status":"firing",\n'
    '    "alerts":[{\n'
    '      "status":"firing",\n'
    '      "labels":{\n'
    '        "alertname":"PodRestartingTooMuch",\n'
    '        "namespace":"demo",\n'
    '        "pod":"crash-app",\n'
    '        "severity":"critical"\n'
    '      },\n'
    '      "annotations":{\n'
    '        "summary":"Pod crash-app CrashLoopBackOff"\n'
    '      }\n'
    "    }]\n"
    "  }'"
)
add_body(doc, "Utilité : Simule exactement ce qu'Alertmanager envoie à n8n. Déclenche le workflow complet : Webhook → Is Firing? → Call SRE Agent → Parse Response → Build Report → Webhook Response. Permet de valider le flux de bout en bout.")

add_h2(doc, "5.4  Vérification des alertes Prometheus et Alertmanager")
cmds_alerts = [
    ("# Ouvrir dans le navigateur :\n# http://localhost:9090/alerts   → Prometheus (liste des alertes firing)\n# http://localhost:9093          → Alertmanager (alertes routées)",
     "Vérifie visuellement que l'alerte PodRestartingTooMuch est bien en état Firing dans Prometheus, et qu'elle est bien routée vers n8n dans Alertmanager."),
    ("kubectl get prometheusrule -n monitoring",
     "Liste les règles d'alertes Prometheus actives dans le cluster."),
    ("kubectl describe prometheusrule -n monitoring pod-restart-alert",
     "Affiche le détail de la règle PodRestartingTooMuch : expression PromQL, seuils, labels."),
]
for cmd, desc in cmds_alerts:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 6. COMMANDES MAINTENANCE QUOTIDIENNE
# ══════════════════════════════════════════════════════════════
add_h1(doc, "6. Commandes de maintenance quotidienne")
add_sep(doc)

add_h2(doc, "6.1  Démarrage de l'environnement")
add_code(doc, "cd C:\\Users\\PC\\Desktop\\SRE-Copilot\n.\\scripts\\start-sre-copilot.ps1")
add_body(doc, "Utilité : Script de démarrage complet. Tue les anciens port-forwards, vérifie l'état du cluster, lance 5 fenêtres PowerShell avec port-forward persistant (auto-relance si coupure) pour Grafana, Prometheus, Alertmanager, n8n et ArgoCD.")

add_h2(doc, "6.2  Arrêt propre des port-forwards")
add_code(doc, "Get-Process kubectl | Stop-Process -Force")
add_body(doc, "Utilité : Tue tous les processus kubectl actifs (port-forwards). À exécuter avant de relancer start-sre-copilot.ps1 pour éviter les conflits de ports.")

add_h2(doc, "6.3  Consultation des logs")
cmds_logs = [
    ("kubectl logs -n sre-agent -l app=sre-agent -f --tail=50",
     "Affiche et suit les logs du SRE Agent en temps réel. Montre les appels Mistral, les outils kubectl exécutés, et les erreurs éventuelles."),
    ("kubectl logs -n n8n -l app=n8n -f --tail=50",
     "Logs de n8n. Utile pour diagnostiquer les problèmes de connexion ou d'exécution du workflow."),
    ("kubectl logs -n monitoring -l app.kubernetes.io/name=alertmanager --tail=30",
     "Logs Alertmanager. Confirme que les alertes sont bien envoyées au webhook n8n."),
    ("kubectl logs -n demo -l app=crash-app --previous",
     "Logs du dernier crash de l'app de test (container terminé). Montre le exit code 1."),
]
for cmd, desc in cmds_logs:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "6.4  Surveillance des pods")
cmds_watch = [
    ("kubectl get pods -n sre-agent -w",
     "Surveille les changements de statut du pod sre-agent en temps réel (-w = watch)."),
    ("kubectl describe pod -n sre-agent <pod-name>",
     "Affiche tous les détails d'un pod : events, volumes montés, variables d'environnement, probes. Essentiel pour diagnostiquer un pod qui ne démarre pas."),
    ("kubectl top pods -n sre-agent",
     "Affiche la consommation CPU et mémoire du pod SRE Agent en temps réel. Nécessite metrics-server."),
]
for cmd, desc in cmds_watch:
    add_code(doc, cmd)
    add_body(doc, f"Utilité : {desc}")

add_h2(doc, "6.5  Rebuild rapide après modification du code")
add_code(doc,
    "cd C:\\Users\\PC\\Desktop\\SRE-Copilot\\sre-agent\n"
    "docker build -t sre-agent:latest . --quiet\n"
    "kind load docker-image sre-agent:latest --name sre-cluster\n"
    "kubectl rollout restart deployment/sre-agent -n sre-agent\n"
    "kubectl rollout status deployment/sre-agent -n sre-agent --timeout=60s"
)
add_body(doc, "Utilité : Séquence complète de rebuild à exécuter à chaque modification du code Python. Le --quiet supprime les logs verbeux du build pour aller plus vite.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 7. DESCRIPTION DES FICHIERS
# ══════════════════════════════════════════════════════════════
add_h1(doc, "7. Description de chaque fichier du projet")
add_sep(doc)

add_h2(doc, "7.1  Fichiers de configuration Kubernetes (racine)")
files_k8s = [
    ("kind-cluster.yaml",
     "Définition du cluster KinD",
     "Déclare la topologie du cluster : 1 noeud control-plane + 2 noeuds worker. "
     "Utilisé une seule fois avec 'kind create cluster --config kind-cluster.yaml'."),
    ("alertmanager-config.yaml",
     "Configuration du routage Alertmanager",
     "Définit les routes d'alertes : les alertes de sévérité critical ou warning sont envoyées "
     "au webhook http://n8n-service.n8n.svc.cluster.local:5678/webhook/sre-alert. "
     "L'alerte Watchdog est ignorée (receiver: null)."),
    ("pod-restart-alert.yaml",
     "Règle PrometheusRule pour CrashLoopBackOff",
     "Crée l'alerte PodRestartingTooMuch : se déclenche quand un pod redémarre plus de 5 fois "
     "en 10 minutes (expression PromQL : increase(kube_pod_container_status_restarts_total[10m]) > 5). "
     "Sévérité : critical."),
    ("demo-app.yaml",
     "Application de test volontairement défectueuse",
     "Déploie un pod busybox avec la commande 'exit 1' dans le namespace demo. "
     "Génère en permanence un CrashLoopBackOff pour déclencher l'alerte. "
     "Cible principale de l'investigation de l'agent SRE."),
    ("n8n-deploy.yaml",
     "Déploiement de n8n",
     "Crée le Deployment n8n (image n8nio/n8n:1.82.3) et un Service NodePort 30567 "
     "dans le namespace n8n. Version fixée à 1.82.3 pour éviter les problèmes d'architecture ARM64."),
    ("sre-agent-deploy.yaml",
     "Déploiement complet du SRE Agent",
     "Fichier principal de déploiement K8s. Contient en un seul fichier : Namespace, "
     "ServiceAccount (sre-agent-sa), ClusterRole + ClusterRoleBinding (droits kubectl), "
     "Secret (tokens GitHub/ArgoCD), ConfigMap (OLLAMA_URL, model, AUTO_REMEDIATE), "
     "PVC (stockage gitops), Deployment (image sre-agent:latest) et Service NodePort 30800."),
]
for fname, title, desc in files_k8s:
    add_h3(doc, f"{fname}  —  {title}")
    add_body(doc, desc)

add_h2(doc, "7.2  Code source du SRE Agent (sre-agent/)")
files_agent = [
    ("main.py",
     "Point d'entrée FastAPI",
     "Définit l'application FastAPI avec 5 endpoints : GET /health (santé), "
     "GET /status (Ollama + modèle), POST /alert (webhook Alertmanager), "
     "POST /analyze (test unitaire), POST /test-crash-loop (endpoint de test intégré). "
     "Configure le lifespan et les logs uvicorn."),
    ("agent.py",
     "Agent LangChain ReAct + Ollama",
     "Coeur de l'intelligence artificielle. Construit un agent LangChain ReAct avec "
     "Mistral 7B via Ollama. Contient le prompt système SRE, la logique d'investigation "
     "(max 6 itérations), le parsing du Final Answer JSON, et la gestion des erreurs. "
     "La fonction run_sre_investigation() est async et utilise run_in_executor() pour "
     "ne pas bloquer l'event loop FastAPI."),
    ("models.py",
     "Modèles Pydantic",
     "Définit les structures de données : AlertLabels, AlertAnnotations, Alert, "
     "AlertManagerPayload (payload webhook Alertmanager), RemediationStep et AgentResponse "
     "(réponse structurée retournée par l'agent avec classification, root_cause, runbook, etc.)."),
    ("Dockerfile",
     "Image Docker du SRE Agent",
     "Construit l'image depuis python:3.12-slim. Installe kubectl (binaire Linux amd64 "
     "depuis dl.k8s.io), copie requirements.txt, installe les dépendances Python, "
     "copie le code source. Point d'entrée : uvicorn main:app --host 0.0.0.0 --port 8000."),
    ("requirements.txt",
     "Dépendances Python",
     "Liste exacte des packages avec versions compatibles : fastapi==0.115.5, "
     "uvicorn[standard]==0.32.1, pydantic==2.9.2, httpx==0.27.2, langchain==0.3.7, "
     "langchain-core==0.3.27 (version clé : >=0.3.20 requis par langchain-ollama), "
     "langchain-ollama==0.2.1, langchain-community==0.3.7, requests==2.32.3."),
    ("tools/kubectl_tools.py",
     "Outils LangChain pour kubectl",
     "8 outils LangChain (@tool) que l'agent peut appeler : get_pod_logs, get_pod_status, "
     "describe_pod, get_deployment_status, get_namespace_events, restart_deployment, "
     "scale_deployment, find_crashing_pods. Utilise le ServiceAccount token in-cluster "
     "(auto-détecté via /var/run/secrets/) pour les appels kubectl. "
     "Inclut _clean() qui supprime les guillemets que Mistral ajoute parfois aux inputs."),
    ("tools/git_tools.py",
     "Outils LangChain pour GitOps",
     "Outils pour créer des Pull Requests GitOps : clone du dépôt, modification du manifest "
     "YAML défectueux, commit et push sur une branche, création de PR via l'API GitHub."),
    ("tools/argocd_tools.py",
     "Outils LangChain pour ArgoCD",
     "Outils pour interagir avec ArgoCD : déclencher une synchronisation (sync), "
     "rollback d'une application, vérification du statut de santé d'une Application ArgoCD."),
    ("tools/__init__.py",
     "Registre des outils",
     "Agrège tous les outils des 3 modules (kubectl, git, argocd) dans une liste "
     "ALL_TOOLS importée par agent.py pour les passer au AgentExecutor LangChain."),
]
for fname, title, desc in files_agent:
    add_h3(doc, f"{fname}  —  {title}")
    add_body(doc, desc)

add_h2(doc, "7.3  Workflows n8n (n8n-workflows/)")
add_h3(doc, "sre-alert-workflow.json  —  Workflow principal n8n")
add_body(doc,
    "Workflow JSON importable dans n8n. Contient 8 noeuds connectés :\n"
    "1. Alertmanager Webhook — écoute POST /webhook/sre-alert\n"
    "2. Is Firing? — filtre les alertes en état 'firing' uniquement\n"
    "3. Call SRE Agent — POST vers l'agent avec timeout 360s\n"
    "4. Parse Agent Response — extrait classification, runbook, prUrl depuis la réponse JSON\n"
    "5. Has PR? — branche selon si une PR GitOps a été créée\n"
    "6. Build Incident Report — formate un rapport lisible\n"
    "7. Audit Log — journalise l'incident dans les logs n8n\n"
    "8. Webhook Response — répond à Alertmanager avec le résultat"
)

add_h2(doc, "7.4  Manifests ArgoCD (argocd/)")
add_h3(doc, "sre-agent-app.yaml  —  Application ArgoCD")
add_body(doc,
    "Définit une Application ArgoCD qui surveille le dossier gitops-repo/apps/sre-agent/ "
    "et synchronise automatiquement le Deployment du SRE Agent. "
    "Permet la mise à jour du SRE Agent via GitOps sans kubectl manuel."
)

add_h2(doc, "7.5  GitOps repo (gitops-repo/)")
files_gitops = [
    ("apps/demo/crash-app.yaml",
     "Manifest de crash-app géré par ArgoCD",
     "Version GitOps du déploiement de crash-app. ArgoCD surveille ce fichier. "
     "Quand l'agent SRE crée une PR pour corriger la commande (exit 1 → sleep 3600), "
     "ArgoCD synchronise ce fichier corrigé vers le cluster automatiquement."),
    ("apps/sre-agent/deployment.yaml",
     "Manifest du SRE Agent géré par ArgoCD",
     "Déploiement GitOps du SRE Agent. Permet les mises à jour de l'agent via PR GitOps "
     "plutôt que par kubectl direct."),
]
for fname, title, desc in files_gitops:
    add_h3(doc, f"{fname}  —  {title}")
    add_body(doc, desc)

add_h2(doc, "7.6  Scripts PowerShell (scripts/)")
files_scripts = [
    ("start-sre-copilot.ps1",
     "Script de démarrage de l'environnement",
     "Tue les anciens kubectl port-forwards, vérifie Docker/KinD/pods, "
     "puis lance chaque service via pf-persistent.ps1 dans des fenêtres séparées : "
     "Grafana (3000), Prometheus (9090), Alertmanager (9093), n8n (5678), ArgoCD (8080), "
     "SRE Agent (8000 si déployé). Ouvre les navigateurs automatiquement."),
    ("pf-persistent.ps1",
     "Port-forward avec auto-relance",
     "Script de port-forward robuste. Boucle while($true) : détecte si le port est occupé, "
     "libère le port si nécessaire, lance kubectl port-forward, relance automatiquement "
     "si la connexion coupe (pod restart, timeout réseau, etc.). "
     "Résout le problème récurrent de Grafana/n8n inaccessibles après un redémarrage de pod."),
    ("deploy-sre-agent.ps1",
     "Script de déploiement du SRE Agent",
     "Automatise les 6 étapes : docker build → kind load → kubectl apply → "
     "kubeconfig secret → rollout status → port-forward. "
     "Arrête en cas d'erreur à chaque étape (if LASTEXITCODE -ne 0)."),
    ("generate_doc.py",
     "Générateur du guide de test (Document 1)",
     "Script Python (python-docx) qui génère SRE-Copilot-Guide-Test.docx : "
     "architecture, description du workflow, étapes de test, URLs et credentials."),
    ("generate_doc2.py",
     "Générateur de la référence commandes/fichiers (ce document)",
     "Script Python (python-docx) qui génère ce document : état du projet, "
     "toutes les commandes utilisées avec leur utilité, et description de chaque fichier."),
]
for fname, title, desc in files_scripts:
    add_h3(doc, f"{fname}  —  {title}")
    add_body(doc, desc)

# Footer
doc.add_paragraph()
add_sep(doc)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("SRE Copilot — Référence Commandes & Fichiers  |  Stage PFE PwC DigiTech  |  Avril 2026")
r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x99,0x99,0x99); r.italic = True

# Save
out = r"C:\Users\PC\Desktop\SRE-Copilot\SRE-Copilot-Commandes-Fichiers.docx"
doc.save(out)
print(f"OK: {out}")
