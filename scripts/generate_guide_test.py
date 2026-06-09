"""
Generates SRE-Copilot-Guide-Test.docx
Professional Word document — Complete A-Z test guide with colors and architecture.
Run: python scripts/generate_guide_test.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Color palette ──────────────────────────────────────────────────────────────
C_NAVY     = RGBColor(0x1F, 0x38, 0x64)
C_BLUE     = RGBColor(0x2E, 0x75, 0xB6)
C_TEAL     = RGBColor(0x00, 0x70, 0x7F)
C_GREEN    = RGBColor(0x37, 0x56, 0x23)
C_ORANGE   = RGBColor(0xC5, 0x50, 0x0C)
C_RED      = RGBColor(0x9C, 0x00, 0x00)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY_TXT = RGBColor(0x40, 0x40, 0x40)
C_GRAY_LT  = RGBColor(0xD6, 0xDC, 0xE4)
C_CODE_TXT = RGBColor(0x1F, 0x1F, 0x1F)

HEX_NAVY      = "1F3864"
HEX_BLUE      = "2E75B6"
HEX_TEAL      = "00707F"
HEX_TABLE_HDR = "1F3864"
HEX_TABLE_ALT = "DEEAF1"
HEX_CODE_BG   = "F2F2F2"
HEX_GREEN_BG  = "E2EFDA"
HEX_ORANGE_BG = "FFF2CC"
HEX_RED_BG    = "FCE4D6"
HEX_BLUE_BG   = "D9E1F2"

PHASE_COLORS = {
    "A": ("DÉMARRAGE",       "1F3864", C_WHITE),
    "B": ("N8N",             "2E75B6", C_WHITE),
    "C": ("ALERTMANAGER",    "00707F", C_WHITE),
    "D": ("SIMULATION A–Z",  "375623", C_WHITE),
    "E": ("VÉRIFICATION",    "843C0C", C_WHITE),
    "F": ("RÉINITIALISATION","595959", C_WHITE),
    "G": ("ROBUSTESSE",      "4472C4", C_WHITE),
    "H": ("OBSERVABILITÉ",   "1F3864", C_WHITE),
    "I": ("KYVERNO",         "7030A0", C_WHITE),
    "J": ("MÉTRIQUES GRAFANA","C5500C", C_WHITE),
}


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_cell_borders(table, color="D6DCE4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, color=None, size=10, font="Calibri",
              align=WD_ALIGN_PARAGRAPH.LEFT, italic=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return run


# ── Typography helpers ─────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.bold = True
    run.font.color.rgb = C_NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), HEX_NAVY)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = C_BLUE
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = C_TEAL
    return p


def body(doc, text, indent=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = color if color else C_GRAY_TXT
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = C_GRAY_TXT
    return p


def code_block(doc, lines, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after  = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"
        lr.font.size = Pt(9)
        lr.bold = True
        lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = C_CODE_TXT
        set_para_bg(p, HEX_CODE_BG)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(4)


def info_box(doc, text, bg_hex, icon="ℹ"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(f"  {icon}  {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = C_GRAY_TXT
    set_para_bg(p, bg_hex)
    return p


def phase_header(doc, letter, name, description):
    color_hex, txt_color = PHASE_COLORS[letter][1], PHASE_COLORS[letter][2]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f"  PHASE {letter} — {name.upper()}")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.bold = True
    run.font.color.rgb = txt_color
    set_para_bg(p, color_hex)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    run2 = p2.add_run(f"  {description}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = C_WHITE
    set_para_bg(p2, color_hex)


def expected_result(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"  ✔  Résultat attendu : {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = C_GREEN
    set_para_bg(p, HEX_GREEN_BG)


def warning_box(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(f"  ⚠  {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = C_ORANGE
    set_para_bg(p, HEX_ORANGE_BG)


def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell_text(hdr.cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(hdr.cells[i], HEX_TABLE_HDR)
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else HEX_TABLE_ALT
        for c_idx, val in enumerate(row):
            cell_text(tr.cells[c_idx], val, size=9)
            set_cell_bg(tr.cells[c_idx], bg)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    set_cell_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── Cover page ─────────────────────────────────────────────────────────────────

def add_cover(doc):
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    set_para_bg(bar, HEX_NAVY)
    r = bar.add_run("  SRE COPILOT")
    r.font.name = "Calibri"; r.font.size = Pt(28); r.bold = True; r.font.color.rgb = C_WHITE

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(0)
    set_para_bg(sub, HEX_TEAL)
    sr = sub.add_run("  GUIDE DE TEST COMPLET — SIMULATION A à Z  |  MISSION 3")
    sr.font.name = "Calibri"; sr.font.size = Pt(13); sr.bold = True; sr.font.color.rgb = C_WHITE

    doc.add_paragraph()

    info_rows = [
        ("Projet",      "PFE IFS 013/26 — PwC DigiTech Cloud Native Operations"),
        ("Objectif",    "Valider le pipeline complet : Alerte → Investigation IA → Remédiation GitOps"),
        ("Stack testée","KinD + Prometheus + Alertmanager + n8n + SRE Agent + Ollama/Mistral"),
        ("Durée test",  "Démarrage : ~3 min | Simulation complète : ~5-8 min par cycle"),
        ("Date",        "2026-05-27"),
    ]
    t = doc.add_table(rows=len(info_rows), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(info_rows):
        bg = HEX_TABLE_ALT if i % 2 == 0 else "FFFFFF"
        cell_text(t.rows[i].cells[0], k, bold=True, color=C_NAVY, size=10)
        cell_text(t.rows[i].cells[1], v, size=10)
        set_cell_bg(t.rows[i].cells[0], bg)
        set_cell_bg(t.rows[i].cells[1], bg)
    for row in t.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.2)
    set_cell_borders(t)

    doc.add_paragraph()

    # Pipeline visual
    heading2(doc, "Vue d'ensemble du Pipeline Testé")
    arrow_t = doc.add_table(rows=2, cols=9)
    arrow_t.alignment = WD_TABLE_ALIGNMENT.CENTER
    components = [
        ("Prometheus\n:9090",    HEX_NAVY),
        ("→",                    "FFFFFF"),
        ("Alertmanager\n:9093",  HEX_TEAL),
        ("→",                    "FFFFFF"),
        ("n8n\n:5678",           "375623"),
        ("→",                    "FFFFFF"),
        ("SRE Agent\n:8001",     HEX_BLUE),
        ("→",                    "FFFFFF"),
        ("ArgoCD\n:8080",        "843C0C"),
    ]
    descs = [
        "Détecte les\nconditions d'alerte", "",
        "Route vers\nn8n (namespace=demo)", "",
        "Filtre + déduplique\n+ appelle l'agent", "",
        "Phase 1+2+3\nInvestigation + Fix", "",
        "Sync GitOps\naprès remédiation",
    ]
    for i, (lbl, bg) in enumerate(components):
        c = arrow_t.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        if bg == "FFFFFF":
            run = p.add_run(lbl)
            run.font.name = "Calibri"; run.font.size = Pt(14); run.bold = True
            run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
        else:
            set_cell_bg(c, bg)
            run = p.add_run(lbl)
            run.font.name = "Calibri"; run.font.size = Pt(9); run.bold = True
            run.font.color.rgb = C_WHITE

        c2 = arrow_t.rows[1].cells[i]
        p2 = c2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(2)
        r2 = p2.add_run(descs[i])
        r2.font.name = "Calibri"; r2.font.size = Pt(8)
        r2.font.color.rgb = C_GRAY_TXT

    for row in arrow_t.rows:
        for i, c in enumerate(row.cells):
            c.width = Inches(0.3 if components[i][1] == "FFFFFF" else 1.1)

    doc.add_page_break()


# ── Prerequisites ──────────────────────────────────────────────────────────────

def add_prerequisites(doc):
    heading1(doc, "Prérequis & Informations Générales")

    heading2(doc, "Outils Requis")
    make_table(doc,
        ["Outil", "Version min.", "Commande de vérification", "Remarque"],
        [
            ("Docker Desktop", "4.x",  "docker info",        "Doit être EN COURS D'EXÉCUTION avant tout"),
            ("KinD",           "0.20", "kind version",       "Kubernetes in Docker — cluster sre-cluster"),
            ("kubectl",        "1.28", "kubectl version --client", "Contexte : kind-sre-cluster"),
            ("Helm",           "3.x",  "helm version",       "Pour kube-prometheus-stack"),
            ("Ollama",         "0.1+", "ollama list",        "Doit tourner nativement sur Windows"),
            ("Mistral 7B",     "—",    "ollama list",        "Doit apparaître dans la liste des modèles"),
            ("Python",         "3.12", "python --version",   "Utilisé par le SRE Agent (Dockerfile)"),
        ],
        col_widths=[1.3, 0.9, 2.2, 2.8]
    )

    heading2(doc, "URLs & Accès")
    make_table(doc,
        ["Service", "URL", "Credentials", "Note"],
        [
            ("SRE Agent",    "http://localhost:8001/docs",   "Aucune",              "Swagger UI — tester les endpoints"),
            ("Grafana",      "http://localhost:3000",        "admin / prom-operator","Dashboard Kubernetes"),
            ("Prometheus",   "http://localhost:9090",        "Aucune",              "Alertes, règles, PromQL"),
            ("Alertmanager", "http://localhost:9093",        "Aucune",              "Receivers, routes"),
            ("n8n",          "http://localhost:5678",        "Aucune",              "Workflow, historique exécutions"),
            ("ArgoCD",       "https://localhost:8080",       "admin / Sre@Copilot2026","GitOps sync"),
            ("Ollama",       "http://localhost:11434",       "Aucune",              "API modèles LLM"),
        ],
        col_widths=[1.2, 1.8, 1.8, 2.8]
    )

    heading2(doc, "Durées Estimées")
    make_table(doc,
        ["Opération", "Durée", "Dépendance"],
        [
            ("Démarrage (cluster existant)",  "~3 min",    "Port-forwards + health checks"),
            ("Démarrage (cluster recréé)",     "~15-20 min","Helm install kube-prometheus-stack"),
            ("crash-app → CrashLoopBackOff",  "~1-2 min",  ">3 restarts + 30s for: dans la rule"),
            ("Investigation complète (LLM)",   "~30-60s",   "Dépend de la charge CPU Mistral"),
            ("Cycle de simulation complet",    "~5-8 min",  "Déployer + attendre alerte + agent"),
        ],
        col_widths=[2.5, 1.0, 4.2]
    )
    doc.add_page_break()


# ── Phase A ────────────────────────────────────────────────────────────────────

def add_phase_a(doc):
    phase_header(doc, "A", "Démarrage de l'Environnement",
                 "Vérifier que le cluster KinD, tous les services et Ollama sont opérationnels.")

    heading3(doc, "Étape A.1 — Lancer le script de démarrage automatique")
    code_block(doc, [
        "# Depuis le répertoire SRE-Copilot/",
        ".\\scripts\\start-sre-copilot.ps1",
    ], label="COMMANDE")
    info_box(doc, "Le script gère : vérification cluster → déploiement si vide → n8n Docker → port-forwards → health checks → ouverture navigateurs.", HEX_BLUE_BG)
    expected_result(doc, "5 lignes vertes   OK   pour SRE Agent, Grafana, Prometheus, Alertmanager, n8n")

    heading3(doc, "Étape A.2 — Vérification manuelle de l'état")
    code_block(doc, [
        "kubectl get nodes",
        "# Attendu : 3 lignes Ready (control-plane + worker + worker2)",
        "",
        "kubectl get pods -n monitoring    # kube-prometheus-stack",
        "kubectl get pods -n argocd        # ArgoCD",
        "kubectl get pods -n sre-agent     # SRE Agent",
        "# Tous Running/Ready",
    ], label="VÉRIFICATION CLUSTER")
    warning_box(doc, "ArgoCD applicationset-controller peut être en CrashLoopBackOff — NON BLOQUANT pour les tests Mission 3.")

    heading3(doc, "Étape A.3 — Vérifier la connectivité Ollama depuis le pod")
    code_block(doc, [
        "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
        "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
        "",
        "kubectl exec -n sre-agent $pod -- python3 -c \"",
        "import urllib.request, json",
        "req = urllib.request.Request('http://host.docker.internal:11434/api/tags')",
        "with urllib.request.urlopen(req, timeout=5) as r:",
        "    data = json.loads(r.read())",
        "    print('Ollama OK. Models:', [m['name'] for m in data.get('models',[])])",
        "\"",
    ], label="TEST OLLAMA DEPUIS POD")
    expected_result(doc, "Ollama OK. Models: ['mistral:latest']")
    info_box(doc, "Pourquoi host.docker.internal ? Les pods KinD s'exécutent dans des conteneurs Docker. Cette adresse DNS spéciale résout vers l'IP du host Windows — là où tourne Ollama nativement.", HEX_ORANGE_BG, "ℹ")


# ── Phase B ────────────────────────────────────────────────────────────────────

def add_phase_b(doc):
    phase_header(doc, "B", "Vérification du Workflow n8n",
                 "Confirmer que le workflow est actif, bien configuré et répond au webhook.")

    heading3(doc, "Étape B.1 — Vérifier le workflow dans l'UI n8n")
    body(doc, "Ouvrir http://localhost:5678 → Workflows → Vérifier que le workflow SRE Alert Pipeline existe et est ACTIF (toggle vert).")
    body(doc, "Le workflow doit contenir exactement 5 nœuds dans cet ordre :")
    for n in ["1. webhook-trigger", "2. filter-dedup", "3. call-sre-agent", "4. parse-response", "5. build-report"]:
        bullet(doc, n)

    heading3(doc, "Étape B.2 — Importer le workflow si absent")
    code_block(doc, [
        "# Dans n8n : + New Workflow → Import from File",
        "# Sélectionner : n8n-workflows/sre-alert-workflow.json",
        "# Activer le toggle en haut à droite",
    ], label="IMPORT WORKFLOW")

    heading3(doc, "Étape B.3 — Vérifier la configuration critique du webhook")
    body(doc, "Cliquer sur le nœud webhook-trigger et vérifier :")
    make_table(doc,
        ["Paramètre", "Valeur attendue", "Pourquoi c'est critique"],
        [
            ("HTTP Method",    "POST",             "Alertmanager envoie toujours en POST."),
            ("Path",           "sre-alert",        "URL complète : /webhook/sre-alert"),
            ("Response Mode",  "Respond immediately (onReceived)",
             "OBLIGATOIRE : répond HTTP 200 avant traitement. Alertmanager a un timeout de 30s. Sans onReceived → retry infini."),
        ],
        col_widths=[1.5, 2.0, 4.2]
    )

    heading3(doc, "Étape B.4 — Tester le webhook directement")
    code_block(doc, [
        "$testPayload = @{",
        "    version = '4'; status = 'firing'; receiver = 'n8n-webhook'",
        "    alerts  = @(@{",
        "        status      = 'firing'",
        "        labels      = @{ alertname='TestAlert'; namespace='demo'; pod='test-pod'; severity='critical' }",
        "        annotations = @{ summary='Test webhook'; description='Ping test' }",
        "        startsAt    = (Get-Date -Format 'o')",
        "    })",
        "} | ConvertTo-Json -Depth 5",
        "",
        "$resp = Invoke-WebRequest -Uri 'http://localhost:5678/webhook/sre-alert' `",
        "    -Method POST -Body $testPayload -ContentType 'application/json'",
        "Write-Host \"Status: $($resp.StatusCode)\"",
    ], label="TEST WEBHOOK")
    expected_result(doc, "Status: 200 + 1 exécution visible dans n8n → Executions")


# ── Phase C ────────────────────────────────────────────────────────────────────

def add_phase_c(doc):
    phase_header(doc, "C", "Vérification de l'Alertmanager",
                 "Confirmer que les alertes du namespace demo sont bien routées vers n8n.")

    heading3(doc, "Étape C.1 — Vérifier les receivers")
    code_block(doc, [
        "# Via API Alertmanager",
        "Invoke-WebRequest 'http://localhost:9093/api/v2/receivers' | `",
        "    Select-Object -ExpandProperty Content",
        "# Doit montrer 2 receivers : 'null' et 'n8n-webhook'",
        "",
        "# Vérifier les routes dans l'UI",
        "Start-Process 'http://localhost:9093'  # Menu : Status → Config",
    ], label="RECEIVERS ALERTMANAGER")
    expected_result(doc, "Receiver n8n-webhook avec URL http://host.docker.internal:5678/webhook/sre-alert")

    heading3(doc, "Étape C.2 — Vérifier la règle PrometheusRule")
    code_block(doc, [
        "# Lister les PrometheusRules",
        "kubectl get prometheusrules -A",
        "# Attendu : pod-restart-alert dans le namespace demo ou monitoring",
        "",
        "# Vérifier dans l'UI Prometheus",
        "Start-Process 'http://localhost:9090/alerts'",
        "# Chercher 'PodRestartingTooMuch' — état 'inactive' tant que crash-app n'est pas déployée",
    ], label="VÉRIFICATION RÈGLE")
    expected_result(doc, "PodRestartingTooMuch apparaît dans Prometheus → Alerts (état inactive ou pending)")

    heading3(doc, "Étape C.3 — Tester l'envoi direct à Alertmanager")
    info_box(doc, "Alertmanager v1 supprimée depuis v0.27.0 — utiliser /api/v2/alerts avec un TABLEAU JSON. La version here-string (@\"...\"@) évite tout problème d'échappement de quotes.", HEX_ORANGE_BG)
    code_block(doc, [
        "$now = (Get-Date -Format 'o')",
        "$end = (Get-Date).AddHours(1).ToString('o')",
        "",
        "$body = @\"",
        "[{",
        "  \"labels\": {",
        "    \"alertname\": \"PodRestartingTooMuch\",",
        "    \"namespace\": \"demo\",",
        "    \"pod\":       \"crash-app\",",
        "    \"severity\":  \"critical\"",
        "  },",
        "  \"annotations\": {",
        "    \"summary\":     \"Test\",",
        "    \"description\": \"Simulation\"",
        "  },",
        "  \"startsAt\": \"$now\",",
        "  \"endsAt\":   \"$end\"",
        "}]",
        "\"@",
        "",
        "Invoke-WebRequest -Uri 'http://localhost:9093/api/v2/alerts' `",
        "    -Method POST -Body $body -ContentType 'application/json'",
        "# Alertmanager route → n8n → SRE Agent (attendre 30-60s)",
    ], label="TEST ALERTMANAGER (API v2)")
    info_box(doc, "Le `\"@` de fermeture du here-string DOIT être en colonne 0 (pas d'indentation), sinon PowerShell renvoie une erreur de parsing.", HEX_BLUE_BG)


# ── Phase D ────────────────────────────────────────────────────────────────────

def add_phase_d(doc):
    phase_header(doc, "D", "Simulation Complète Bout en Bout",
                 "Tester le pipeline complet : pod crash → alerte → investigation IA → remédiation GitOps.")

    info_box(doc, "DEUX MÉTHODES disponibles. Option 1 (test rapide) recommandée pour la démo. Option 2 (flux complet) pour valider tout le pipeline de bout en bout.", HEX_BLUE_BG)

    heading2(doc, "Option 1 — Test Rapide via /test-crash-loop (Recommandé)")
    body(doc, "Cette méthode appelle directement le SRE Agent sans passer par Alertmanager/n8n. Idéale pour valider rapidement la Phase 3.")

    heading3(doc, "D.1 — Déployer crash-app")
    code_block(doc, [
        "kubectl apply -f demo-app.yaml",
        "",
        "# Surveiller l'état jusqu'à CrashLoopBackOff (~1-2 minutes)",
        "kubectl get pods -n demo -w",
        "# Ctrl+C quand vous voyez : STATUS = CrashLoopBackOff",
    ], label="DÉPLOYER CRASH-APP")
    expected_result(doc, "crash-app-XXXXX   0/1   CrashLoopBackOff   3   90s")

    heading3(doc, "D.2 — Déclencher l'investigation complète")
    code_block(doc, [
        "Write-Host 'Démarrage investigation (30-60 secondes)...'",
        "$startTime = Get-Date",
        "",
        "$resp = Invoke-WebRequest -Uri 'http://localhost:8001/test-crash-loop' `",
        "    -Method POST -TimeoutSec 120",
        "",
        "$elapsed = (Get-Date) - $startTime",
        "Write-Host \"Durée : $($elapsed.TotalSeconds.ToString('F1'))s\"",
        "",
        "# Afficher le résultat formaté",
        "$result = $resp.Content | ConvertFrom-Json",
        "$result | ConvertTo-Json -Depth 5",
    ], label="DÉCLENCHER INVESTIGATION")
    expected_result(doc, "classification=CrashLoopBackOff | auto_remediation_applied=true | git_pr_url=local:branch/fix/...")

    heading2(doc, "Option 2 — Flux Complet via Alertmanager (Pipeline Entier)")
    body(doc, "Teste tout le pipeline : Prometheus détecte → Alertmanager route → n8n filtre → SRE Agent répond.")

    heading3(doc, "D.3 — Attendre l'alerte automatique Prometheus")
    code_block(doc, [
        "# L'alerte se déclenche quand : restarts > 3 pendant 30 secondes",
        "# Vérifier dans Prometheus http://localhost:9090/alerts",
        "# Rafraîchir jusqu'à voir PodRestartingTooMuch en ROUGE (firing)",
        "",
        "# Vérification API",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/alerts' | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | `",
        "    ForEach-Object { $_.data.alerts } | `",
        "    Where-Object { $_.labels.alertname -eq 'PodRestartingTooMuch' }",
    ], label="ATTENDRE ALERTE PROMETHEUS")

    heading3(doc, "D.4 — Surveiller l'exécution dans n8n")
    body(doc, "Dans http://localhost:5678 → Workflow → onglet Executions :")
    make_table(doc,
        ["État", "Signification", "Action"],
        [
            ("Running",   "Investigation en cours (LLM tourne)",     "Attendre, ne pas relancer"),
            ("Succeeded", "Pipeline complet sans erreur",             "Cliquer pour voir les résultats de chaque nœud"),
            ("Failed",    "Erreur dans un nœud",                     "Cliquer → identifier le nœud rouge → lire l'erreur"),
            ("(aucune)",  "Alerte filtrée par filter-dedup (dedup active)", "Attendre 10 min ou relancer le workflow"),
        ],
        col_widths=[1.2, 2.8, 3.7]
    )

    heading3(doc, "D.5 — Surveiller les logs SRE Agent en temps réel")
    code_block(doc, [
        "kubectl logs -f deployment/sre-agent -n sre-agent",
        "",
        "# Vous devez voir ces lignes dans cet ordre :",
        "# [Phase 1] kubectl collection — demo/crash-app",
        "# [Phase 1] Collected 4 sections",
        "# [Phase 2] LLM JSON-mode analysis",
        "# [Phase 2] Classification: CrashLoopBackOff | Fix: scale_to_zero",
        "# [Phase 3] AUTO_REMEDIATE=true — applying fix_type=scale_to_zero",
        "# [Phase 3] Complete: {'auto_remediation_applied': True, 'argocd_sync_triggered': True}",
    ], label="LOGS SRE AGENT")
    expected_result(doc, "Les 3 phases apparaissent dans les logs sans erreur")


# ── Phase E ────────────────────────────────────────────────────────────────────

def add_phase_e(doc):
    phase_header(doc, "E", "Vérification des Résultats",
                 "Confirmer que la remédiation a bien été appliquée, le commit Git créé, et ArgoCD appelé.")

    heading3(doc, "Étape E.1 — Vérifier la remédiation sur le pod")
    code_block(doc, [
        "# Le déploiement crash-app doit être à 0 replicas",
        "kubectl get deployment crash-app -n demo",
        "# Attendu : READY 0/0, UP-TO-DATE 0",
        "",
        "kubectl get pods -n demo",
        "# Attendu : 'No resources found in demo namespace.'",
    ], label="VÉRIFICATION REMÉDIATION")
    expected_result(doc, "crash-app   0/0   0   0   — deployment scaled to 0 replicas")

    heading3(doc, "Étape E.2 — Vérifier les commits dans le gitops-repo")
    code_block(doc, [
        "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
        "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
        "",
        "# Historique des commits",
        "kubectl exec -n sre-agent $pod -- git -C /gitops-repo log --oneline",
        "# Attendu : fix(demo): auto-remediate CrashLoopBackOff on crash-app",
        "",
        "# Voir le contenu du dernier commit",
        "kubectl exec -n sre-agent $pod -- git -C /gitops-repo show --stat HEAD",
        "",
        "# Lister les fichiers d'incidents créés",
        "kubectl exec -n sre-agent $pod -- ls /gitops-repo/incidents/demo/",
    ], label="VÉRIFICATION GIT")
    expected_result(doc, "1+ commits visibles | fichier incidents/demo/crash-app-TIMESTAMP.yaml créé")

    heading3(doc, "Étape E.3 — Vérifier l'appel API ArgoCD")
    code_block(doc, [
        "$b64 = kubectl get secret sre-agent-secrets -n sre-agent -o jsonpath='{.data.ARGOCD_TOKEN}'",
        "$token = [System.Text.Encoding]::UTF8.GetString([System.Convert]::FromBase64String($b64))",
        "",
        "$headers = @{ Authorization = \"Bearer $token\" }",
        "Invoke-WebRequest -Uri 'https://localhost:8080/api/v1/applications' `",
        "    -Headers $headers -SkipCertificateCheck | Select-Object -ExpandProperty Content",
        "# HTTP 200 confirme que le token permanent est valide",
    ], label="TEST TOKEN ARGOCD")

    heading3(doc, "Étape E.4 — Résultat JSON Complet Attendu")
    code_block(doc, [
        "{",
        "  \"alert_name\": \"PodRestartingTooMuch\",",
        "  \"namespace\": \"demo\",",
        "  \"severity\": \"critical\",",
        "  \"classification\": \"CrashLoopBackOff\",",
        "  \"root_cause\": \"Container crashes immediately after start (exit code 1). Kubernetes restarts in exponential-backoff loop.\",",
        "  \"auto_remediation_applied\": true,",
        "  \"remediation_details\": \"Scaled demo/crash-app to 0 replicas. Output: deployment.apps/crash-app scaled\",",
        "  \"git_pr_url\": \"local:branch/fix/crash-app-crashloopbackoff-20260527-XXXXXX\",",
        "  \"argocd_sync_triggered\": true,",
        "  \"runbook\": [",
        "    {\"order\": 1, \"action\": \"Check previous logs\",  \"command\": \"kubectl logs crash-app-XXX -n demo --previous\"},",
        "    {\"order\": 2, \"action\": \"Describe pod\",          \"command\": \"kubectl describe pod crash-app-XXX -n demo\"},",
        "    {\"order\": 3, \"action\": \"Fix or scale to zero\",  \"command\": \"kubectl scale deployment crash-app --replicas=0 -n demo\"}",
        "  ]",
        "}",
    ], label="RÉPONSE JSON ATTENDUE")


# ── Phase F ────────────────────────────────────────────────────────────────────

def add_phase_f(doc):
    phase_header(doc, "F", "Réinitialisation pour une Nouvelle Simulation",
                 "Remettre l'environnement dans l'état initial pour rejouer un test.")

    code_block(doc, [
        "# Remettre crash-app à 1 replica (relance le crash loop)",
        "kubectl scale deployment crash-app --replicas=1 -n demo",
        "",
        "# Vérifier que crash-app redémarre",
        "kubectl get pods -n demo -w",
        "",
        "# Réinitialisation complète propre",
        "kubectl delete -f demo-app.yaml",
        "Start-Sleep -Seconds 5",
        "kubectl apply -f demo-app.yaml",
    ], label="RÉINITIALISATION")

    warning_box(doc, "DÉDUPLICATION n8n active 10 min : si vous relancez immédiatement le même test, n8n l'ignorera silencieusement (TTL du cache). Attendez 10 minutes ou redémarrez le workflow n8n.")

    info_box(doc, "Pour tester plusieurs fois rapidement : utiliser /test-crash-loop directement (contourne la déduplication n8n).", HEX_BLUE_BG)


# ── Phase G ────────────────────────────────────────────────────────────────────

def add_phase_g(doc):
    phase_header(doc, "G", "Tests de Robustesse",
                 "Valider le comportement de l'agent face aux cas limites et aux pannes.")

    heading3(doc, "Test G.1 — LLM hors ligne (fallback keyword)")
    body(doc, "Arrêter Ollama (Gestionnaire des tâches → terminer ollama.exe), puis déclencher une investigation.")
    code_block(doc, [
        "# Déclencher avec Ollama arrêté",
        "$resp = Invoke-WebRequest -Uri 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 120",
        "$result = $resp.Content | ConvertFrom-Json",
        "Write-Host \"Classification : $($result.classification)\"",
        "# Attendu : CrashLoopBackOff (via keyword fallback, pas LLM)",
        "",
        "# Vérifier dans investigation_log",
        "$result.investigation_log | Where-Object { $_ -match 'fallback' }",
        "# Attendu : [llm:fallback] CrashLoopBackOff",
    ], label="TEST FALLBACK KEYWORD")
    expected_result(doc, "classification=CrashLoopBackOff même sans LLM — fallback keyword actif")
    info_box(doc, "Le fallback keyword analyse le texte brut des données kubectl (logs, describe, events) pour détecter 'crashloopbackoff', 'oomkilled', 'imagepullbackoff'. Garantit une classification même en cas de panne Ollama.", HEX_BLUE_BG)

    heading3(doc, "Test G.2 — Double alerte (déduplication n8n)")
    code_block(doc, [
        "$payload = @{ version='4'; status='firing'; receiver='n8n-webhook'",
        "    alerts=@(@{ status='firing'",
        "        labels=@{alertname='PodRestartingTooMuch';namespace='demo';pod='crash-app';severity='critical'}",
        "        annotations=@{summary='Test dedup'} ; startsAt=(Get-Date -Format 'o')",
        "    })",
        "} | ConvertTo-Json -Depth 5",
        "",
        "# Envoyer 2 fois à 5 secondes d'intervalle",
        "Invoke-WebRequest 'http://localhost:5678/webhook/sre-alert' -Method POST -Body $payload -ContentType 'application/json'",
        "Start-Sleep -Seconds 5",
        "Invoke-WebRequest 'http://localhost:5678/webhook/sre-alert' -Method POST -Body $payload -ContentType 'application/json'",
    ], label="TEST DÉDUPLICATION")
    expected_result(doc, "Seulement 1 exécution n8n complète — la 2ème est bloquée par le cache TTL 10min")

    heading3(doc, "Test G.3 — Alerte hors namespace demo (doit être ignorée)")
    code_block(doc, [
        "$payload = @{ version='4'; status='firing'; receiver='n8n-webhook'",
        "    alerts=@(@{ status='firing'",
        "        labels=@{alertname='PodRestartingTooMuch';namespace='monitoring';pod='test';severity='critical'}",
        "        annotations=@{summary='Hors demo'} ; startsAt=(Get-Date -Format 'o')",
        "    })",
        "} | ConvertTo-Json -Depth 5",
        "",
        "Invoke-WebRequest 'http://localhost:5678/webhook/sre-alert' -Method POST -Body $payload -ContentType 'application/json'",
    ], label="TEST FILTRE NAMESPACE")
    expected_result(doc, "Exécution n8n Succeeded mais avec résultat vide — alerte ignorée par filter-dedup")


# ── Phase H ────────────────────────────────────────────────────────────────────

def add_phase_h(doc):
    phase_header(doc, "H", "Observabilité & Dashboards",
                 "Exploiter Grafana, Prometheus et n8n pour observer le comportement du système.")

    heading3(doc, "Grafana — http://localhost:3000 (admin / prom-operator)")
    make_table(doc,
        ["Dashboard", "Ce qu'il montre", "Navigation"],
        [
            ("Kubernetes / Pods",          "CPU, mémoire, restarts par pod et namespace", "Dashboards → Browse → Kubernetes"),
            ("Kubernetes / Nodes",         "Ressources des nœuds KinD (CPU, disque)",     "Dashboards → Browse → Node Exporter"),
            ("Alertmanager Overview",      "Alertes actives, silences, receivers",         "Dashboards → Browse → Alertmanager"),
        ],
        col_widths=[2.0, 2.8, 2.9]
    )
    body(doc, "Requêtes PromQL utiles dans Explore :")
    code_block(doc, [
        "# Restarts du pod crash-app en temps réel",
        "kube_pod_container_status_restarts_total{namespace='demo'}",
        "",
        "# Taux de restart sur 5 minutes",
        "rate(kube_pod_container_status_restarts_total{namespace='demo'}[5m])",
        "",
        "# Alertes actives namespace demo",
        "ALERTS{namespace='demo', alertstate='firing'}",
    ], label="PROMQL UTILES")

    heading3(doc, "n8n — Historique des Exécutions")
    body(doc, "Chaque exécution contient le résultat de chaque nœud. Pour voir le rapport complet :")
    body(doc, "Workflows → Cliquer sur le workflow → Executions (icône horloge) → Cliquer sur une exécution Succeeded → Cliquer sur le nœud build-report pour voir le rapport Markdown.")

    heading3(doc, "Prometheus — Vérifier l'état des alertes")
    code_block(doc, [
        "# Toutes les alertes actives",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/alerts' | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | ConvertTo-Json -Depth 5",
        "",
        "# État des règles PrometheusRule",
        "Start-Process 'http://localhost:9090/rules'",
    ], label="ÉTAT PROMETHEUS")


# ── Phase I (Kyverno - Mission 4) ─────────────────────────────────────────────

def add_phase_i(doc):
    phase_header(doc, "I", "Sécurité Kyverno (Mission 4)",
                 "Vérifier que les 4 politiques de sécurité détectent les violations sur les workloads.")

    heading3(doc, "Étape I.1 — Vérifier l'installation Kyverno")
    code_block(doc, [
        "# Les 4 controllers doivent être Running",
        "kubectl get pods -n kyverno",
        "",
        "# Attendu : 4 lignes",
        "#   kyverno-admission-controller-XXX    1/1 Running",
        "#   kyverno-background-controller-XXX   1/1 Running",
        "#   kyverno-cleanup-controller-XXX      1/1 Running",
        "#   kyverno-reports-controller-XXX      1/1 Running",
    ], label="VÉRIFICATION KYVERNO")
    expected_result(doc, "4 pods Kyverno en état Running")

    heading3(doc, "Étape I.2 — Vérifier les 4 politiques déployées")
    code_block(doc, [
        "kubectl get clusterpolicies",
        "",
        "# Attendu :",
        "# disallow-latest-tag              true   true   True   Ready",
        "# disallow-privileged-containers   true   true   True   Ready",
        "# require-labels                   true   true   True   Ready",
        "# require-resource-limits          true   true   True   Ready",
    ], label="VÉRIFICATION POLITIQUES")
    expected_result(doc, "4 ClusterPolicies avec READY=True")

    heading3(doc, "Étape I.3 — Déployer les 4 pods de test (violations volontaires)")
    body(doc, "Chaque pod viole une politique différente. En mode Audit, ils sont créés mais génèrent des PolicyReports avec FAIL.")
    code_block(doc, [
        "kubectl apply -f kyverno/test-violations/",
        "",
        "# Liste les pods créés",
        "kubectl get pods -n demo",
        "",
        "# Attendu :",
        "# violation-latest        ContainerCreating  (image:latest)",
        "# violation-no-labels     ContainerCreating  (pas de labels app/owner)",
        "# violation-no-limits     ContainerCreating  (pas de resources.limits)",
        "# violation-privileged    ContainerCreating  (privileged:true)",
    ], label="DÉPLOIEMENT VIOLATIONS")

    heading3(doc, "Étape I.4 — Lire les PolicyReports générés")
    code_block(doc, [
        "# Attendre 5s pour que reports-controller traite",
        "Start-Sleep -Seconds 5",
        "",
        "# Vue d'ensemble",
        "kubectl get policyreport -n demo",
        "",
        "# Attendu : chaque pod a 1 FAIL + 4 PASS",
        "# (un pod ne viole qu'une politique, mais 4 sont évaluées)",
    ], label="LECTURE REPORTS")
    expected_result(doc, "4 PolicyReports (un par pod test) + 2 pour crash-app/Deployment et ReplicaSet")

    heading3(doc, "Étape I.5 — Compter les violations par politique")
    code_block(doc, [
        "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty items | `",
        "    Select-Object -ExpandProperty results -ErrorAction SilentlyContinue | `",
        "    Where-Object { $_.result -eq 'fail' } | `",
        "    Group-Object -Property policy | `",
        "    Select-Object Name, Count | Format-Table -AutoSize",
        "",
        "# Attendu (avec crash-app présente) :",
        "# disallow-latest-tag                  3   (test + crash-app Deploy + ReplicaSet)",
        "# disallow-privileged-containers       1",
        "# require-labels                       3",
        "# require-resource-limits              3",
    ], label="STATS VIOLATIONS")

    heading3(doc, "Étape I.6 — Détail des violations sur crash-app")
    body(doc, "crash-app viole 3 politiques sans modification — montre l'apport de Kyverno sur les workloads existants.")
    code_block(doc, [
        "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty items | `",
        "    Where-Object { $_.scope.name -eq 'crash-app' -and $_.scope.kind -eq 'Deployment' } | `",
        "    Select-Object -ExpandProperty results | `",
        "    Where-Object { $_.result -eq 'fail' } | `",
        "    Select-Object policy, message | Format-List",
    ], label="DÉTAIL CRASH-APP")
    expected_result(doc, "3 violations : disallow-latest-tag, require-labels (owner manquant), require-resource-limits (cpu manquant)")

    heading3(doc, "Étape I.7 — (Optionnel) Tester le mode Enforce")
    body(doc, "En mode Enforce, Kyverno BLOQUE la création des ressources non-conformes.")
    code_block(doc, [
        "# Éditer kyverno/policies/01-disallow-privileged.yaml",
        "# Changer : validationFailureAction: Audit  →  Enforce",
        "kubectl apply -f kyverno/policies/01-disallow-privileged.yaml",
        "",
        "# Tenter de recréer un pod privilegied",
        "kubectl delete pod violation-privileged -n demo",
        "kubectl apply -f kyverno/test-violations/violation-privileged.yaml",
        "",
        "# Attendu (mode Enforce) :",
        "# Error from server: admission webhook denied the request:",
        "# disallow-privileged-containers: validation error: Les conteneurs privilégiés sont interdits.",
    ], label="MODE ENFORCE")
    warning_box(doc, "Le mode Enforce bloque AUSSI les pods créés par les controllers (Deployment, DaemonSet…). À utiliser uniquement quand le cluster est totalement conforme.")

    heading3(doc, "Étape I.8 — Nettoyage (optionnel)")
    code_block(doc, [
        "# Supprimer les pods de test",
        "kubectl delete -f kyverno/test-violations/",
        "",
        "# Les PolicyReports correspondants disparaissent automatiquement",
        "kubectl get policyreport -n demo",
    ], label="CLEANUP")


# ── Phase J (Mission 5 - Métriques + Grafana) ─────────────────────────────────

def add_phase_j(doc):
    phase_header(doc, "J", "Métriques Prometheus & Dashboard Grafana (Mission 5)",
                 "Vérifier que l'agent expose ses KPIs, que Prometheus les scrape, et que le dashboard Grafana est auto-importé.")

    heading3(doc, "Étape J.1 — Tester l'endpoint /metrics du SRE Agent")
    code_block(doc, [
        "$r = Invoke-WebRequest 'http://localhost:8001/metrics' -TimeoutSec 5",
        "$r.StatusCode    # Attendu : 200",
        "$r.Headers.'Content-Type'  # Attendu : text/plain; version=0.0.4",
        "",
        "# Voir les premières métriques sre_*",
        "$r.Content -split \"`n\" | Select-String '^sre_' | Select-Object -First 10",
    ], label="TEST /metrics")
    expected_result(doc, "HTTP 200 + format Prometheus avec sre_agent_info, sre_alerts_total, sre_investigation_duration_seconds_bucket, etc.")

    heading3(doc, "Étape J.2 — Vérifier le ServiceMonitor")
    code_block(doc, [
        "kubectl get servicemonitor sre-agent -n monitoring",
        "# Attendu : 1 ligne, age > 1 minute",
        "",
        "# Vérifier les labels (release: monitoring est obligatoire)",
        "kubectl get servicemonitor sre-agent -n monitoring -o jsonpath='{.metadata.labels}'",
    ], label="SERVICEMONITOR")
    expected_result(doc, "ServiceMonitor visible avec label release=monitoring")

    heading3(doc, "Étape J.3 — Vérifier que Prometheus scrape l'agent")
    code_block(doc, [
        "$t = (Invoke-WebRequest 'http://localhost:9090/api/v1/targets' -TimeoutSec 5).Content | ConvertFrom-Json",
        "$t.data.activeTargets | Where-Object { $_.labels.job -eq 'sre-agent-service' } | `",
        "    Select-Object scrapeUrl, health, lastScrape, lastError | Format-List",
    ], label="VERIFICATION PROMETHEUS")
    expected_result(doc, "health=up | lastError vide | lastScrape récent (< 30s)")

    heading3(doc, "Étape J.4 — Requêtes PromQL de base")
    code_block(doc, [
        "# Métadonnées de l'agent",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/query?query=sre_agent_info' | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty data",
        "",
        "# Compteur d'alertes par classification",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/query?query=sum%20by%20(classification)%20(sre_alerts_total)' | `",
        "    Select-Object -ExpandProperty Content",
    ], label="PROMQL")
    expected_result(doc, "Au moins 1 série retournée par requête (après quelques investigations)")

    heading3(doc, "Étape J.5 — Vérifier que le dashboard Grafana est importé")
    code_block(doc, [
        "$basic = [Convert]::ToBase64String([Text.Encoding]::ASCII.GetBytes('admin:prom-operator'))",
        "$h = @{ Authorization = \"Basic $basic\" }",
        "Invoke-WebRequest 'http://localhost:3000/api/search?query=SRE' -Headers $h | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | `",
        "    Select-Object title, uid, url",
        "",
        "# Ouvrir directement le dashboard",
        "Start-Process 'http://localhost:3000/d/sre-copilot/sre-copilot-kpis'",
    ], label="DASHBOARD GRAFANA")
    expected_result(doc, "1 dashboard 'SRE Copilot — KPIs' avec uid=sre-copilot et 9 panels visibles")

    heading3(doc, "Étape J.6 — Déclencher une investigation pour peupler les métriques")
    body(doc, "Le dashboard s'enrichit au fur et à mesure des investigations. Lance-en au moins 2-3 pour avoir des données significatives.")
    code_block(doc, [
        "# Re-deployer crash-app si scaled à 0",
        "kubectl scale deployment crash-app -n demo --replicas=1",
        "",
        "# Lancer 1 investigation et attendre (~3 min)",
        "Invoke-WebRequest 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 360",
        "",
        "# Rafraîchir le dashboard Grafana — les panels se remplissent",
    ], label="PEUPLER LES METRIQUES")
    expected_result(doc, "Panels Stat (Total/Success rate/MTTR), donut Classifications et bar Remédiations affichent des valeurs")


# ── Checklist ─────────────────────────────────────────────────────────────────

def add_checklist(doc):
    doc.add_page_break()
    heading1(doc, "Checklist de Validation Complète — Missions 1 à 5")

    items = [
        ("A.1",  "Script start-sre-copilot.ps1 terminé sans erreur (5 lignes OK vertes)"),
        ("A.2",  "3 nœuds KinD Ready (control-plane + 2 workers)"),
        ("A.2",  "Tous les pods monitoring/argocd/sre-agent en Running"),
        ("A.3",  "Ollama accessible depuis le pod : 'Ollama OK. Models: [mistral:latest]'"),
        ("B.1",  "Workflow n8n actif avec 5 nœuds dans le bon ordre"),
        ("B.3",  "responseMode = 'Respond immediately (onReceived)' sur webhook-trigger"),
        ("B.4",  "Test direct webhook n8n → HTTP 200 reçu"),
        ("C.1",  "Receiver n8n-webhook visible dans Alertmanager"),
        ("C.2",  "Règle PodRestartingTooMuch chargée dans Prometheus"),
        ("D.1",  "crash-app déployée et en état CrashLoopBackOff"),
        ("D.2",  "Investigation /test-crash-loop terminée en < 90 secondes"),
        ("E.1",  "crash-app scaled to 0 replicas (Phase 3 remédiation)"),
        ("E.2",  "Commit git créé dans /gitops-repo/incidents/demo/"),
        ("E.3",  "Token ArgoCD permanent → API ArgoCD répond HTTP 200"),
        ("E.4",  "JSON : classification=CrashLoopBackOff, auto_remediation_applied=true"),
        ("G.1",  "Fallback keyword actif : classification OK même sans Ollama"),
        ("G.2",  "Déduplication n8n : 2ème alerte ignorée dans les 10 min"),
        ("G.3",  "Alerte hors namespace demo ignorée par le filtre n8n"),
        ("I.1",  "4 controllers Kyverno en Running"),
        ("I.2",  "4 ClusterPolicies avec READY=True"),
        ("I.3",  "4 pods de test créés (mode Audit ne bloque pas)"),
        ("I.4",  "PolicyReports générés avec 1 FAIL par pod test"),
        ("I.5",  "Stats par politique : 1-3 violations selon la politique"),
        ("I.6",  "crash-app détectée avec 3 violations (latest, owner, cpu)"),
        ("J.1",  "Endpoint /metrics répond HTTP 200 en format Prometheus"),
        ("J.2",  "ServiceMonitor sre-agent présent avec label release=monitoring"),
        ("J.3",  "Prometheus target health=up sur job sre-agent-service"),
        ("J.4",  "Requête PromQL sre_agent_info retourne au moins 1 série"),
        ("J.5",  "Dashboard 'SRE Copilot — KPIs' visible dans Grafana (uid: sre-copilot)"),
        ("J.6",  "Après 1 investigation, les 9 panels Grafana affichent des données"),
    ]

    t = doc.add_table(rows=1 + len(items), cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(["#", "Phase", "Validation", "Statut"]):
        cell_text(t.rows[0].cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(t.rows[0].cells[i], HEX_TABLE_HDR)
    for r, (phase, desc) in enumerate(items):
        bg = "FFFFFF" if r % 2 == 0 else HEX_TABLE_ALT
        cell_text(t.rows[r+1].cells[0], str(r+1), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(t.rows[r+1].cells[1], phase, bold=True, color=C_NAVY, size=9)
        cell_text(t.rows[r+1].cells[2], desc, size=9)
        cell_text(t.rows[r+1].cells[3], "[ ]", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        for ci in range(4):
            set_cell_bg(t.rows[r+1].cells[ci], bg)
    for row in t.rows:
        row.cells[0].width = Inches(0.3)
        row.cells[1].width = Inches(0.5)
        row.cells[2].width = Inches(5.5)
        row.cells[3].width = Inches(0.5)
    set_cell_borders(t)


# ── Troubleshooting ────────────────────────────────────────────────────────────

def add_troubleshooting(doc):
    doc.add_paragraph()
    heading1(doc, "Dépannage Rapide")

    make_table(doc,
        ["Symptôme", "Cause probable", "Solution"],
        [
            ("classification: Unknown dans la réponse",
             "Ollama hors ligne ou timeout LLM",
             "Vérifier ollama list. Redémarrer : ollama serve. Tester http://localhost:8001/status"),
            ("n8n timeout sur call-sre-agent",
             "Port-forward SRE Agent sur 8001 non actif",
             "kubectl port-forward --address 0.0.0.0 svc/sre-agent-service -n sre-agent 8001:8000"),
            ("Alertmanager ne route pas vers n8n",
             "Secret alertmanager-monitoring-kube-prometheus-alertmanager mal configuré",
             "kubectl apply -f alertmanager-config.yaml puis redéployer le secret"),
            ("argocd_sync_triggered: false",
             "Application demo-app non enregistrée dans ArgoCD",
             "NON BLOQUANT pour Mission 3. Le token est valide — il faut créer l'App ArgoCD."),
            ("n8n exécution bloquée — aucune execution visible",
             "Déduplication active (TTL 10min sur alerte déjà traitée)",
             "Attendre 10 min OU utiliser /test-crash-loop qui bypass n8n."),
            ("Port 8001 ECONNREFUSED dans n8n",
             "port-forward sans --address 0.0.0.0 (non accessible depuis Docker)",
             "Tuer le port-forward existant et relancer avec --address 0.0.0.0"),
            ("CrashLoopBackOff ne déclenche pas d'alerte",
             "Seuil > 3 restarts non atteint ou délai for: 30s pas encore écoulé",
             "Attendre 2-3 minutes. Vérifier dans Prometheus → Alerts l'état pending."),
            ("Token ArgoCD expire (toutes les 24h)",
             "Ancien token de session admin (sub: admin:login avec exp)",
             "Relancer .\\scripts\\generate-argocd-token.ps1 pour générer un token permanent."),
            ("SRE Agent pod en Pending",
             "Image sre-agent:latest non chargée dans KinD",
             "docker build -t sre-agent:latest sre-agent/ puis kind load docker-image sre-agent:latest --name sre-cluster"),
            ("kubectl ERROR: certificate signed by unknown authority dans pod",
             "Certificat ArgoCD auto-signé",
             "ARGOCD_INSECURE=true dans le ConfigMap — déjà configuré. Vérifier que le pod a redémarré."),
            ("Kyverno: 'admission webhook denied the request'",
             "Politique en mode Enforce qui bloque le déploiement",
             "Identifier la politique en cause dans le message. Soit corriger le pod, soit repasser la politique en Audit."),
            ("kubectl get policyreport renvoie une liste vide",
             "Pas de pods créés depuis l'installation Kyverno OU background scan pas encore lancé",
             "Attendre 1-2 min après création d'un pod. Le reports-controller scanne périodiquement."),
            ("Politique 'disallow-latest-tag' loggue FAIL sur des images busybox/nginx sans tag",
             "Image sans tag = :latest implicite pour Kubernetes",
             "Spécifier un tag explicite : image: busybox:1.36 au lieu de image: busybox"),
        ],
        col_widths=[1.9, 2.0, 3.8]
    )


def main():
    out_path = os.path.join(os.path.dirname(__file__), "..", "SRE-Copilot-Guide-Test.docx")
    out_path = os.path.abspath(out_path)

    doc = Document()

    # A4 layout
    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    doc.styles['Normal'].font.name = "Calibri"
    doc.styles['Normal'].font.size = Pt(10)

    add_cover(doc)
    add_prerequisites(doc)
    add_phase_a(doc)
    add_phase_b(doc)
    add_phase_c(doc)
    add_phase_d(doc)
    add_phase_e(doc)
    add_phase_f(doc)
    add_phase_g(doc)
    add_phase_h(doc)
    add_phase_i(doc)
    add_phase_j(doc)
    add_checklist(doc)
    add_troubleshooting(doc)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("PFE IFS 013/26 — PwC DigiTech Cloud Native Operations | SRE Copilot | Guide de Test Mission 3 | 2026-05-27")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xD6, 0xDC, 0xE4)
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
