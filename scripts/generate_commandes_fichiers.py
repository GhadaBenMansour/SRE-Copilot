"""
Generates SRE-Copilot-Commandes-Fichiers.docx
Professional Word document with colors, tables, and architecture.
Run: python scripts/generate_commandes_fichiers.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Color palette ──────────────────────────────────────────────────────────────
C_NAVY      = RGBColor(0x1F, 0x38, 0x64)   # dark navy — main headings
C_BLUE      = RGBColor(0x2E, 0x75, 0xB6)   # medium blue — sub-headings
C_TEAL      = RGBColor(0x00, 0x70, 0x7F)   # teal — accent
C_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)   # white
C_GREEN     = RGBColor(0x37, 0x56, 0x23)   # dark green — complete
C_ORANGE    = RGBColor(0xC5, 0x50, 0x0C)   # orange — warning
C_GRAY_TXT  = RGBColor(0x40, 0x40, 0x40)   # dark gray — body text
C_GRAY_LITE = RGBColor(0xD6, 0xDC, 0xE4)   # light gray — dividers
C_CODE_TXT  = RGBColor(0x1F, 0x1F, 0x1F)   # near-black — code text

HEX_NAVY       = "1F3864"
HEX_BLUE       = "2E75B6"
HEX_TEAL       = "00707F"
HEX_TABLE_HDR  = "1F3864"
HEX_TABLE_ALT  = "DEEAF1"
HEX_TABLE_EVEN = "FFFFFF"
HEX_CODE_BG    = "F2F2F2"
HEX_GREEN_BG   = "E2EFDA"
HEX_ORANGE_BG  = "FCE4D6"
HEX_COVER_BG   = "1F3864"


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_cell_borders(table):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'D6DCE4')
                tcBorders.append(border)
            tcPr.append(tcBorders)


def no_space_before(para):
    para.paragraph_format.space_before = Pt(0)


def cell_text(cell, text, bold=False, color=None, size=10, font="Calibri", align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


# ── Section heading helpers ────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = C_NAVY
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), HEX_NAVY)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.bold = True
    run.font.color.rgb = C_BLUE
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = C_TEAL
    return p


def body_text(doc, text, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = C_GRAY_TXT
    return p


def code_block(doc, lines, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"
        lr.font.size = Pt(9)
        lr.bold = True
        lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = C_CODE_TXT
        set_para_bg(p, HEX_CODE_BG)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(4)


def badge(doc, text, bg_hex, text_color=C_WHITE):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"  {text}  ")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.bold = True
    run.font.color.rgb = text_color
    set_para_bg(p, bg_hex)
    return p


# ── Table builders ─────────────────────────────────────────────────────────────

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell_text(hdr.cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(hdr.cells[i], HEX_TABLE_HDR)

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        bg = HEX_TABLE_EVEN if r_idx % 2 == 0 else HEX_TABLE_ALT
        for c_idx, val in enumerate(row):
            cell_text(tr.cells[c_idx], val, size=9)
            set_cell_bg(tr.cells[c_idx], bg)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)

    set_cell_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


# ── Cover page ─────────────────────────────────────────────────────────────────

def add_cover(doc):
    # Top colored bar
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after = Pt(0)
    set_para_bg(bar, HEX_COVER_BG)
    bar_run = bar.add_run("  SRE COPILOT")
    bar_run.font.name = "Calibri"
    bar_run.font.size = Pt(28)
    bar_run.bold = True
    bar_run.font.color.rgb = C_WHITE

    subtitle_bar = doc.add_paragraph()
    subtitle_bar.paragraph_format.space_before = Pt(0)
    subtitle_bar.paragraph_format.space_after = Pt(0)
    set_para_bg(subtitle_bar, HEX_NAVY)
    sub_run = subtitle_bar.add_run("  RÉFÉRENCE COMMANDES & ARCHITECTURE")
    sub_run.font.name = "Calibri"
    sub_run.font.size = Pt(14)
    sub_run.bold = True
    sub_run.font.color.rgb = C_WHITE

    doc.add_paragraph()

    info_rows = [
        ("Projet",       "PFE IFS 013/26 — PwC DigiTech Cloud Native Operations"),
        ("Objectif",     "Agent IA SRE : remédiation GitOps pilotée par Ollama/Mistral"),
        ("Stack",        "KinD + Prometheus + ArgoCD + n8n + FastAPI + Ollama"),
        ("Environnement","Local Windows — Docker Desktop + KinD"),
        ("Mise à jour",  "2026-05-27"),
        ("Statut",       "Mission 3 Complete — Missions 4 & 5 à venir"),
    ]
    t = doc.add_table(rows=len(info_rows), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(info_rows):
        bg = HEX_TABLE_ALT if i % 2 == 0 else HEX_TABLE_EVEN
        cell_text(t.rows[i].cells[0], k, bold=True, color=C_NAVY, size=10)
        cell_text(t.rows[i].cells[1], v, size=10)
        set_cell_bg(t.rows[i].cells[0], bg)
        set_cell_bg(t.rows[i].cells[1], bg)
    for row in t.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(5.0)
    set_cell_borders(t)
    doc.add_page_break()


# ── Section 1: Architecture ────────────────────────────────────────────────────

def add_architecture(doc):
    heading1(doc, "1. Architecture du Pipeline")
    body_text(doc, "Le pipeline SRE Copilot suit un flux déterministe en 3 phases. Chaque composant a un rôle précis et une raison d'être documentée ci-dessous.")
    doc.add_paragraph()

    # Pipeline flow as a table
    t = doc.add_table(rows=1, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    boxes = [
        ("ALERTMANAGER\n:9093", HEX_NAVY, "Détecte les\nalertes Prometheus"),
        ("N8N WEBHOOK\n:5678", HEX_TEAL, "Filtre + dédup\n(namespace=demo)"),
        ("SRE AGENT\n:8001", HEX_BLUE, "Pipeline\n3 phases"),
        ("OLLAMA\nMistral 7B", "375623", "LLM local\nJSON mode"),
        ("ARGOCD\n:8080", "843C0C", "GitOps\nSync"),
    ]
    arrows = ["→", "→", "→", "→"]
    for i, (lbl, bg, sub) in enumerate(boxes):
        c = t.rows[0].cells[i]
        set_cell_bg(c, bg)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(lbl)
        r1.font.name = "Calibri"
        r1.font.size = Pt(9)
        r1.bold = True
        r1.font.color.rgb = C_WHITE
        p2 = c.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(sub)
        r2.font.name = "Calibri"
        r2.font.size = Pt(8)
        r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    for row in t.rows:
        for c in row.cells:
            c.width = Inches(1.3)

    doc.add_paragraph()

    # Phase details table
    heading2(doc, "Détail des 3 Phases de l'Agent")
    phases = [
        ("Phase 1",  "Collecte kubectl (déterministe)",
         "get pods/logs/describe/events — JAMAIS de LLM",
         "Données fiables avant l'analyse. Pas de dépendance LLM pour la collecte.",
         "Toujours"),
        ("Phase 2",  "Analyse LLM (Ollama JSON mode)",
         "ChatOllama format='json', temperature=0, num_predict=1024",
         "format='json' force Mistral à sortir du JSON valide. Fallback keyword si LLM KO.",
         "Toujours"),
        ("Phase 3",  "Remédiation GitOps",
         "scale_to_zero OU rollout restart → git commit → ArgoCD sync",
         "Démontre la boucle GitOps complète du cahier des charges PFE.",
         "AUTO_REMEDIATE=true seulement"),
    ]
    t2 = doc.add_table(rows=1 + len(phases), cols=5)
    t2.style = 'Table Grid'
    hdrs = ["Phase", "Nom", "Action", "Justification", "Condition"]
    for i, h in enumerate(hdrs):
        cell_text(t2.rows[0].cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(t2.rows[0].cells[i], HEX_TABLE_HDR)
    for r, row_data in enumerate(phases):
        bg = HEX_TABLE_EVEN if r % 2 == 0 else HEX_TABLE_ALT
        for c, val in enumerate(row_data):
            cell_text(t2.rows[r+1].cells[c], val, size=9)
            set_cell_bg(t2.rows[r+1].cells[c], bg)
    widths = [0.7, 1.5, 2.0, 2.3, 1.2]
    for row in t2.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    set_cell_borders(t2)
    doc.add_paragraph()


# ── Section 2: Structure des fichiers ─────────────────────────────────────────

def add_structure(doc):
    heading1(doc, "2. Structure du Projet & Rôle des Fichiers")

    code_block(doc, [
        "SRE-Copilot/",
        "├── credentials.txt              # Tokens & mots de passe (ne PAS versionner)",
        "├── kind-cluster.yaml            # Cluster KinD (3 nœuds, API port 6550)",
        "├── alertmanager-config.yaml     # Routage alertes Alertmanager → n8n",
        "├── pod-restart-alert.yaml       # PrometheusRule : PodRestartingTooMuch",
        "├── demo-app.yaml                # Déploiement crash-app (busybox exit 1)",
        "├── sre-agent-deploy.yaml        # Stack complète SRE Agent (SA+RBAC+PVC+Deploy+SVC)",
        "├── scripts/",
        "│   ├── start-sre-copilot.ps1   # Démarrage complet (cluster+stack+port-forwards)",
        "│   └── generate-argocd-token.ps1  # Génère token ArgoCD permanent",
        "├── sre-agent/",
        "│   ├── Dockerfile               # Python 3.12-slim + kubectl + git",
        "│   ├── requirements.txt         # FastAPI, LangChain, Ollama, Pydantic",
        "│   ├── main.py                  # FastAPI : /alert /analyze /health /status /test-crash-loop",
        "│   ├── agent.py                 # Pipeline 3 phases (collecte → LLM → GitOps)",
        "│   ├── models.py                # Pydantic : Alert, AlertManagerPayload, AgentResponse",
        "│   └── tools/",
        "│       ├── __init__.py          # Registre : READONLY_TOOLS, ALL_TOOLS",
        "│       ├── kubectl_tools.py     # Outils kubectl (lecture + écriture)",
        "│       ├── git_tools.py         # Commits gitops-repo + push GitHub optionnel",
        "│       └── argocd_tools.py      # API ArgoCD (sync, rollback, status)",
        "└── n8n-workflows/",
        "    └── sre-alert-workflow.json  # Workflow n8n 5 nœuds (importable UI)",
    ], label="ARBORESCENCE DU PROJET")

    heading2(doc, "Justification des Choix de Fichiers")
    make_table(doc,
        ["Fichier", "Rôle", "Pourquoi ce choix"],
        [
            ("kind-cluster.yaml", "Cluster Kubernetes local 3 nœuds, port API 6550",
             "KinD = Kubernetes in Docker, reproductible sur tout laptop. Port 6550 évite les conflits avec la plage réservée Windows (> 62000)."),
            ("alertmanager-config.yaml", "Route UNIQUEMENT les alertes namespace=demo vers n8n",
             "Filtre à la source = pas de bruit. En production : matcher des équipes/services précis."),
            ("pod-restart-alert.yaml", "Alerte PodRestartingTooMuch (> 3 restarts en 30s)",
             "Seuil bas pour faciliter la démo. Production : > 5 restarts en 5min avec fenêtre glissante."),
            ("demo-app.yaml", "Pod crash-app (busybox exit 1) = crash délibéré",
             "Image minimale (50Mi), crash immédiat, reproductible. Pas de dépendances applicatives."),
            ("sre-agent-deploy.yaml", "Déploiement K8s complet avec ServiceAccount+RBAC+PVC",
             "ServiceAccount in-cluster évite de monter le kubeconfig. PVC = gitops-repo persistant entre redémarrages."),
            ("agent.py", "Pipeline 3 phases déterministe",
             "Remplace LangChain ReAct (instable, > 50% échecs). Pipeline fixe = 3-5x plus rapide, > 95% fiabilité."),
            ("Dockerfile", "Python 3.12-slim + kubectl + git",
             "git ajouté pour les commits Phase 3. kubectl embarqué = l'agent n'a pas besoin de kubeconfig monté."),
        ],
        col_widths=[1.5, 2.0, 4.0]
    )


# ── Section 3: Variables d'environnement ──────────────────────────────────────

def add_env_vars(doc):
    heading1(doc, "3. Variables d'Environnement & Configuration")

    heading2(doc, "ConfigMap sre-agent-config (non-sensitif)")
    make_table(doc,
        ["Variable", "Valeur par défaut", "Rôle & Justification"],
        [
            ("OLLAMA_BASE_URL", "http://host.docker.internal:11434",
             "host.docker.internal = DNS spécial Docker qui résout vers l'adresse IP du host Windows depuis un pod KinD."),
            ("OLLAMA_MODEL", "mistral",
             "Mistral 7B = meilleur rapport qualité/performance pour les tâches SRE en mode JSON. Open-source, local, pas d'API key."),
            ("AUTO_REMEDIATE", "true",
             "Active la Phase 3. false = mode observation seule (investigation + runbook sans action). Permet de tester sans risque."),
            ("ARGOCD_SERVER", "https://argocd-server.argocd.svc.cluster.local",
             "DNS cluster interne. Depuis le pod sre-agent, ArgoCD est accessible sans port-forward."),
            ("ARGOCD_INSECURE", "true",
             "ArgoCD utilise un certificat auto-signé. true = ignore la validation TLS (environnement local uniquement)."),
            ("ARGOCD_APP_NAME", "demo-app",
             "Application ArgoCD à synchroniser après remédiation. Déclenche le redeploiement depuis git."),
            ("GITOPS_REPO_PATH", "/gitops-repo",
             "Chemin du dépôt git local (PVC monté). Les commits d'incidents y sont stockés sous incidents/<namespace>/."),
        ],
        col_widths=[1.8, 2.0, 3.9]
    )

    heading2(doc, "Secret sre-agent-secrets (sensitif — injecté via kubectl)")
    make_table(doc,
        ["Variable", "Valeur", "Rôle"],
        [
            ("GITHUB_TOKEN", "vide (optionnel)", "Push vers GitHub pour créer une vraie PR. Vide = commits locaux uniquement."),
            ("ARGOCD_TOKEN",  "JWT permanent sre-copilot:apiKey",
             "Token API ArgoCD sans expiry (expiresIn=0). Généré via generate-argocd-token.ps1."),
            ("GIT_USER_NAME",  "SRE-Copilot", "Auteur des commits gitops-repo."),
            ("GIT_USER_EMAIL", "sre-copilot@local", "Email auteur git."),
        ],
        col_widths=[1.8, 2.2, 3.7]
    )

    heading2(doc, "Token ArgoCD Permanent — Justification")
    badge(doc, "  TOKEN TYPE : sre-copilot:apiKey  |  EXPIRE : JAMAIS  |  RBAC : get/sync/update applications", HEX_TEAL)
    body_text(doc, "Le token de session JWT standard d'ArgoCD (/api/v1/session) expire en 24h. À chaque redémarrage du pod SRE Agent, l'investigation échouerait sur l'étape ArgoCD sync. La solution : créer un compte dédié 'sre-copilot' avec la capacité 'apiKey', générer un token permanent (expiresIn=0 = pas de champ 'exp' dans le JWT), et l'injecter dans le Secret Kubernetes. Ce token survit aux redémarrages du cluster.")


# ── Section 4: RBAC ────────────────────────────────────────────────────────────

def add_rbac(doc):
    heading1(doc, "4. RBAC Kubernetes — Permissions SRE Agent")

    body_text(doc, "Le principe du moindre privilège est appliqué : l'agent a uniquement les permissions nécessaires à ses 3 phases. Les opérations d'écriture (scale, restart) ne sont pas exposées comme outils LangChain.")
    doc.add_paragraph()

    make_table(doc,
        ["Ressource", "Namespace", "Verbes autorisés", "Phase", "Justification"],
        [
            ("pods, pods/log, events, namespaces, services", "Tous", "get, list, watch",
             "Phase 1", "Collecte kubectl — lecture seule des données d'investigation."),
            ("deployments, replicasets, statefulsets, daemonsets", "Tous", "get, list, watch, update, patch",
             "Phase 3", "Rollout restart du déploiement après classification."),
            ("deployments/scale", "Tous", "get, update, patch",
             "Phase 3", "Scale_to_zero pour stopper un CrashLoopBackOff immédiatement."),
            ("metrics.k8s.io / pods, nodes", "Tous", "get, list",
             "Phase 1", "Métriques CPU/mémoire pour analyse OOMKilled."),
        ],
        col_widths=[2.0, 0.8, 1.5, 0.8, 2.6]
    )

    heading2(doc, "Sécurité des Opérations d'Écriture")
    body_text(doc, "Les fonctions restart_deployment() et scale_deployment() dans kubectl_tools.py ne sont PAS décorées avec @tool. Elles ne peuvent donc jamais être appelées par un LLM via ReAct ou tool-calling. Seul le code Python de la Phase 3 peut les invoquer explicitement, après validation de la classification et du fix_type proposé par le LLM.")


# ── Section 5: Endpoints FastAPI ───────────────────────────────────────────────

def add_endpoints(doc):
    heading1(doc, "5. Endpoints FastAPI — SRE Agent")

    make_table(doc,
        ["Méthode", "Endpoint", "Usage", "Description"],
        [
            ("GET",  "/health",          "Liveness probe K8s",
             "Retourne {status: ok}. Utilisé par Kubernetes pour détecter les pods mort."),
            ("GET",  "/status",          "Vérification Ollama",
             "Vérifie la connectivité Ollama et la disponibilité du modèle Mistral. Utile au démarrage."),
            ("POST", "/alert",           "Webhook Alertmanager",
             "Point d'entrée principal. Accepte un payload AlertManagerPayload. Traite toutes les alertes 'firing' séquentiellement (Semaphore(1))."),
            ("POST", "/analyze",         "Test direct",
             "Analyse un seul objet Alert directement. Utile pour tester sans passer par Alertmanager."),
            ("POST", "/test-crash-loop", "Demo/simulation",
             "Simule une alerte CrashLoopBackOff sur le pod crash-app réel du namespace demo. Résout le nom du pod dynamiquement."),
        ],
        col_widths=[0.8, 1.5, 1.6, 3.8]
    )

    heading2(doc, "Modèle de Réponse (AgentResponse)")
    make_table(doc,
        ["Champ", "Type", "Description"],
        [
            ("classification",          "str",        "CrashLoopBackOff | OOMKilled | ImagePullBackOff | ConfigError | ResourceExhausted | NetworkError | Unknown"),
            ("root_cause",              "str",        "Explication technique de la cause racine, référençant les données kubectl collectées."),
            ("runbook",                 "list",       "3 étapes kubectl ordonnées : action, command, description."),
            ("auto_remediation_applied","bool",       "true si scale_to_zero ou restart a été exécuté avec succès."),
            ("remediation_details",     "str | null", "Description de l'action appliquée + output kubectl."),
            ("git_pr_url",              "str | null", "URL de la PR (GitHub) ou référence locale (local:branch/fix/...)."),
            ("argocd_sync_triggered",   "bool",       "true si l'API ArgoCD a été appelée avec succès pour synchroniser demo-app."),
            ("investigation_log",       "list[str]",  "Log interne de toutes les étapes : kubectl sections, LLM, remédiation."),
        ],
        col_widths=[2.0, 1.2, 4.5]
    )


# ── Section 6: Commandes ──────────────────────────────────────────────────────

def add_commands(doc):
    heading1(doc, "6. Référence des Commandes")

    # 6.1 Démarrage
    heading2(doc, "6.1  Démarrage de l'Environnement")
    code_block(doc, [
        "# Démarrage complet automatique (recommandé)",
        ".\\scripts\\start-sre-copilot.ps1",
        "",
        "# Exporter le kubeconfig manuellement",
        "kind export kubeconfig --name sre-cluster",
        "",
        "# Vérifier le contexte kubectl actif",
        "kubectl config current-context",
        "# Résultat attendu : kind-sre-cluster",
    ], label="DÉMARRAGE")

    # 6.2 État du cluster
    heading2(doc, "6.2  Vérification de l'État du Cluster")
    code_block(doc, [
        "kubectl get nodes                  # 3 nœuds Ready attendus",
        "kubectl get pods -A                # Tous les pods",
        "kubectl get pods -n monitoring     # Prometheus, Grafana, Alertmanager",
        "kubectl get pods -n argocd         # ArgoCD components",
        "kubectl get pods -n sre-agent      # SRE Agent",
        "kubectl get pods -n demo           # crash-app (quand déployée)",
        "kubectl get deployments -A         # État des déploiements",
    ], label="ÉTAT DU CLUSTER")

    # 6.3 KinD
    heading2(doc, "6.3  Gestion du Cluster KinD")
    code_block(doc, [
        "kind get clusters                                          # Lister les clusters",
        "kind create cluster --name sre-cluster --config kind-cluster.yaml  # Créer",
        "kind delete cluster --name sre-cluster                    # Supprimer",
        "kind load docker-image sre-agent:latest --name sre-cluster # Charger image",
    ], label="KIND")

    # 6.4 Build image
    heading2(doc, "6.4  Build & Déploiement de l'Image SRE Agent")
    code_block(doc, [
        "# Build l'image Docker",
        "cd sre-agent",
        "docker build -t sre-agent:latest .",
        "cd ..",
        "",
        "# Charger dans KinD (obligatoire — imagePullPolicy: Never)",
        "kind load docker-image sre-agent:latest --name sre-cluster",
        "",
        "# Redémarrer le pod pour prendre la nouvelle image",
        "kubectl rollout restart deployment/sre-agent -n sre-agent",
        "kubectl rollout status deployment/sre-agent -n sre-agent --timeout=120s",
    ], label="BUILD IMAGE")

    # 6.5 Port-forwards
    heading2(doc, "6.5  Port-Forwards Manuels")
    code_block(doc, [
        "# Grafana      -> http://localhost:3000",
        "kubectl port-forward svc/monitoring-grafana -n monitoring 3000:80",
        "",
        "# Prometheus   -> http://localhost:9090",
        "kubectl port-forward svc/monitoring-kube-prometheus-prometheus -n monitoring 9090:9090",
        "",
        "# Alertmanager -> http://localhost:9093",
        "kubectl port-forward svc/monitoring-kube-prometheus-alertmanager -n monitoring 9093:9093",
        "",
        "# ArgoCD       -> https://localhost:8080",
        "kubectl port-forward svc/argocd-server -n argocd 8080:443",
        "",
        "# SRE Agent    -> http://localhost:8001 (--address pour accès depuis Docker/n8n)",
        "kubectl port-forward --address 0.0.0.0 svc/sre-agent-service -n sre-agent 8001:8000",
    ], label="PORT-FORWARDS")

    # 6.6 Tests
    heading2(doc, "6.6  Tests des Endpoints SRE Agent")
    code_block(doc, [
        "# Health check",
        "Invoke-WebRequest -Uri 'http://localhost:8001/health' | Select-Object -ExpandProperty Content",
        "",
        "# Statut Ollama",
        "Invoke-WebRequest -Uri 'http://localhost:8001/status' | Select-Object -ExpandProperty Content",
        "",
        "# Simulation CrashLoop complète (30-60s)",
        "$resp = Invoke-WebRequest -Uri 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 120",
        "$resp.Content | ConvertFrom-Json | ConvertTo-Json -Depth 5",
    ], label="TESTS ENDPOINTS")

    # 6.7 Alertmanager
    heading2(doc, "6.7  Envoi Manuel d'une Alerte")
    code_block(doc, [
        "# Envoyer directement au webhook n8n",
        "$payload = @{",
        "    version = '4'; status = 'firing'; receiver = 'n8n-webhook'",
        "    alerts = @(@{",
        "        status = 'firing'",
        "        labels = @{ alertname='PodRestartingTooMuch'; namespace='demo'",
        "                    pod='crash-app'; severity='critical' }",
        "        annotations = @{ summary='Test'; description='Simulation' }",
        "        startsAt = (Get-Date -Format 'o')",
        "    })",
        "} | ConvertTo-Json -Depth 5",
        "Invoke-WebRequest -Uri 'http://localhost:5678/webhook/sre-alert' -Method POST `",
        "    -Body $payload -ContentType 'application/json'",
    ], label="ENVOI ALERTE")

    # 6.8 crash-app
    heading2(doc, "6.8  Gestion de crash-app (Pod de Test)")
    code_block(doc, [
        "kubectl apply -f demo-app.yaml                          # Déployer",
        "kubectl get pods -n demo -w                             # Surveiller (CrashLoopBackOff ~2min)",
        "kubectl logs -n demo -l app=crash-app --previous        # Logs du crash",
        "kubectl scale deployment crash-app --replicas=1 -n demo # Remettre à 1 replica",
        "kubectl delete -f demo-app.yaml                         # Supprimer",
    ], label="CRASH-APP")

    # 6.9 ArgoCD token
    heading2(doc, "6.9  Gestion du Token ArgoCD Permanent")
    code_block(doc, [
        "# Régénérer le token (si cluster recréé)",
        ".\\scripts\\generate-argocd-token.ps1",
        "",
        "# Vérifier le token dans le secret",
        "$b64 = kubectl get secret sre-agent-secrets -n sre-agent -o jsonpath='{.data.ARGOCD_TOKEN}'",
        "$token = [System.Text.Encoding]::UTF8.GetString([System.Convert]::FromBase64String($b64))",
        "",
        "# Décoder le JWT pour vérifier l'absence d'expiry",
        "$payload = $token.Split('.')[1]",
        "$padded = $payload + '=' * ((4 - $payload.Length % 4) % 4)",
        "[System.Text.Encoding]::UTF8.GetString([System.Convert]::FromBase64String($padded))",
        "# Attendu : {\"iss\":\"argocd\",\"sub\":\"sre-copilot:apiKey\",...} — PAS de champ exp",
    ], label="TOKEN ARGOCD")

    # 6.10 Debugging
    heading2(doc, "6.10  Monitoring & Debugging")
    code_block(doc, [
        "# Logs SRE Agent en temps réel",
        "kubectl logs -f deployment/sre-agent -n sre-agent",
        "",
        "# Commits dans gitops-repo",
        "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
        "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
        "kubectl exec -n sre-agent $pod -- git -C /gitops-repo log --oneline",
        "",
        "# Alertes actives Prometheus",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/alerts' | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | ConvertTo-Json -Depth 5",
        "",
        "# Events Kubernetes namespace demo",
        "kubectl get events -n demo --sort-by='.lastTimestamp' -w",
    ], label="DEBUGGING")


# ── Section 7: Workflow n8n ────────────────────────────────────────────────────

def add_n8n(doc):
    heading1(doc, "7. Workflow n8n — Détail des 5 Nœuds")

    make_table(doc,
        ["#", "Nœud", "Type", "Configuration clé", "Justification du choix"],
        [
            ("1", "webhook-trigger",  "Webhook",     "httpMethod: POST, path: sre-alert, responseMode: onReceived",
             "onReceived = réponse 200 immédiate. Evite le retry Alertmanager (timeout 30s)."),
            ("2", "filter-dedup",     "Code (JS)",   "Filtre namespace=demo + cache TTL 10min pour déduplication",
             "Sans dédup : la même alerte lance 10+ investigations simultanées pendant que le LLM tourne."),
            ("3", "call-sre-agent",   "HTTP Request","POST http://host.docker.internal:8001/alert, timeout 600s",
             "host.docker.internal = host Windows depuis Docker n8n. Timeout 600s > durée max investigation (60s)."),
            ("4", "parse-response",   "Code (JS)",   "Extrait classification, root_cause, auto_remediation_applied",
             "Enrichit le rapport avec des icônes de sévérité et formate pour l'affichage n8n."),
            ("5", "build-report",     "Code (JS)",   "Rapport Markdown structuré dans le résultat de l'exécution",
             "Visible dans l'historique n8n — sert de log d'audit des incidents traités."),
        ],
        col_widths=[0.3, 1.3, 1.0, 2.2, 2.9]
    )

    heading2(doc, "Settings Critiques du Workflow")
    make_table(doc,
        ["Paramètre", "Valeur", "Pourquoi"],
        [
            ("maxConcurrentExecutions", "1",           "1 seule investigation à la fois. Protège Mistral 7B (CPU-bound) de la saturation."),
            ("responseMode",            "onReceived",  "Répond HTTP 200 immédiatement, sans attendre la fin du workflow."),
            ("saveDataSuccessExecution","all",          "Conserve toutes les exécutions dans l'historique n8n pour l'audit."),
            ("executionOrder",          "v1",           "Ordre d'exécution déterministe garantissant la séquence correcte des nœuds."),
        ],
        col_widths=[2.0, 1.5, 4.2]
    )


# ── Section 8: Kyverno (Mission 4) ─────────────────────────────────────────────

def add_kyverno(doc):
    heading1(doc, "8. Mission 4 — Sécurité Kyverno (Policy as Code)")

    body_text(doc, "Kyverno applique des politiques de sécurité au niveau du cluster Kubernetes via des admission webhooks. Chaque création/mise à jour de ressource est validée contre les ClusterPolicies. Le mode Audit génère des PolicyReports sans bloquer, idéal pour évaluer l'impact avant durcissement.")
    doc.add_paragraph()

    heading2(doc, "Architecture Kyverno (4 controllers)")
    make_table(doc,
        ["Controller", "Rôle", "Pourquoi"],
        [
            ("admission-controller",  "Intercepte chaque requête API K8s (CREATE/UPDATE)",
             "Point d'entrée des validations. Webhook ValidatingAdmissionWebhook."),
            ("background-controller", "Audite les ressources EXISTANTES périodiquement",
             "Détecte les workloads déjà déployés avant l'installation des politiques."),
            ("reports-controller",    "Génère et maintient les PolicyReports",
             "Stockage des violations en CR (Custom Resource) lisible via kubectl."),
            ("cleanup-controller",    "Supprime ressources orphelines selon CleanupPolicy",
             "Gestion automatique des cycles de vie. Non utilisé dans ce PFE."),
        ],
        col_widths=[2.0, 2.5, 3.2]
    )

    heading2(doc, "Les 4 ClusterPolicies déployées")
    make_table(doc,
        ["Politique", "Vérifie", "Mode", "Justification métier"],
        [
            ("disallow-privileged-containers",
             "securityContext.privileged ≠ true",
             "Audit", "Privileged container = root accès noyau host. Container breakout possible."),
            ("disallow-latest-tag",
             "image:* sans :latest et tag obligatoire",
             "Audit", "':latest' = ambigu, non reproductible, rollback impossible. Traçabilité GitOps cassée."),
            ("require-labels",
             "Présence labels 'app' + 'owner'",
             "Audit", "Filtrage Prometheus, attribution coûts, audit qui possède quoi."),
            ("require-resource-limits",
             "resources.limits.memory + .cpu présents",
             "Audit", "Pod sans limits = noeud OOM. Garantit la QoS Guaranteed."),
        ],
        col_widths=[2.2, 2.0, 0.7, 2.8]
    )

    heading2(doc, "Structure des fichiers")
    code_block(doc, [
        "kyverno/",
        "├── policies/",
        "│   ├── 01-disallow-privileged.yaml",
        "│   ├── 02-disallow-latest-tag.yaml",
        "│   ├── 03-require-labels.yaml",
        "│   └── 04-require-resource-limits.yaml",
        "└── test-violations/",
        "    ├── violation-privileged.yaml",
        "    ├── violation-latest.yaml",
        "    ├── violation-no-labels.yaml",
        "    └── violation-no-limits.yaml",
    ], label="ARBORESCENCE KYVERNO")

    heading2(doc, "Commandes Kyverno")

    body_text(doc, "Installation (1 fois) :")
    code_block(doc, [
        "helm repo add kyverno https://kyverno.github.io/kyverno/",
        "helm repo update",
        "helm install kyverno kyverno/kyverno -n kyverno --create-namespace --timeout 5m",
        "",
        "# Attendre que les 4 controllers soient Ready",
        "kubectl wait --for=condition=Ready pods --all -n kyverno --timeout=180s",
    ], label="INSTALLATION")

    body_text(doc, "Déploiement des politiques :")
    code_block(doc, [
        "kubectl apply -f kyverno/policies/",
        "",
        "# Vérifier que les 4 politiques sont Ready",
        "kubectl get clusterpolicies",
        "# Attendu : 4 lignes avec READY=True",
    ], label="POLITIQUES")

    body_text(doc, "Tests de violation :")
    code_block(doc, [
        "# Appliquer 4 pods qui violent chacun une politique",
        "kubectl apply -f kyverno/test-violations/",
        "",
        "# Lister les pods (en mode Audit, ils sont créés sans blocage)",
        "kubectl get pods -n demo",
        "",
        "# Attendre 5s puis lister les PolicyReports générés",
        "Start-Sleep -Seconds 5",
        "kubectl get policyreport -n demo",
    ], label="TESTS VIOLATIONS")

    body_text(doc, "Inspection des PolicyReports :")
    code_block(doc, [
        "# Vue d'ensemble : pour chaque ressource auditée, count pass/fail/warn",
        "kubectl get policyreport -n demo",
        "",
        "# Détail d'un PolicyReport (avec les messages d'erreur)",
        "kubectl describe policyreport <UID> -n demo",
        "",
        "# Compter les violations par politique (PowerShell)",
        "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty items | `",
        "    Select-Object -ExpandProperty results | `",
        "    Where-Object { $_.result -eq 'fail' } | `",
        "    Group-Object -Property policy | Format-Table Name, Count",
    ], label="INSPECTION REPORTS")

    body_text(doc, "Passage en mode Enforce (blocage strict) — quand la base est saine :")
    code_block(doc, [
        "# Éditer chaque YAML dans kyverno/policies/ :",
        "# Remplacer  validationFailureAction: Audit",
        "# Par        validationFailureAction: Enforce",
        "",
        "kubectl apply -f kyverno/policies/",
        "",
        "# Désormais, tout pod en violation est REJETÉ à la création",
        "# Test : kubectl apply -f kyverno/test-violations/violation-privileged.yaml",
        "# → Error from server: admission webhook denied the request",
    ], label="MODE ENFORCE")

    heading2(doc, "Pourquoi le mode Audit d'abord ?")
    body_text(doc, "Activer Enforce immédiatement bloquerait la création de tout workload non conforme — y compris la crash-app existante qui sert aux démos Mission 3. Le mode Audit permet :")
    bullet_lines = [
        "• Découvrir l'ampleur des violations sur le cluster existant",
        "• Discuter avec les équipes avant de durcir",
        "• Corriger progressivement les workloads sans interruption",
        "• Mesurer l'impact (taux de conformité par namespace)",
    ]
    for line in bullet_lines:
        body_text(doc, line, indent=True)

    heading2(doc, "Cas pédagogique : violations sur crash-app")
    body_text(doc, "Sans modifier crash-app, Kyverno a détecté automatiquement 3 violations sur ce déploiement existant :")
    make_table(doc,
        ["Politique violée", "Cause exacte", "Correction recommandée"],
        [
            ("disallow-latest-tag",
             "image: busybox (sans tag = :latest implicite)",
             "Spécifier un tag immutable, ex: busybox:1.36"),
            ("require-labels",
             "Seul label 'app: crash-app', 'owner' absent",
             "Ajouter labels: { app: crash-app, owner: pfe-ghada }"),
            ("require-resource-limits",
             "limits.memory: 50Mi mais limits.cpu absent",
             "Ajouter limits.cpu, ex: cpu: 50m"),
        ],
        col_widths=[2.5, 3.0, 2.5]
    )
    body_text(doc, "C'est précisément la valeur de Kyverno : auditer les workloads existants pour révéler les dérives qu'on aurait laissé passer.")


# ── Section 9 : Mission 5 — Métriques Prometheus + Dashboard Grafana ──────────

def add_metrics(doc):
    heading1(doc, "9. Mission 5 — Métriques Prometheus & Dashboard Grafana")

    body_text(doc, "Cette mission ferme la boucle de l'observabilité : le SRE Agent expose ses propres KPIs via /metrics, Prometheus les scrape via un ServiceMonitor, Grafana les visualise sur un dashboard auto-importé.")
    doc.add_paragraph()

    heading2(doc, "Architecture observabilité")
    code_block(doc, [
        "SRE Agent (FastAPI)",
        "  └── /metrics (Prometheus format, 0.0.4)",
        "       └── exposé via Service sre-agent-service (port 8000)",
        "            └── ServiceMonitor sre-agent (label release: monitoring)",
        "                 └── Prometheus scrape toutes les 30s",
        "                      └── Grafana datasource Prometheus",
        "                           └── Dashboard 'SRE Copilot — KPIs'",
    ], label="FLUX OBSERVABILITÉ")

    heading2(doc, "Métriques exposées")
    make_table(doc,
        ["Métrique", "Type", "Labels", "Description"],
        [
            ("sre_alerts_total",                       "Counter",
             "classification, severity, namespace",
             "Total d'alertes reçues et analysées par l'agent"),
            ("sre_remediated_total",                   "Counter",
             "classification, fix_type, result",
             "Total des remédiations tentées (success | skipped | failed)"),
            ("sre_argocd_sync_total",                  "Counter",
             "result",
             "Appels ArgoCD sync (success | skipped)"),
            ("sre_llm_fallback_total",                 "Counter",
             "reason",
             "Bascules en keyword fallback (llm_error | invalid_classification)"),
            ("sre_investigation_duration_seconds",     "Histogram",
             "(buckets fixes)",
             "Distribution des durées d'investigation Phase 1+2+3"),
            ("sre_agent_info",                         "Gauge",
             "version, ollama_model",
             "Métadonnées de l'agent (toujours 1)"),
        ],
        col_widths=[2.4, 1.0, 2.3, 2.5]
    )

    heading2(doc, "Dashboard Grafana — 9 panels")
    make_table(doc,
        ["#", "Panel", "Type", "Requête PromQL"],
        [
            ("1", "Alertes totales reçues",         "Stat",        "sum(sre_alerts_total)"),
            ("2", "Taux de remédiation auto",        "Stat (%)",    "sum(sre_remediated_total{result=\"success\"}) / clamp_min(sum(sre_alerts_total), 1)"),
            ("3", "MTTR médian (1h)",                "Stat (s)",    "histogram_quantile(0.5, sum(rate(sre_investigation_duration_seconds_bucket[1h])) by (le))"),
            ("4", "ArgoCD sync réussis",             "Stat",        "sum(sre_argocd_sync_total{result=\"success\"})"),
            ("5", "Keyword fallbacks",               "Stat",        "sum(sre_llm_fallback_total)"),
            ("6", "Taux d'alertes par classification","Time series", "sum by (classification) (rate(sre_alerts_total[5m]) * 3600)"),
            ("7", "Répartition classifications",     "Donut",       "sum by (classification) (sre_alerts_total)"),
            ("8", "Durée investigation p50/p95",     "Time series", "histogram_quantile(0.5/0.95, ... seconds_bucket)"),
            ("9", "Remédiations par type",           "Bar chart",   "sum by (fix_type, result) (sre_remediated_total)"),
        ],
        col_widths=[0.3, 2.2, 1.2, 4.3]
    )

    heading2(doc, "Structure des fichiers")
    code_block(doc, [
        "sre-agent/",
        "├── metrics.py                # Définition des compteurs/histogram",
        "└── main.py                   # Endpoint GET /metrics",
        "",
        "monitoring/",
        "├── grafana-dashboard.json    # Dashboard JSON Grafana (9 panels)",
        "└── grafana-dashboard-cm.yaml # ConfigMap pour auto-import via sidecar",
        "",
        "sre-agent-deploy.yaml         # ServiceMonitor inclus en bas du fichier",
    ], label="FICHIERS MISSION 5")

    heading2(doc, "Commandes de déploiement & vérification")

    body_text(doc, "Déploiement (déjà inclus dans start-sre-copilot.ps1 v2) :")
    code_block(doc, [
        "# 1. ServiceMonitor (s'applique avec sre-agent-deploy.yaml)",
        "kubectl apply -f sre-agent-deploy.yaml",
        "",
        "# 2. Dashboard Grafana via ConfigMap labeled grafana_dashboard=1",
        "kubectl create configmap sre-copilot-dashboard `",
        "    --from-file=grafana-dashboard.json=monitoring/grafana-dashboard.json `",
        "    -n monitoring --dry-run=client -o yaml | kubectl apply -f -",
        "kubectl label configmap sre-copilot-dashboard -n monitoring `",
        "    grafana_dashboard=1 --overwrite",
    ], label="DEPLOIEMENT")

    body_text(doc, "Vérification :")
    code_block(doc, [
        "# Tester l'endpoint /metrics",
        "Invoke-WebRequest 'http://localhost:8001/metrics' | Select-Object -ExpandProperty Content",
        "",
        "# Vérifier que Prometheus scrape l'agent",
        "(Invoke-WebRequest 'http://localhost:9090/api/v1/targets').Content | `",
        "    ConvertFrom-Json | Select-Object -ExpandProperty data | `",
        "    Select-Object -ExpandProperty activeTargets | `",
        "    Where-Object { $_.labels.job -eq 'sre-agent-service' } | `",
        "    Select-Object scrapeUrl, health, lastScrape",
        "# Attendu : health=up",
        "",
        "# Requête PromQL",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/query?query=sre_agent_info' | `",
        "    Select-Object -ExpandProperty Content",
        "",
        "# Ouvrir le dashboard Grafana",
        "Start-Process 'http://localhost:3000/d/sre-copilot/sre-copilot-kpis'",
    ], label="VERIFICATION")

    heading2(doc, "Pourquoi cette mission complète le projet")
    body_text(doc, "Sans cette mission, l'agent serait un 'opaque box' : on ne sait pas combien d'alertes il traite, ni combien de temps Mistral met, ni quel taux de remédiation. Avec Mission 5 :")
    bullet_lines = [
        "• MTTR mesurable et historique (alimente les KPI PFE)",
        "• Tableau de bord exécutif lisible par un manager non-technique",
        "• Détection des dérives (taux de fallback en hausse = Mistral instable)",
        "• Justification chiffrée pour le passage en production (taux de succès > X%)",
    ]
    for line in bullet_lines:
        body_text(doc, line, indent=True)


def main():
    out_path = os.path.join(os.path.dirname(__file__), "..", "SRE-Copilot-Commandes-Fichiers.docx")
    out_path = os.path.abspath(out_path)

    doc = Document()

    # Page setup: A4
    section = doc.sections[0]
    section.page_width  = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default font
    doc.styles['Normal'].font.name = "Calibri"
    doc.styles['Normal'].font.size = Pt(10)

    add_cover(doc)
    add_architecture(doc)
    add_structure(doc)
    add_env_vars(doc)
    add_rbac(doc)
    add_endpoints(doc)
    add_commands(doc)
    add_n8n(doc)
    add_kyverno(doc)
    add_metrics(doc)

    # Footer note
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("PFE IFS 013/26 — PwC DigiTech Cloud Native Operations | SRE Copilot | 2026-06-03")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY_LITE
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
