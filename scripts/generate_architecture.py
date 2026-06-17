"""
Generates SRE-Copilot-Architecture.docx
Document de référence : architecture complète, rôle de chaque fichier, commandes.
Conçu pour la défense PFE devant le jury.

Run: python scripts/generate_architecture.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Palette ────────────────────────────────────────────────────────────────────
C_NAVY     = RGBColor(0x1F, 0x38, 0x64)
C_BLUE     = RGBColor(0x2E, 0x75, 0xB6)
C_TEAL     = RGBColor(0x00, 0x70, 0x7F)
C_GREEN    = RGBColor(0x37, 0x56, 0x23)
C_ORANGE   = RGBColor(0xC5, 0x50, 0x0C)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT     = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED    = RGBColor(0x6E, 0x7B, 0x8B)
C_GRAY_LT  = RGBColor(0xD6, 0xDC, 0xE4)

HEX_NAVY      = "1F3864"
HEX_TEAL      = "00707F"
HEX_BLUE      = "2E75B6"
HEX_TABLE_HDR = "1F3864"
HEX_TABLE_ALT = "DEEAF1"
HEX_CODE_BG   = "F2F2F2"
HEX_GREEN_BG  = "E2EFDA"
HEX_ORANGE_BG = "FFF2CC"
HEX_BLUE_BG   = "D9E1F2"


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_cell_borders(table, color="D6DCE4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
                tcBorders.append(b)
            tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, color=None, size=10, font="Calibri",
              align=WD_ALIGN_PARAGRAPH.LEFT, mono=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.clear(); p.alignment = align
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Courier New" if mono else font
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return r


# ── Typo helpers ───────────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(18); r.bold = True; r.font.color.rgb = C_NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8'); bot.set(qn('w:space'), '2'); bot.set(qn('w:color'), HEX_NAVY)
    pBdr.append(bot); pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(13); r.bold = True; r.font.color.rgb = C_BLUE
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.bold = True; r.font.color.rgb = C_TEAL
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = C_TEXT
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.6)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = C_TEXT
    return p

def callout(doc, text, bg_hex, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(8); lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"; lr.font.size = Pt(9); lr.bold = True; lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"  {text}  ")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = C_TEXT
    set_para_bg(p, bg_hex)
    return p

def code_block(doc, lines, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6); lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"; lr.font.size = Pt(9); lr.bold = True; lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.3)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"; r.font.size = Pt(9); r.font.color.rgb = C_TEXT
        set_para_bg(p, HEX_CODE_BG)
    sp = doc.add_paragraph(); sp.paragraph_format.space_after = Pt(4)

def make_table(doc, headers, rows, col_widths=None, mono_cols=None):
    """mono_cols : liste d'indices de colonnes en monospace."""
    mono_cols = mono_cols or []
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(t.rows[0].cells[i], HEX_TABLE_HDR)
    for r, row in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else HEX_TABLE_ALT
        for c, val in enumerate(row):
            cell_text(t.rows[r+1].cells[c], val, size=10, mono=(c in mono_cols))
            set_cell_bg(t.rows[r+1].cells[c], bg)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    set_cell_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ════════════════════════════════════════════════════════════════════════════
# CONTENU
# ════════════════════════════════════════════════════════════════════════════

def add_cover(doc):
    bar = doc.add_paragraph(); set_para_bg(bar, HEX_NAVY)
    r = bar.add_run("  SRE Copilot — Architecture")
    r.font.name = "Calibri"; r.font.size = Pt(26); r.bold = True; r.font.color.rgb = C_WHITE

    sub = doc.add_paragraph(); set_para_bg(sub, HEX_TEAL)
    sr = sub.add_run("  Guide complet pour maitriser et defendre le projet")
    sr.font.name = "Calibri"; sr.font.size = Pt(13); sr.font.color.rgb = C_WHITE

    doc.add_paragraph()

    info = [
        ("Pour",      "Soutenance PFE devant le jury PwC."),
        ("Objectif",  "Comprendre chaque fichier, chaque script, chaque commande, sans ambiguite."),
        ("Format",    "Tableaux clairs, exemples concrets, aucun jargon non explique."),
        ("Auteur",    "Ghada Ben Mansour — PFE IFS 013/26"),
        ("Date",      "2026-06-10"),
    ]
    t = doc.add_table(rows=len(info), cols=2); t.style = 'Table Grid'
    for i, (k, v) in enumerate(info):
        bg = HEX_TABLE_ALT if i % 2 == 0 else "FFFFFF"
        cell_text(t.rows[i].cells[0], k, bold=True, color=C_NAVY, size=11)
        cell_text(t.rows[i].cells[1], v, size=11)
        set_cell_bg(t.rows[i].cells[0], bg); set_cell_bg(t.rows[i].cells[1], bg)
    for row in t.rows:
        row.cells[0].width = Inches(1.3); row.cells[1].width = Inches(5.4)
    set_cell_borders(t)
    doc.add_page_break()


# ─── Section 1 — Vue d'ensemble ───────────────────────────────────────────────

def add_overview(doc):
    heading1(doc, "1.  Le projet en une page")

    body(doc, "SRE Copilot est un agent IA qui automatise la reponse aux incidents Kubernetes. Quand une alerte tombe (un pod qui crash, par exemple), l'agent :")
    bullet(doc, "1. Reçoit l'alerte d'Alertmanager via n8n")
    bullet(doc, "2. Enquete avec kubectl (logs, events, describe)")
    bullet(doc, "3. Demande a Mistral 7B (LLM local) de classifier et proposer un fix")
    bullet(doc, "4. Applique le fix (scale ou restart) + commit Git + sync ArgoCD")
    bullet(doc, "5. Documente tout dans une dashboard centralisee")

    heading2(doc, "Architecture en une image")
    code_block(doc, [
        "Prometheus  →  Alertmanager  →  n8n webhook  →  SRE Agent",
        "                                                    │",
        "                                                    ├─ Phase 1: kubectl (collecte)",
        "                                                    ├─ Phase 2: Mistral 7B (JSON)",
        "                                                    └─ Phase 3: kubectl scale",
        "                                                       + git commit",
        "                                                       + ArgoCD sync",
        "                                                            │",
        "                                            +---------------+---------------+",
        "                                            │                               │",
        "                                       /metrics (Prom)              /dashboard (HTML)",
        "                                            │                               │",
        "                                            └──→ Grafana                    └──→ Browser",
    ], label="FLUX DE BOUT EN BOUT")

    heading2(doc, "Les 5 missions du PFE")
    make_table(doc,
        ["#", "Mission", "Livrable principal"],
        [
            ("1", "Infrastructure",           "Cluster KinD + Prometheus + Grafana + ArgoCD + n8n"),
            ("2", "Pipeline d'alertes",       "Alertmanager route les alertes vers n8n, qui appelle l'agent"),
            ("3", "Agent IA + auto-remediation", "Pipeline 3 phases (kubectl → Mistral → action) + dashboard"),
            ("4", "Securite Kyverno",         "4 ClusterPolicies qui auditent les workloads en continu"),
            ("5", "Observabilite",            "Endpoint /metrics + dashboard Grafana avec 9 panels KPI"),
        ],
        col_widths=[0.3, 2.0, 4.5]
    )


# ─── Section 2 — Arborescence ─────────────────────────────────────────────────

def add_tree(doc):
    heading1(doc, "2.  Arborescence du projet")

    body(doc, "Voici la structure complete du projet, avec un commentaire sur chaque element :")
    code_block(doc, [
        "SRE-Copilot/",
        "│",
        "├─ Fichiers racine (manifests Kubernetes principaux)",
        "│  ├─ kind-cluster.yaml              Cluster KinD : 1 control-plane + 2 workers",
        "│  ├─ alertmanager-config.yaml       Routage alertes vers n8n (filtre namespace=demo)",
        "│  ├─ pod-restart-alert.yaml         PrometheusRule : detecter PodRestartingTooMuch",
        "│  ├─ demo-app.yaml                  Pod 'crash-app' qui crash volontairement (test)",
        "│  ├─ sre-agent-deploy.yaml          Deploiement complet du SRE Agent (RBAC, Service…)",
        "│  ├─ credentials.txt                Mots de passe & tokens (NON versionne)",
        "│  └─ .gitignore                     Exclut credentials, .docx, rapport, etc.",
        "│",
        "├─ sre-agent/                         Code Python de l'agent IA",
        "│  ├─ Dockerfile                     Image Python 3.12 + kubectl + git",
        "│  ├─ requirements.txt               FastAPI, LangChain, Ollama, prometheus_client",
        "│  ├─ main.py                        FastAPI : endpoints HTTP (alert, dashboard, metrics…)",
        "│  ├─ agent.py                       Pipeline 3 phases (collecte / LLM / GitOps)",
        "│  ├─ metrics.py                     Definition des 6 metriques Prometheus",
        "│  ├─ models.py                      Schemas Pydantic (Alert, AgentResponse…)",
        "│  └─ tools/",
        "│     ├─ kubectl_tools.py            Wrapper kubectl avec auth in-cluster",
        "│     ├─ git_tools.py                Commit + push GitOps",
        "│     └─ argocd_tools.py             API ArgoCD (sync, rollback)",
        "│",
        "├─ kyverno/                          Mission 4 — politiques de securite",
        "│  ├─ policies/                      4 ClusterPolicies en mode Audit",
        "│  │  ├─ 01-disallow-privileged.yaml",
        "│  │  ├─ 02-disallow-latest-tag.yaml",
        "│  │  ├─ 03-require-labels.yaml",
        "│  │  └─ 04-require-resource-limits.yaml",
        "│  └─ test-violations/               4 pods de test qui violent chacun 1 politique",
        "│",
        "├─ argocd/                           Mission 3 — GitOps loop",
        "│  ├─ demo-app-application.yaml      Application ArgoCD qui watch ce repo GitHub",
        "│  ├─ rbac-cm-patch.yaml             RBAC : compte sre-copilot peut sync",
        "│  └─ sre-agent-app.yaml             (Historique — non utilise actuellement)",
        "│",
        "├─ monitoring/                       Mission 5 — observabilite",
        "│  ├─ grafana-dashboard.json         Dashboard 9 panels (JSON Grafana)",
        "│  └─ grafana-dashboard-cm.yaml      ConfigMap d'auto-import (sidecar Grafana)",
        "│",
        "├─ n8n-workflows/",
        "│  └─ sre-alert-workflow.json        Workflow n8n 5 noeuds (webhook → agent → report)",
        "│",
        "├─ scripts/                          Scripts d'orchestration (PowerShell + Python)",
        "│  ├─ start-sre-copilot.ps1          Demarrage complet de la stack",
        "│  ├─ stop-sre-copilot.ps1           Arret des port-forwards",
        "│  ├─ reset-demo.ps1                 Nettoyage entre 2 demos",
        "│  ├─ run-demo.ps1                   Lancement d'une demo guidee",
        "│  ├─ generate-argocd-token.ps1      Generation du token ArgoCD permanent",
        "│  └─ generate_*.py                  Generateurs de documentation .docx",
        "│",
        "└─ rapport/                          Redaction LaTeX du rapport PFE",
        "   ├─ main.tex                       Source LaTeX",
        "   ├─ main.pdf                       PDF compile",
        "   ├─ figures/roadmap.png / .svg     Roadmap visuelle pour le rapport",
        "   └─ chapters/                      Sous-chapitres (a remplir)",
    ], label="STRUCTURE COMPLETE")


# ─── Section 3 — Scripts ──────────────────────────────────────────────────────

def add_scripts(doc):
    heading1(doc, "3.  Tous les scripts — quand et comment les utiliser")

    body(doc, "Le projet contient 10 scripts. Les uns servent au runtime (PowerShell), les autres a la documentation (Python).")

    heading2(doc, "Scripts d'orchestration runtime (PowerShell)")
    make_table(doc,
        ["Script", "Quand l'utiliser", "Ce qu'il fait"],
        [
            ("start-sre-copilot.ps1",
             "Au demarrage (premiere fois OU apres recreation cluster)",
             "Cree le cluster KinD, deploie Prometheus + ArgoCD + Kyverno + SRE Agent + dashboard Grafana, injecte le token ArgoCD, demarre n8n et les 6 port-forwards. Tout en 1 commande."),
            ("stop-sre-copilot.ps1",
             "A la fin de la journee",
             "Tue uniquement les port-forwards. Le cluster reste actif pour reprise rapide. Possibilite de supprimer aussi le cluster (interactif)."),
            ("reset-demo.ps1",
             "Entre 2 demos",
             "Tue les jobs PowerShell en background, silence Alertmanager 1h, scale crash-app a 0. Option --ClearHistory pour vider la dashboard."),
            ("run-demo.ps1",
             "Au debut d'une demo devant le jury",
             "Verifie la stack (health checks), leve les silences, deploie crash-app, ouvre la dashboard et n8n, suit les logs en direct."),
            ("generate-argocd-token.ps1",
             "Manuellement si le token expire ou si le secret est ecrase",
             "Cree le compte sre-copilot dans ArgoCD avec capacite apiKey, genere un token permanent (sans exp), l'injecte dans le Secret sre-agent-secrets, restart le SRE Agent."),
        ],
        col_widths=[2.0, 2.5, 3.0]
    )

    heading2(doc, "Generateurs de documentation (Python)")
    body(doc, "Ces scripts produisent les .docx du projet a partir de python-docx. Aucune dependance autre que pip install python-docx Pillow.")
    make_table(doc,
        ["Script", "Genere", "Contenu"],
        [
            ("generate_architecture.py",       "SRE-Copilot-Architecture.docx",       "Ce document lui-meme."),
            ("generate_commandes_fichiers.py", "SRE-Copilot-Commandes-Fichiers.docx", "Reference exhaustive des commandes et fichiers."),
            ("generate_guide_test.py",         "SRE-Copilot-Guide-Test.docx",         "Guide de test A-Z par phase."),
            ("generate_scenario_e2e.py",       "SRE-Copilot-Scenario-E2E.docx",       "Scenario complet Mission 1 → 5 (24 etapes numerotees)."),
            ("generate_explication.py",        "SRE-Copilot-Explication.docx",        "Explications pedagogiques (CrashLoop vs ConfigError, Mistral, etc.)."),
            ("generate_roadmap.py",            "rapport/figures/roadmap.png",         "Roadmap visuelle 5 missions (PNG haute resolution)."),
        ],
        col_widths=[2.5, 2.5, 2.5]
    )

    heading2(doc, "Comment invoquer les scripts")
    code_block(doc, [
        "# Scripts PowerShell (depuis la racine du projet)",
        ".\\scripts\\start-sre-copilot.ps1",
        ".\\scripts\\reset-demo.ps1",
        ".\\scripts\\reset-demo.ps1 -ClearHistory   # variante : vide aussi l'historique dashboard",
        ".\\scripts\\run-demo.ps1",
        ".\\scripts\\stop-sre-copilot.ps1",
        ".\\scripts\\generate-argocd-token.ps1",
        "",
        "# Scripts Python (regenerer la documentation)",
        "python scripts\\generate_architecture.py",
        "python scripts\\generate_commandes_fichiers.py",
        "python scripts\\generate_guide_test.py",
        "python scripts\\generate_scenario_e2e.py",
        "python scripts\\generate_explication.py",
        "python scripts\\generate_roadmap.py",
    ], label="COMMANDES")

    callout(doc,
        "Important : seul start-sre-copilot.ps1 appelle automatiquement generate-argocd-token.ps1 en interne. Les autres scripts sont independants et a invoquer manuellement.",
        HEX_BLUE_BG, label="ATTENTION")


# ─── Section 4 — Manifests YAML ───────────────────────────────────────────────

def add_yaml(doc):
    heading1(doc, "4.  Manifests YAML — role de chacun")

    heading2(doc, "Configuration cluster et infrastructure")
    make_table(doc,
        ["Fichier", "Role"],
        [
            ("kind-cluster.yaml",         "Definit le cluster KinD : 3 noeuds (1 control-plane + 2 workers), port API 6550 pour eviter conflits."),
            ("alertmanager-config.yaml",  "Configure Alertmanager : route les alertes namespace=demo + severity>=warning vers le webhook n8n. group_wait 1m, repeat_interval 12h."),
            ("pod-restart-alert.yaml",    "PrometheusRule : declenche l'alerte 'PodRestartingTooMuch' quand un pod a > 3 restarts en 30 secondes."),
            ("demo-app.yaml",             "Deploiement 'crash-app' : un pod busybox qui execute 'exit 1' immediatement. Sert a declencher des CrashLoopBackOff pour les tests."),
            ("sre-agent-deploy.yaml",     "Stack complete du SRE Agent : Namespace + ServiceAccount + ClusterRole + ClusterRoleBinding + ConfigMap + PersistentVolumeClaim + Deployment + Service + ServiceMonitor. Le Secret est gere separement (par generate-argocd-token.ps1)."),
        ],
        col_widths=[2.4, 4.6]
    )

    heading2(doc, "Mission 3 — ArgoCD")
    make_table(doc,
        ["Fichier", "Role"],
        [
            ("argocd/demo-app-application.yaml",
             "Application ArgoCD nommee 'demo-app' qui watch le repo GitHub du projet. Quand l'agent appelle sync, ArgoCD relit demo-app.yaml et reapplique."),
            ("argocd/rbac-cm-patch.yaml",
             "Patch sur le ConfigMap argocd-rbac-cm : autorise le compte 'sre-copilot' a faire get/sync/update sur les Applications."),
        ],
        col_widths=[2.8, 4.2]
    )

    heading2(doc, "Mission 4 — Kyverno (4 ClusterPolicies)")
    make_table(doc,
        ["Politique", "Ce qu'elle interdit", "Severite"],
        [
            ("disallow-privileged-containers", "Conteneurs avec securityContext.privileged=true (acces root au noyau host).", "high"),
            ("disallow-latest-tag",            "Images sans tag ou avec :latest (non-reproductibles, rollback impossible).", "medium"),
            ("require-labels",                 "Pods sans labels 'app' et 'owner' obligatoires.",                            "medium"),
            ("require-resource-limits",        "Pods sans limits.memory + limits.cpu (risque de saturer le noeud).",         "high"),
        ],
        col_widths=[2.4, 3.6, 1.0]
    )

    heading2(doc, "Mission 5 — Observabilite")
    make_table(doc,
        ["Fichier", "Role"],
        [
            ("monitoring/grafana-dashboard.json",   "Dashboard Grafana 9 panels : KPIs (alertes, success rate, MTTR), graphes time-series, donut classifications, bar remediations."),
            ("monitoring/grafana-dashboard-cm.yaml","ConfigMap labelle 'grafana_dashboard=1' que le sidecar Grafana detecte automatiquement et importe."),
        ],
        col_widths=[2.8, 4.2]
    )

    heading2(doc, "Workflow n8n")
    make_table(doc,
        ["Fichier", "Role"],
        [
            ("n8n-workflows/sre-alert-workflow.json",
             "Workflow 5 noeuds : Alertmanager Webhook → Filter & Dedup (TTL 10min) → Call SRE Agent (POST /alert) → Parse Response → Build Incident Report. responseMode=onReceived pour repondre 200 immediatement."),
        ],
        col_widths=[2.8, 4.2]
    )


# ─── Section 5 — Code Python ──────────────────────────────────────────────────

def add_python(doc):
    heading1(doc, "5.  Code Python du SRE Agent — role de chaque fichier")

    body(doc, "Le code de l'agent fait ~2300 lignes reparties en 8 fichiers. Chaque fichier a un role unique et explicite.")

    make_table(doc,
        ["Fichier", "Lignes", "Role"],
        [
            ("Dockerfile",
             "23",
             "Image Python 3.12-slim + kubectl + git. Build local puis 'kind load' dans le cluster."),
            ("requirements.txt",
             "13",
             "Dependances : FastAPI, uvicorn, Pydantic, httpx, LangChain (langchain-ollama, langchain-community), requests (pour ArgoCD), prometheus_client."),
            ("main.py",
             "989",
             "Serveur FastAPI. Endpoints : /health, /status, /alert (webhook Alertmanager), /analyze, /test-crash-loop, /dashboard (HTML), /api/incidents, /api/stats, /metrics (Prometheus). Inclut le HTML complet de la dashboard."),
            ("agent.py",
             "530",
             "Cur de l'agent : pipeline 3 phases. Phase 1 : collecte kubectl. Phase 2 : ChatOllama JSON mode + validator de classification. Phase 3 : scale_to_zero / restart + git commit + ArgoCD sync conditionnel."),
            ("metrics.py",
             "92",
             "Definition de 6 metriques Prometheus (Counter, Histogram, Gauge). Fonction record_investigation() appelee a chaque fin de pipeline."),
            ("models.py",
             "68",
             "Schemas Pydantic : Alert, AlertManagerPayload, AgentResponse, RemediationStep. Validation stricte des donnees a l'entree."),
            ("tools/kubectl_tools.py",
             "214",
             "Wrapper kubectl. Auth automatique : ServiceAccount token si in-cluster, kubeconfig local sinon. Outils LangChain en lecture seule (get_pod_logs, describe…). Fonctions ecriture (scale, restart) en Python pur pour eviter les abuses LLM."),
            ("tools/git_tools.py",
             "164",
             "Operations git locales et push GitHub optionnel. Cree un commit par incident dans /gitops-repo/incidents/."),
            ("tools/argocd_tools.py",
             "165",
             "API ArgoCD via Python requests : list / get_status / sync / rollback. Authentification par Bearer token (sre-copilot:apiKey)."),
        ],
        col_widths=[2.0, 0.6, 4.4]
    )

    heading2(doc, "Garde-fous du code")
    make_table(doc,
        ["Garde-fou", "Ou", "Pourquoi"],
        [
            ("ALLOWED_CLASSIFICATIONS",
             "agent.py",
             "7 classes autorisees (CrashLoopBackOff, OOMKilled, ImagePullBackOff, ConfigError, ResourceExhausted, NetworkError, Unknown). Si Mistral inverte une classe (ex: 'PolicyViolation'), keyword_fallback normalise."),
            ("ARGOCD_SYNC_AFTER_FIX = {'restart'}",
             "agent.py",
             "Empeche la boucle infinie : scale_to_zero ne declenche PAS ArgoCD sync (sinon ArgoCD reapplique replicas=1 = annule la remediation)."),
            ("asyncio.Semaphore(1)",
             "main.py",
             "Une seule investigation Mistral a la fois (CPU-bound). Sequentialise les alertes en rafale."),
            ("format='json' dans ChatOllama",
             "agent.py",
             "Force Mistral a produire du JSON valide (sinon impossible a parser)."),
            ("Validator post-LLM",
             "agent.py",
             "Si Mistral genere une classification hors liste → keyword_fallback prend le relai."),
            ("Skip rollout restart si deployment absent",
             "generate-argocd-token.ps1",
             "Lors du 1er deploiement, le SRE Agent n'existe pas encore : le script doit ne pas planter."),
        ],
        col_widths=[2.4, 1.8, 3.0]
    )


# ─── Section 6 — Comment lancer le projet ─────────────────────────────────────

def add_launch(doc):
    heading1(doc, "6.  Comment lancer le projet — etape par etape")

    heading2(doc, "Demarrage initial (1ere fois OU apres avoir supprime le cluster)")
    code_block(doc, [
        "# Verifier que Docker Desktop tourne (mandatory)",
        "docker info",
        "",
        "# Verifier qu'Ollama tourne et que Mistral est tire",
        "ollama list   # doit montrer 'mistral'",
        "# Si pas la : ollama pull mistral",
        "",
        "# Demarrage complet de la stack (15-20 min la 1ere fois, 30s ensuite)",
        ".\\scripts\\start-sre-copilot.ps1",
    ], label="DEMARRAGE INITIAL")
    body(doc, "Le script affiche un OK vert pour chaque etape. A la fin : 5 lignes 'OK' pour les health checks (SRE Agent, Grafana, Prometheus, Alertmanager, n8n).")

    heading2(doc, "Pour une demo standard (devant le jury)")
    code_block(doc, [
        "# 1. Reset l'etat (5 secondes)",
        ".\\scripts\\reset-demo.ps1",
        "",
        "# 2. Lancer la demo (~90 secondes pour deployer + suivre les logs)",
        ".\\scripts\\run-demo.ps1",
        "",
        "# Le script ouvre 2 navigateurs (dashboard SRE Agent + n8n executions)",
        "# Puis lance kubectl logs -f sur le SRE Agent.",
        "# Tu vois en direct : Phase 1 → Phase 2 (Mistral, 3 min) → Phase 3 → Complete.",
    ], label="CYCLE DE DEMO")

    heading2(doc, "Pour arreter en fin de journee")
    code_block(doc, [
        "# Stoppe les port-forwards mais conserve le cluster",
        ".\\scripts\\stop-sre-copilot.ps1",
        "",
        "# Pour reprendre : juste relancer start-sre-copilot.ps1",
        "# Si le cluster est detecte, il saute le deploiement et relance juste les port-forwards.",
    ], label="ARRET")

    heading2(doc, "URLs et credentials a retenir")
    make_table(doc,
        ["Service", "URL", "Login"],
        [
            ("SRE Agent dashboard", "http://localhost:8001/dashboard",   "Aucun"),
            ("SRE Agent /metrics",  "http://localhost:8001/metrics",     "Aucun"),
            ("Grafana KPI",         "http://localhost:3000/d/sre-copilot/", "admin / prom-operator"),
            ("Prometheus",          "http://localhost:9090",             "Aucun"),
            ("Alertmanager",        "http://localhost:9093",             "Aucun"),
            ("n8n",                 "http://localhost:5678",             "Aucun"),
            ("ArgoCD",              "https://localhost:8080",            "admin / Sre@Copilot2026"),
        ],
        col_widths=[2.0, 3.2, 2.0]
    )


# ─── Section 7 — Defense devant le jury ───────────────────────────────────────

def add_defense(doc):
    heading1(doc, "7.  Comment defendre le projet devant le jury")

    heading2(doc, "Les 5 messages cles a faire passer")
    make_table(doc,
        ["#", "Message", "Preuve concrete a montrer"],
        [
            ("1",
             "Le projet automatise vraiment la reponse aux incidents (de bout en bout, sans intervention humaine).",
             "Une demo : alerte → 3 phases → fix applique → commit Git → sync ArgoCD → dashboard."),
            ("2",
             "L'agent est intelligent (IA Mistral) mais aussi robuste (validator, fallback, garde-fous anti-boucle).",
             "Les 6 garde-fous du code (tableau Section 5). Montrer le code agent.py."),
            ("3",
             "100% open-source et 100% local — aucune dependance cloud, aucun cout API.",
             "ollama list (mistral), helm list (kube-prom + Kyverno), kubectl get nodes (KinD)."),
            ("4",
             "Securite by design : Kyverno audite tous les workloads en continu, RBAC strict.",
             "kubectl get clusterpolicies → 4 lignes Ready. kubectl get policyreport -n demo."),
            ("5",
             "Observabilite production-ready : metriques Prometheus + dashboard Grafana + dashboard custom.",
             "Ouvrir Grafana sur d/sre-copilot/. Montrer les 9 panels animes."),
        ],
        col_widths=[0.3, 3.0, 3.7]
    )

    heading2(doc, "Demo type en 5 minutes (script de presentation)")
    code_block(doc, [
        "MINUTE 0-1 : Contexte",
        "  'Ce projet est un agent IA SRE qui automatise la reponse aux incidents",
        "   Kubernetes. Il replace les outils Microsoft (Azure OpenAI) par",
        "   des outils open-source (Ollama, Mistral).'",
        "  → Montrer la roadmap (rapport/figures/roadmap.png)",
        "",
        "MINUTE 1-2 : Architecture",
        "  'L'agent suit un pipeline 3 phases : collecte kubectl, analyse",
        "   IA via Mistral, remediation GitOps. Tout est trace dans Git",
        "   et synchronise via ArgoCD.'",
        "  → Schema flux de bout en bout (Section 1)",
        "",
        "MINUTE 2-3 : Demo live",
        "  .\\scripts\\reset-demo.ps1",
        "  .\\scripts\\run-demo.ps1",
        "  → Le jury voit les logs defiler : Phase 1 (rapide), Phase 2 (Mistral),",
        "    Phase 3 (action). Total ~3 min.",
        "",
        "MINUTE 3-4 : Resultat dans la dashboard",
        "  → Ouvrir http://localhost:8001/dashboard",
        "  → Cliquer sur l'incident le plus recent : payload JSON, runbook,",
        "    classification, fix applique.",
        "",
        "MINUTE 4-5 : KPIs et securite",
        "  → Ouvrir Grafana sur d/sre-copilot/ → 9 panels : MTTR, success rate…",
        "  → kubectl get policyreport -n demo → Kyverno audite en continu.",
    ], label="SCRIPT 5 MINUTES")

    heading2(doc, "Questions probables du jury et reponses")
    make_table(doc,
        ["Question probable", "Reponse claire"],
        [
            ("Pourquoi Ollama et pas OpenAI ?",
             "Open-source, gratuit, 100% local (pas de fuite de logs clients), reproductible chez le client sans negocier d'API key. Le tradeoff est la lenteur (Mistral CPU = 3 min) mais c'est acceptable pour des incidents non-temps-reel."),
            ("Pourquoi 3 phases separees ?",
             "Decouplage : Phase 1 (kubectl) est deterministe et fiable. Phase 2 (IA) peut echouer mais on a un fallback keyword. Phase 3 (action) ne se declenche QUE si Phase 2 propose un fix valide. Approche plus robuste que LangChain ReAct qui melangeait tout."),
            ("Qu'est-ce qui empeche une boucle infinie agent ↔ ArgoCD ?",
             "ARGOCD_SYNC_AFTER_FIX = {'restart'} dans agent.py. Pour scale_to_zero, l'agent NE DECLENCHE PAS sync ArgoCD (sinon ArgoCD reapplique replicas=1 et annule notre remediation). Garde-fou explicite dans le code."),
            ("Comment l'agent sait quel fix appliquer ?",
             "Mistral est instruit dans le prompt de retourner un proposed_fix.type parmi : scale_to_zero (pour CrashLoopBackOff), restart (transient), ou none (manuel requis). Le code Python applique mecaniquement selon le type."),
            ("Que se passe-t-il si Mistral est hors ligne ?",
             "Le keyword_fallback prend le relai : il cherche des mots-cles ('crashloopbackoff', 'oomkilled'…) dans les logs kubectl. Pas aussi fin que Mistral mais l'agent reste fonctionnel."),
            ("Pourquoi Kyverno en mode Audit et pas Enforce ?",
             "En Enforce, Kyverno bloque la creation de tout pod non conforme — y compris crash-app (necessaire pour les demos). Audit genere des PolicyReports sans bloquer, ce qui permet de visualiser l'ampleur des violations avant de durcir progressivement."),
            ("Combien d'alertes simultanees l'agent peut-il gerer ?",
             "Sequentiellement : 1 a la fois grace au Semaphore(1). En theorie ~20/heure (avec Mistral CPU). En production avec GPU ou Mistral cloud, ~600/heure."),
            ("Comment passer en production ?",
             "Remplacer Ollama local par une API Mistral cloud (URL change, le reste reste identique). Activer Kyverno Enforce. Brancher Alertmanager sur le vrai Prometheus de prod (PagerDuty, Slack en sortie). Ajouter un repo GitHub real pour les commits."),
        ],
        col_widths=[2.5, 4.5]
    )

    heading2(doc, "Points forts a mettre en avant")
    bullet(doc, "Architecture decouplee (FastAPI + Ollama + ArgoCD + Kyverno + Prometheus = 5 outils standards de l'ecosysteme cloud-native).")
    bullet(doc, "Pipeline 3 phases deterministe — plus simple a expliquer et debugger qu'un ReAct LLM.")
    bullet(doc, "Observabilite end-to-end : tout est mesure, exposable, alertable.")
    bullet(doc, "GitOps loop complete : remediation → commit → ArgoCD sync. Boucle ITIL fermee.")
    bullet(doc, "Garde-fous explicites : validator, anti-loop, fallback, semaphore, RBAC strict.")
    bullet(doc, "100% reproductible : 1 commande (start-sre-copilot.ps1) recree TOUT.")

    callout(doc,
        "Conseil final : ne pas dire 'l'IA est magique'. Dire 'l'IA est UN OUTIL parmi d'autres dans un pipeline robuste'. Le jury appreciera la rigueur d'avoir mis des garde-fous deterministes autour du LLM.",
        HEX_GREEN_BG, label="ASTUCE DEFENSE")


# ─── Section 8 — Aide-memoire ────────────────────────────────────────────────

def add_cheatsheet(doc):
    doc.add_page_break()
    heading1(doc, "8.  Aide-memoire — commandes les plus utiles")

    heading2(doc, "Verifier l'etat du systeme")
    code_block(doc, [
        "# Etat du cluster",
        "kubectl get nodes",
        "kubectl get pods --all-namespaces",
        "",
        "# SRE Agent",
        "Invoke-WebRequest 'http://localhost:8001/health' | Select-Object -Exp Content",
        "Invoke-WebRequest 'http://localhost:8001/status' | Select-Object -Exp Content",
        "Invoke-WebRequest 'http://localhost:8001/api/stats' | Select-Object -Exp Content",
        "",
        "# Prometheus targets",
        "Invoke-WebRequest 'http://localhost:9090/api/v1/targets' | Select-Object -Exp Content",
        "",
        "# Kyverno policies actives",
        "kubectl get clusterpolicies",
        "kubectl get policyreport -n demo",
        "",
        "# ArgoCD application",
        "kubectl get application demo-app -n argocd",
    ], label="DIAGNOSTIC")

    heading2(doc, "Declencher une investigation manuellement")
    code_block(doc, [
        "# Methode 1 : appel direct (synchrone, attend la reponse)",
        "$resp = Invoke-WebRequest 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 360",
        "$resp.Content | ConvertFrom-Json",
        "",
        "# Methode 2 : via Alertmanager (passe par tout le pipeline)",
        "$now = (Get-Date -Format 'o'); $end = (Get-Date).AddHours(1).ToString('o')",
        "$body = @\"",
        "[{",
        "  \"labels\": { \"alertname\": \"PodRestartingTooMuch\", \"namespace\": \"demo\",",
        "                \"pod\": \"crash-app\", \"severity\": \"critical\" },",
        "  \"annotations\": { \"summary\": \"Test\", \"description\": \"Demo\" },",
        "  \"startsAt\": \"$now\", \"endsAt\": \"$end\"",
        "}]",
        "\"@",
        "Invoke-WebRequest 'http://localhost:9093/api/v2/alerts' -Method POST `",
        "    -Body $body -ContentType 'application/json'",
    ], label="DECLENCHER UN TEST")

    heading2(doc, "Voir ce qui se passe en direct")
    code_block(doc, [
        "# Logs SRE Agent en streaming",
        "kubectl logs -f deployment/sre-agent -n sre-agent",
        "",
        "# Events Kubernetes en streaming",
        "kubectl get events -n demo --sort-by='.lastTimestamp' -w",
        "",
        "# Commits Git dans le pod",
        "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
        "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
        "kubectl exec -n sre-agent $pod -- git -C /gitops-repo log --oneline",
    ], label="OBSERVABILITE LIVE")


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════

def main():
    out_path = os.path.join(os.path.dirname(__file__), "..", "SRE-Copilot-Architecture.docx")
    out_path = os.path.abspath(out_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)
    section.top_margin  = Cm(2.0); section.bottom_margin = Cm(2.0)
    doc.styles['Normal'].font.name = "Calibri"
    doc.styles['Normal'].font.size = Pt(11)

    add_cover(doc)
    add_overview(doc)
    add_tree(doc)
    add_scripts(doc)
    add_yaml(doc)
    add_python(doc)
    add_launch(doc)
    add_defense(doc)
    add_cheatsheet(doc)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("PFE IFS 013/26 — PwC DigiTech Cloud Native Operations | SRE Copilot | Architecture & Defense | 2026-06-10")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.italic = True; r.font.color.rgb = C_GRAY_LT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
