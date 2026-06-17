"""
Generates SRE-Copilot-Test-Validation.docx
Scenario de test A → Z structure par composant, avec :
  - Utilite de chaque etape
  - Outils utilises
  - Commande(s) exacte(s)
  - Resultat attendu

Run: python scripts/generate_test_validation.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Palette ────────────────────────────────────────────────────────────────────
C_NAVY     = RGBColor(0x1F, 0x38, 0x64)
C_BLUE     = RGBColor(0x2E, 0x75, 0xB6)
C_TEAL     = RGBColor(0x00, 0x70, 0x7F)
C_GREEN    = RGBColor(0x37, 0x56, 0x23)
C_ORANGE   = RGBColor(0xC5, 0x50, 0x0C)
C_PURPLE   = RGBColor(0x70, 0x30, 0xA0)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT     = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED    = RGBColor(0x6E, 0x7B, 0x8B)
C_GRAY_LT  = RGBColor(0xD6, 0xDC, 0xE4)

HEX_NAVY      = "1F3864"
HEX_TEAL      = "00707F"
HEX_BLUE      = "2E75B6"
HEX_TABLE_HDR = "1F3864"
HEX_TABLE_ALT = "DEEAF1"
HEX_CODE_BG   = "F2F2F2"
HEX_GREEN_BG  = "E2EFDA"
HEX_ORANGE_BG = "FFF2CC"
HEX_BLUE_BG   = "D9E1F2"
HEX_PURPLE_BG = "EBE4F2"

# Couleurs par phase
PHASE_COLORS = {
    1:  ("PRE-FLIGHT",          "595959"),
    2:  ("STACK STARTUP",       "1F3864"),
    3:  ("INFRASTRUCTURE",      "2E75B6"),
    4:  ("PIPELINE D'ALERTES",  "00707F"),
    5:  ("AGENT IA",            "375623"),
    6:  ("REMEDIATION",         "C5500C"),
    7:  ("SECURITE KYVERNO",    "7030A0"),
    8:  ("OBSERVABILITE",       "843C0C"),
    9:  ("ROBUSTESSE",          "4472C4"),
    10: ("CLEANUP",             "595959"),
}


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def set_cell_borders(table, color="D6DCE4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top','left','bottom','right','insideH','insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
                tcBorders.append(b)
            tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, color=None, size=10, font="Calibri",
              align=WD_ALIGN_PARAGRAPH.LEFT, mono=False):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]; p.clear(); p.alignment = align
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = "Courier New" if mono else font
    r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color


# ── Typo helpers ───────────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(16); r.bold = True; r.font.color.rgb = C_NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8'); bot.set(qn('w:space'), '2'); bot.set(qn('w:color'), HEX_NAVY)
    pBdr.append(bot); pPr.append(pBdr)

def body(doc, text, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = C_TEXT
    if italic: r.italic = True

def phase_banner(doc, phase_num, total_steps):
    name, hex_color = PHASE_COLORS[phase_num]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"  PHASE {phase_num} — {name}   ({total_steps} etape{'s' if total_steps > 1 else ''})")
    r.font.name = "Calibri"; r.font.size = Pt(15); r.bold = True; r.font.color.rgb = C_WHITE
    set_para_bg(p, hex_color)
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(6)

def step_box(doc, step_num, title, utility, tools, commands, expected):
    """
    Une etape complete avec : titre numerote, utilite, outils, commandes, resultat.
    Tout dans un tableau pour avoir une mise en page solide.
    """
    # Titre
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    rnum = p.add_run(f"#{step_num:02d}  ")
    rnum.font.name = "Calibri"; rnum.font.size = Pt(13); rnum.bold = True; rnum.font.color.rgb = C_TEAL
    rt = p.add_run(title)
    rt.font.name = "Calibri"; rt.font.size = Pt(13); rt.bold = True; rt.font.color.rgb = C_NAVY

    # Utilite + Outils dans un mini-tableau 2 colonnes
    t = doc.add_table(rows=2, cols=2); t.style = 'Table Grid'
    cell_text(t.rows[0].cells[0], "Utilite", bold=True, color=C_WHITE, size=9)
    set_cell_bg(t.rows[0].cells[0], HEX_NAVY)
    cell_text(t.rows[0].cells[1], utility, size=10)
    set_cell_bg(t.rows[0].cells[1], "FFFFFF")
    cell_text(t.rows[1].cells[0], "Outils", bold=True, color=C_WHITE, size=9)
    set_cell_bg(t.rows[1].cells[0], HEX_TEAL)
    cell_text(t.rows[1].cells[1], tools, size=10)
    set_cell_bg(t.rows[1].cells[1], "FFFFFF")
    for row in t.rows:
        row.cells[0].width = Inches(1.2); row.cells[1].width = Inches(5.5)
    set_cell_borders(t)

    # Commande(s) - bloc code monospace
    lp = doc.add_paragraph()
    lp.paragraph_format.space_before = Pt(6); lp.paragraph_format.space_after = Pt(0)
    lr = lp.add_run("  Commande")
    lr.font.name = "Calibri"; lr.font.size = Pt(9); lr.bold = True; lr.font.color.rgb = C_WHITE
    set_para_bg(lp, HEX_TEAL)

    for line in commands:
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.left_indent = Cm(0.3)
        cr = cp.add_run(line if line else " ")
        cr.font.name = "Courier New"; cr.font.size = Pt(9); cr.font.color.rgb = C_TEXT
        set_para_bg(cp, HEX_CODE_BG)

    # Resultat attendu en vert
    ep = doc.add_paragraph()
    ep.paragraph_format.space_before = Pt(4); ep.paragraph_format.space_after = Pt(8)
    er = ep.add_run(f"  Resultat attendu : {expected}")
    er.font.name = "Calibri"; er.font.size = Pt(10); er.bold = True; er.font.color.rgb = C_GREEN
    set_para_bg(ep, HEX_GREEN_BG)

def callout(doc, text, bg_hex, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(8); lp.paragraph_format.space_after = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"; lr.font.size = Pt(9); lr.bold = True; lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f"  {text}  ")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = C_TEXT
    set_para_bg(p, bg_hex)

def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(t.rows[0].cells[i], HEX_TABLE_HDR)
    for r, row in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else HEX_TABLE_ALT
        for c, val in enumerate(row):
            cell_text(t.rows[r+1].cells[c], val, size=10)
            set_cell_bg(t.rows[r+1].cells[c], bg)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    set_cell_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════════════════════════════════
# COUVERTURE
# ════════════════════════════════════════════════════════════════════════════

def add_cover(doc):
    bar = doc.add_paragraph(); set_para_bg(bar, HEX_NAVY)
    r = bar.add_run("  SRE Copilot — Test de Validation A à Z")
    r.font.name = "Calibri"; r.font.size = Pt(24); r.bold = True; r.font.color.rgb = C_WHITE

    sub = doc.add_paragraph(); set_para_bg(sub, HEX_TEAL)
    sr = sub.add_run("  Tester chaque composant, methodiquement, sans rien oublier")
    sr.font.name = "Calibri"; sr.font.size = Pt(13); sr.font.color.rgb = C_WHITE

    doc.add_paragraph()

    info = [
        ("Objectif",     "Valider que chaque brique du projet fonctionne et que l'integration end-to-end tient."),
        ("Format",       "10 phases, 38 etapes. Chaque etape a son utilite, ses outils, sa commande, son resultat attendu."),
        ("Duree totale", "~30-40 minutes (Mistral est lent : 3 min par investigation IA)."),
        ("Prerequis",    "Docker Desktop, Ollama avec Mistral, kubectl, kind, helm installes."),
        ("Date",         "2026-06-11"),
    ]
    t = doc.add_table(rows=len(info), cols=2); t.style = 'Table Grid'
    for i, (k, v) in enumerate(info):
        bg = HEX_TABLE_ALT if i % 2 == 0 else "FFFFFF"
        cell_text(t.rows[i].cells[0], k, bold=True, color=C_NAVY, size=11)
        cell_text(t.rows[i].cells[1], v, size=11)
        set_cell_bg(t.rows[i].cells[0], bg); set_cell_bg(t.rows[i].cells[1], bg)
    for row in t.rows:
        row.cells[0].width = Inches(1.4); row.cells[1].width = Inches(5.3)
    set_cell_borders(t)

    doc.add_paragraph()

    # Sommaire des phases
    body(doc, "Vue d'ensemble du scenario :")
    make_table(doc,
        ["Phase", "Composant", "Etapes", "Duree"],
        [
            ("1",  "Pre-flight (prerequis)",     "01-03", "1 min"),
            ("2",  "Stack startup",              "04-07", "1-2 min (cluster existant) / 15 min (nouveau)"),
            ("3",  "Infrastructure (par composant)", "08-13", "3 min"),
            ("4",  "Pipeline d'alertes",         "14-16", "2 min"),
            ("5",  "Agent IA end-to-end",        "17-20", "5 min (dont 3 min Mistral)"),
            ("6",  "Verification remediation",   "21-24", "1 min"),
            ("7",  "Securite Kyverno",           "25-28", "2 min"),
            ("8",  "Observabilite Grafana",      "29-31", "2 min"),
            ("9",  "Tests de robustesse",        "32-35", "5 min"),
            ("10", "Cleanup et reset",           "36-38", "1 min"),
        ],
        col_widths=[0.5, 2.6, 1.0, 2.6]
    )
    doc.add_page_break()


# ════════════════════════════════════════════════════════════════════════════
# PHASES
# ════════════════════════════════════════════════════════════════════════════

def add_phase_1(doc):
    phase_banner(doc, 1, 3)
    body(doc, "Avant de toucher au projet, on s'assure que les 3 prerequis externes sont en place. Si l'un manque, rien ne marchera.")

    step_box(doc, 1, "Verifier que Docker Desktop tourne",
        utility="KinD utilise Docker pour creer le cluster. Sans Docker, aucune autre etape ne fonctionnera.",
        tools="Docker CLI (docker info)",
        commands=["docker info | Select-String 'Server Version'"],
        expected="Une ligne 'Server Version: XX.X.X' apparait. Si erreur 'pipe' ou 'connection refused', demarrer Docker Desktop.")

    step_box(doc, 2, "Verifier qu'Ollama tourne et que Mistral est present",
        utility="L'agent IA appelle Ollama en local pour faire l'inference. Sans le modele Mistral charge, le fallback keyword sera utilise (degradation de qualite).",
        tools="Ollama CLI (ollama list) + HTTP (Invoke-WebRequest)",
        commands=[
            "ollama list",
            "# OU via API HTTP :",
            "Invoke-WebRequest 'http://localhost:11434/api/tags' | Select-Object -Exp Content",
        ],
        expected="Une ligne 'mistral:latest' apparait dans la liste. Si manquant : ollama pull mistral.")

    step_box(doc, 3, "Verifier les outils CLI installes",
        utility="Les scripts PowerShell appellent kubectl, kind, helm. Si l'un manque, start-sre-copilot.ps1 plantera silencieusement.",
        tools="kubectl, kind, helm",
        commands=[
            "kubectl version --client",
            "kind version",
            "helm version",
        ],
        expected="Chaque commande affiche une version (1.28+ pour kubectl, 0.20+ pour kind, 3.x pour helm).")


def add_phase_2(doc):
    phase_banner(doc, 2, 4)
    body(doc, "On lance la stack complete. Une fois faite, tout le reste du test passe par cette stack en cours d'execution.")

    step_box(doc, 4, "Lancer start-sre-copilot.ps1",
        utility="Script unique qui cree le cluster KinD (si absent), deploie monitoring + ArgoCD + Kyverno + SRE Agent + dashboards, injecte le token ArgoCD, demarre n8n et tous les port-forwards.",
        tools="PowerShell + scripts internes (start-sre-copilot.ps1 appelle generate-argocd-token.ps1)",
        commands=[
            "cd c:\\Users\\PC\\Desktop\\SRE-Copilot",
            ".\\scripts\\start-sre-copilot.ps1",
        ],
        expected="A la fin : 5 lignes 'OK' vertes pour les health checks (SRE Agent, Grafana, Prometheus, Alertmanager, n8n). Aucun pod en CrashLoopBackOff.")

    step_box(doc, 5, "Verifier les 3 noeuds du cluster KinD",
        utility="Si un noeud est NotReady, certains pods ne pourront pas etre schedules.",
        tools="kubectl",
        commands=["kubectl get nodes"],
        expected="3 lignes 'Ready' : sre-cluster-control-plane, sre-cluster-worker, sre-cluster-worker2.")

    step_box(doc, 6, "Verifier que tous les pods sont Running",
        utility="Si un pod cle est en CrashLoopBackOff (ex : SRE Agent), les tests suivants echoueront.",
        tools="kubectl",
        commands=[
            "kubectl get pods -n monitoring   # 8 pods attendus",
            "kubectl get pods -n argocd       # 7 pods attendus",
            "kubectl get pods -n sre-agent    # 1 pod attendu",
            "kubectl get pods -n kyverno      # 4 pods attendus",
        ],
        expected="Tous les pods en 'Running'. (argocd-applicationset-controller est desactive par start-sre-copilot.ps1 car la CRD ApplicationSet n'est pas livree avec install.yaml et on n'en a pas besoin).")

    step_box(doc, 7, "Verifier que les 6 port-forwards sont actifs",
        utility="Sans port-forwards, on ne peut pas tester les services depuis le host. Les scripts les lancent automatiquement mais ils peuvent tomber.",
        tools="PowerShell (Get-NetTCPConnection)",
        commands=[
            "3000,9090,9093,8080,8001,5678 | ForEach-Object {",
            "  $l = Get-NetTCPConnection -LocalPort $_ -State Listen -ErrorAction SilentlyContinue",
            "  if ($l) { \"OK   :$_\" } else { \"FAIL :$_\" }",
            "}",
        ],
        expected="6 lignes 'OK' pour les ports 3000 (Grafana), 9090 (Prometheus), 9093 (Alertmanager), 8080 (ArgoCD), 8001 (SRE Agent), 5678 (n8n).")


def add_phase_3(doc):
    phase_banner(doc, 3, 6)
    body(doc, "On verifie chaque service infrastructure individuellement avec un test specifique.")

    step_box(doc, 8, "SRE Agent — health + status",
        utility="Confirmer que l'API FastAPI repond et qu'Ollama est joignable depuis le pod.",
        tools="Invoke-WebRequest",
        commands=[
            "Invoke-WebRequest 'http://localhost:8001/health' | Select-Object -Exp Content",
            "Invoke-WebRequest 'http://localhost:8001/status' | Select-Object -Exp Content",
        ],
        expected="/health renvoie {\"status\":\"ok\"}. /status confirme 'ollama':'reachable' et 'model_available':true.")

    step_box(doc, 9, "Prometheus — verifier que sre-agent est scrape",
        utility="Si Prometheus ne scrape pas l'agent, les KPIs Grafana resteront vides. Le ServiceMonitor doit etre decouvert.",
        tools="API Prometheus (/api/v1/targets)",
        commands=[
            "$t = (Invoke-WebRequest 'http://localhost:9090/api/v1/targets').Content | ConvertFrom-Json",
            "$t.data.activeTargets | Where-Object { $_.labels.job -eq 'sre-agent-service' } | `",
            "    Select-Object scrapeUrl, health, lastScrape | Format-List",
        ],
        expected="1 target affiche : health='up', scrapeUrl='http://10.244.X.X:8000/metrics', lastScrape recent (<30s).")

    step_box(doc, 10, "Alertmanager — verifier la config (receivers)",
        utility="Le receiver 'n8n-webhook' doit pointer vers host.docker.internal:5678. Sans ca, les alertes n'iront pas a n8n.",
        tools="API Alertmanager v2 (/api/v2/status)",
        commands=[
            "$cfg = (Invoke-WebRequest 'http://localhost:9093/api/v2/status').Content | ConvertFrom-Json",
            "$cfg.config.original -split \"`n\" | Select-String 'n8n-webhook|url:' | Select-Object -First 5",
        ],
        expected="Bloc 'n8n-webhook' avec url='http://host.docker.internal:5678/webhook/sre-alert'.")

    step_box(doc, 11, "n8n — workflow actif",
        utility="Si le workflow n'est pas ACTIVE, les webhooks recus par n8n ne declencheront rien.",
        tools="Navigateur (UI n8n)",
        commands=[
            "Start-Process 'http://localhost:5678/workflows'",
            "# Dans l'UI : verifier le toggle 'Active' (vert) sur le workflow 'SRE Alert Pipeline'.",
            "# Cliquer sur le noeud 'Alertmanager Webhook' :",
            "#   - HTTP Method : POST",
            "#   - Path        : sre-alert",
            "#   - Respond     : Immediately",
        ],
        expected="Workflow actif avec 5 noeuds visibles, responseMode='Respond immediately'.")

    step_box(doc, 12, "ArgoCD — Application demo-app enregistree",
        utility="Si l'Application n'existe pas, l'agent ne pourra pas faire sync (HTTP 404).",
        tools="kubectl",
        commands=[
            "kubectl get application demo-app -n argocd",
        ],
        expected="1 ligne avec SYNC=OutOfSync ou Synced, HEALTH=Healthy ou Progressing. NAME='demo-app'.")

    step_box(doc, 13, "Kyverno — les 4 ClusterPolicies sont Ready",
        utility="Si une policy n'est pas Ready, les workloads ne seront pas audites contre celle-ci.",
        tools="kubectl",
        commands=["kubectl get clusterpolicies"],
        expected="4 lignes avec READY=True : disallow-privileged-containers, disallow-latest-tag, require-labels, require-resource-limits.")


def add_phase_4(doc):
    phase_banner(doc, 4, 3)
    body(doc, "On verifie que le pipeline Alertmanager → n8n → SRE Agent fonctionne en envoyant une alerte de test.")

    step_box(doc, 14, "Envoyer une alerte test a Alertmanager (API v2)",
        utility="Simule une alerte Prometheus sans attendre que crash-app crash. Permet de valider rapidement le routage.",
        tools="API Alertmanager v2 (POST /api/v2/alerts)",
        commands=[
            "$now = (Get-Date -Format 'o')",
            "$end = (Get-Date).AddHours(1).ToString('o')",
            "$body = @\"",
            "[{",
            "  \"labels\": {",
            "    \"alertname\": \"PodRestartingTooMuch\",",
            "    \"namespace\": \"demo\",",
            "    \"pod\":       \"crash-app\",",
            "    \"severity\":  \"critical\"",
            "  },",
            "  \"annotations\": { \"summary\": \"Test validation\", \"description\": \"E2E test\" },",
            "  \"startsAt\": \"$now\", \"endsAt\": \"$end\"",
            "}]",
            "\"@",
            "Invoke-WebRequest 'http://localhost:9093/api/v2/alerts' -Method POST `",
            "    -Body $body -ContentType 'application/json'",
        ],
        expected="StatusCode 200, body vide. L'alerte sera visible dans http://localhost:9093 (Alerts).")

    step_box(doc, 15, "Voir l'execution se declencher dans n8n",
        utility="Verifier que l'alerte a bien atteint n8n et qu'il a appele l'agent.",
        tools="UI n8n",
        commands=[
            "Start-Process 'http://localhost:5678/executions'",
            "# Une nouvelle execution doit apparaitre, statut 'Running' puis 'Succeeded'.",
            "# Duree typique : 3-4 minutes (Mistral est le goulot).",
        ],
        expected="1 execution recente, tous les 5 noeuds verts a la fin. Cliquer dessus montre la reponse de l'agent.")

    step_box(doc, 16, "Verifier que l'agent a recu l'alerte (logs)",
        utility="Confirmation finale : la chaine AM → n8n → agent est intacte.",
        tools="kubectl logs",
        commands=[
            "kubectl logs deployment/sre-agent -n sre-agent --tail=30 | Select-String 'Webhook received|Phase'",
        ],
        expected="Lignes 'Webhook received — N firing alert(s)' et 'Queuing investigation'. Puis Phase 1, 2, 3 qui s'enchainent.")


def add_phase_5(doc):
    phase_banner(doc, 5, 4)
    body(doc, "On declenche une vraie investigation IA et on observe les 3 phases en direct.")

    step_box(doc, 17, "Reset propre + redeployer crash-app",
        utility="Partir d'un etat connu : pas de pod trainant, alertes silencieuses levees, crash-app a 1 replica.",
        tools="reset-demo.ps1 + foreach PowerShell + kubectl scale",
        commands=[
            ".\\scripts\\reset-demo.ps1",
            "",
            "# Lever les silences Alertmanager (foreach + @() pour deroulement fiable)",
            "$silences = @((Invoke-WebRequest 'http://localhost:9093/api/v2/silences').Content | ConvertFrom-Json)",
            "foreach ($s in $silences) {",
            "    if ($s.status.state -eq 'active') {",
            "        Invoke-WebRequest \"http://localhost:9093/api/v2/silence/$($s.id)\" -Method DELETE | Out-Null",
            "        Write-Host \"  Silence supprime : $($s.id)\"",
            "    }",
            "}",
            "",
            "# Redeployer crash-app",
            "kubectl scale deployment crash-app -n demo --replicas=1",
            "Start-Sleep -Seconds 90",
            "kubectl get pods -n demo -l app=crash-app",
        ],
        expected="crash-app-XXX en CrashLoopBackOff avec 3+ restarts apres 90 secondes.")

    step_box(doc, 18, "Declencher /test-crash-loop (timeout 360s)",
        utility="Endpoint qui declenche le pipeline 3 phases sur crash-app. On utilise Invoke-WebRequest (et non run-demo.ps1) pour recuperer le JSON de retour et inspecter programmatiquement les champs (classification, fix_type, argocd_sync_triggered). run-demo.ps1 streame les logs visuellement mais ne donne pas le JSON.",
        tools="Invoke-WebRequest (timeout long) — alternative visuelle : .\\scripts\\run-demo.ps1",
        commands=[
            "$start = Get-Date",
            "$resp = Invoke-WebRequest 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 360",
            "$dur = (Get-Date) - $start",
            "$r = $resp.Content | ConvertFrom-Json",
            "\"Duree : $([math]::Round($dur.TotalSeconds,1))s\"",
            "$r | Select-Object classification, auto_remediation_applied, duration_seconds, `",
            "    git_pr_url, argocd_sync_triggered | Format-List",
        ],
        expected="Apres 150-250s : classification=CrashLoopBackOff ou ConfigError, auto_remediation_applied=True, git_pr_url=local:branch/fix/..., argocd_sync_triggered=True (si fix_type=restart).")

    step_box(doc, 19, "Verifier les 3 phases dans les logs",
        utility="Confirmer que le pipeline est bien sequentiel (Phase 1 deterministe → Phase 2 IA → Phase 3 action).",
        tools="kubectl logs",
        commands=[
            "kubectl logs deployment/sre-agent -n sre-agent --tail=40 | Select-String 'Phase' | Select-Object -Last 8",
        ],
        expected="Sequence : [Phase 1] kubectl collection → [Phase 1] Collected 4 sections → [Phase 2] LLM JSON-mode analysis → [Phase 2] Classification: X | Fix: Y → [Phase 3] AUTO_REMEDIATE=true → [Phase 3] Complete.")

    step_box(doc, 20, "Verifier que l'incident apparait dans la dashboard",
        utility="La dashboard est le visage utilisateur final. Si l'incident n'y apparait pas, la persistance est cassee.",
        tools="Navigateur",
        commands=[
            "Start-Process 'http://localhost:8001/dashboard'",
            "# Le nouvel incident doit apparaitre en haut de la liste 'Incidents recents'.",
            "# Cliquer dessus pour voir le payload JSON complet.",
        ],
        expected="L'incident est visible dans la dashboard : classification, severity 'critical', badge 'Remedie' vert, timestamp recent.")


def add_phase_6(doc):
    phase_banner(doc, 6, 4)
    body(doc, "On verifie que la remediation a tenu : pod scale, commit Git fait, ArgoCD appele, metriques mises a jour.")

    step_box(doc, 21, "crash-app a 0 replicas (si fix=scale_to_zero) ou redemarre (si restart)",
        utility="Confirmer que l'action de Phase 3 a vraiment ete appliquee au cluster.",
        tools="kubectl get deployment",
        commands=[
            "kubectl get deployment crash-app -n demo",
            "kubectl get pods -n demo -l app=crash-app",
        ],
        expected="Si scale_to_zero : READY 0/0. Si restart : un nouveau pod (AGE recent < 5 min).")

    step_box(doc, 22, "Commit Git documente dans le pod",
        utility="Tracabilite : chaque incident produit un commit dans /gitops-repo/incidents/.",
        tools="kubectl exec + git",
        commands=[
            "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
            "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
            "kubectl exec -n sre-agent $pod -- git -C /gitops-repo log --oneline | Select-Object -First 5",
            "kubectl exec -n sre-agent $pod -- ls /gitops-repo/incidents/demo/",
        ],
        expected="1+ commit recent du type 'fix(demo): auto-remediate <classification> on crash-app'. 1+ fichier YAML d'incident.")

    step_box(doc, 23, "ArgoCD a recu l'ordre sync (si fix_type=restart)",
        utility="Verifier que la boucle GitOps est complete : agent → Git → ArgoCD → cluster.",
        tools="kubectl get application",
        commands=[
            "kubectl get application demo-app -n argocd",
            "# Pour voir l'historique de syncs :",
            "kubectl get application demo-app -n argocd -o jsonpath='{.status.operationState.phase}'",
        ],
        expected="SYNC=Synced (ou OutOfSync apres scale_to_zero). operationState.phase='Succeeded'.")

    step_box(doc, 24, "Metriques Prometheus mises a jour",
        utility="Mission 5 : les compteurs doivent s'incrementer apres chaque investigation.",
        tools="Endpoint /metrics + PromQL",
        commands=[
            "$m = (Invoke-WebRequest 'http://localhost:8001/metrics').Content",
            "$m -split \"`n\" | Select-String '^sre_alerts_total|^sre_remediated_total|^sre_argocd_sync_total|^sre_investigation_duration_seconds_count'",
        ],
        expected="Compteurs >= 1 sur sre_alerts_total, sre_remediated_total. sre_investigation_duration_seconds_count >= 1.")


def add_phase_7(doc):
    phase_banner(doc, 7, 4)
    body(doc, "On valide Mission 4 : Kyverno detecte les violations sur des pods de test.")

    step_box(doc, 25, "Deployer les 4 pods de violation",
        utility="Chaque pod viole 1 politique. En mode Audit, ils sont quand meme crees mais generent des PolicyReports avec FAIL.",
        tools="kubectl apply",
        commands=["kubectl apply -f kyverno/test-violations/"],
        expected="4 pods crees : violation-privileged, violation-latest, violation-no-labels, violation-no-limits.")

    step_box(doc, 26, "Lire les PolicyReports generes",
        utility="Le reports-controller cree 1 PolicyReport par ressource auditee.",
        tools="kubectl get policyreport",
        commands=[
            "Start-Sleep -Seconds 5",
            "kubectl get policyreport -n demo",
        ],
        expected="Plusieurs PolicyReports : 4 pour les pods violation-* (FAIL=1, PASS=4) + 2 pour crash-app Deployment/ReplicaSet.")

    step_box(doc, 27, "Compter les violations par politique",
        utility="Statistiques utiles pour la defense devant le jury.",
        tools="kubectl + ConvertFrom-Json + Group-Object",
        commands=[
            "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
            "    Select-Object -Exp items | Select-Object -Exp results -ErrorAction SilentlyContinue | `",
            "    Where-Object { $_.result -eq 'fail' } | Group-Object -Property policy | `",
            "    Select-Object Name, Count | Format-Table -AutoSize",
        ],
        expected="4 lignes : 1+ violation par politique. crash-app contribue meme aux compteurs (latest tag, owner, cpu).")

    step_box(doc, 28, "Verifier les details d'une violation",
        utility="Comprendre EXACTEMENT pourquoi Kyverno a echoue la validation.",
        tools="kubectl describe policyreport",
        commands=[
            "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
            "    Select-Object -Exp items | Where-Object { $_.scope.name -eq 'violation-privileged' } | `",
            "    Select-Object -Exp results | Where-Object { $_.result -eq 'fail' } | `",
            "    Select-Object policy, message | Format-List",
        ],
        expected="1 ligne avec policy='disallow-privileged-containers' et message explicite ('Les conteneurs privilegies sont interdits...').")


def add_phase_8(doc):
    phase_banner(doc, 8, 3)
    body(doc, "On valide Mission 5 : dashboard Grafana auto-importe avec les bons KPIs.")

    step_box(doc, 29, "Verifier le dashboard est importe dans Grafana",
        utility="Le sidecar Grafana doit decouvrir le ConfigMap labellise et l'importer.",
        tools="API Grafana",
        commands=[
            "$basic = [Convert]::ToBase64String([Text.Encoding]::ASCII.GetBytes('admin:prom-operator'))",
            "$h = @{ Authorization = \"Basic $basic\" }",
            "(Invoke-WebRequest 'http://localhost:3000/api/search?query=SRE' -Headers $h).Content | `",
            "    ConvertFrom-Json | Select-Object title, uid, url",
        ],
        expected="1 dashboard 'SRE Copilot — KPIs' avec uid='sre-copilot'.")

    step_box(doc, 30, "Ouvrir le dashboard et verifier les 9 panels",
        utility="Validation visuelle : tous les panels doivent afficher des donnees coherentes apres les investigations.",
        tools="Navigateur",
        commands=[
            "Start-Process 'http://localhost:3000/d/sre-copilot/sre-copilot-kpis'",
            "# Login : admin / prom-operator",
            "# Verifier les 9 panels :",
            "#   1. Alertes totales (stat)        2. Taux de remediation auto (%)",
            "#   3. MTTR median (s)               4. ArgoCD sync OK (stat)",
            "#   5. Keyword fallbacks (stat)      6. Taux d'alertes par classification (time series)",
            "#   7. Repartition classifications   8. Duree p50/p95 (time series)",
            "#   9. Remediations par type+result (bar)",
        ],
        expected="Dashboard avec 9 panels remplis (au moins 1 valeur > 0 sur chaque, sauf si pas d'investigations).")

    step_box(doc, 31, "Tester une requete PromQL manuelle",
        utility="Valider que Prometheus est bien la source de verite pour les KPIs.",
        tools="API Prometheus + PromQL",
        commands=[
            "$q = 'sum by (classification) (sre_alerts_total)'",
            "$url = 'http://localhost:9090/api/v1/query?query=' + [uri]::EscapeDataString($q)",
            "(Invoke-WebRequest $url).Content | ConvertFrom-Json | Select-Object -Exp data | `",
            "    Select-Object -Exp result | Format-Table -AutoSize",
        ],
        expected="1 ou plusieurs lignes : une par classification (CrashLoopBackOff, ConfigError…) avec leur valeur.")


def add_phase_9(doc):
    phase_banner(doc, 9, 4)
    body(doc, "Tests de robustesse : valider les garde-fous (fallback, deduplication, filtre namespace, anti-boucle).")

    step_box(doc, 32, "Garde-fou 1 : fallback keyword (Mistral hors ligne)",
        utility="Verifier que l'agent reste fonctionnel si Ollama est down.",
        tools="Gestionnaire des taches + /test-crash-loop",
        commands=[
            "# Arreter Ollama dans le Gestionnaire des taches (terminer ollama.exe)",
            "# Puis declencher une investigation :",
            "$resp = Invoke-WebRequest 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 60",
            "$r = $resp.Content | ConvertFrom-Json",
            "$r.classification",
            "$r.investigation_log | Where-Object { $_ -match 'fallback' }",
            "# Redemarrer Ollama apres le test : ollama serve",
        ],
        expected="classification valide (CrashLoopBackOff/ConfigError/Unknown), une ligne 'llm:fallback' dans investigation_log. Temps < 10 secondes.")

    step_box(doc, 33, "Garde-fou 2 : deduplication n8n (envoi double)",
        utility="Verifier qu'une alerte deja traitee dans les 10 dernieres minutes est ignoree (evite la rafale).",
        tools="API Alertmanager v2",
        commands=[
            "# Envoyer la meme alerte 2 fois a 5 secondes d'intervalle",
            "$now = (Get-Date -Format 'o'); $end = (Get-Date).AddHours(1).ToString('o')",
            "$body = @\"",
            "[{\"labels\":{\"alertname\":\"PodRestartingTooMuch\",\"namespace\":\"demo\",\"pod\":\"crash-app\",\"severity\":\"critical\"},",
            "\"annotations\":{\"summary\":\"Dedup test\"},\"startsAt\":\"$now\",\"endsAt\":\"$end\"}]",
            "\"@",
            "Invoke-WebRequest 'http://localhost:9093/api/v2/alerts' -Method POST -Body $body -ContentType 'application/json'",
            "Start-Sleep -Seconds 5",
            "Invoke-WebRequest 'http://localhost:9093/api/v2/alerts' -Method POST -Body $body -ContentType 'application/json'",
            "# Dans http://localhost:5678/executions : seule la 1ere doit etre traitee.",
        ],
        expected="Une seule execution n8n 'Succeeded' (la 2eme est bloquee par filter-dedup).")

    step_box(doc, 34, "Garde-fou 3 : filtre namespace (alerte hors demo)",
        utility="Verifier que les alertes du namespace 'monitoring' (ou autre) ne declenchent PAS l'agent.",
        tools="API Alertmanager v2 + UI n8n",
        commands=[
            "$now = (Get-Date -Format 'o'); $end = (Get-Date).AddHours(1).ToString('o')",
            "$body = @\"",
            "[{\"labels\":{\"alertname\":\"PodRestartingTooMuch\",\"namespace\":\"monitoring\",\"pod\":\"x\",\"severity\":\"critical\"},",
            "\"annotations\":{\"summary\":\"Test\"},\"startsAt\":\"$now\",\"endsAt\":\"$end\"}]",
            "\"@",
            "Invoke-WebRequest 'http://localhost:9093/api/v2/alerts' -Method POST -Body $body -ContentType 'application/json'",
        ],
        expected="StatusCode 200 mais AUCUNE execution n8n declenchee (Alertmanager filtre par namespace=demo).")

    step_box(doc, 35, "Garde-fou 4 : pas de boucle ArgoCD apres scale_to_zero",
        utility="Le bug le plus subtil que le code adresse : verifier qu'apres un scale_to_zero, ArgoCD ne re-cree PAS le pod.",
        tools="kubectl get pods (observation prolongee)",
        commands=[
            "# Apres une investigation avec fix_type=scale_to_zero :",
            "for ($i = 0; $i -lt 6; $i++) {",
            "    Start-Sleep -Seconds 10",
            "    \"($($i*10+10)s) replicas=$(kubectl get deployment crash-app -n demo -o jsonpath='{.spec.replicas}')\"",
            "}",
        ],
        expected="replicas=0 pendant les 60 secondes. crash-app NE revient PAS automatiquement (preuve que ARGOCD_SYNC_AFTER_FIX={'restart'} marche).")


def add_phase_10(doc):
    phase_banner(doc, 10, 3)
    body(doc, "Remettre l'environnement propre pour le prochain test ou pour rendre au repos.")

    step_box(doc, 36, "Reset standard (garde l'historique)",
        utility="Stopper la boucle d'alertes en cours, scale crash-app a 0. L'historique des incidents dans la dashboard reste consultable.",
        tools="reset-demo.ps1",
        commands=[".\\scripts\\reset-demo.ps1"],
        expected="3 etapes 'OK' : jobs PowerShell tues, silence Alertmanager cree, crash-app a 0.")

    step_box(doc, 37, "Reset complet avec effacement de l'historique",
        utility="Utile avant une demo officielle : la dashboard sera vide, comme neuve.",
        tools="reset-demo.ps1 -ClearHistory",
        commands=[".\\scripts\\reset-demo.ps1 -ClearHistory"],
        expected="Idem que step 36 + vidage de /gitops-repo/incidents/responses.jsonl. Dashboard vide au prochain refresh.")

    step_box(doc, 38, "Arret complet (fin de journee)",
        utility="Liberer les ports du host et reduire la charge CPU. Le cluster KinD reste en place pour reprise rapide.",
        tools="stop-sre-copilot.ps1",
        commands=[
            ".\\scripts\\stop-sre-copilot.ps1",
            "# Le script demande si on veut aussi supprimer le cluster (par defaut : non).",
        ],
        expected="6 port-forwards fermes, 6 ports liberes. Cluster KinD toujours actif (re-execution start-sre-copilot.ps1 = 30s).")


# ─── Checklist finale ─────────────────────────────────────────────────────────

def add_checklist(doc):
    doc.add_page_break()
    heading1(doc, "Checklist finale — 38 etapes")

    body(doc, "Coche au fur et a mesure pour t'assurer que rien n'a ete oublie.")

    items = [
        ("1",  "Docker Desktop tourne"),
        ("2",  "Ollama + Mistral disponibles"),
        ("3",  "kubectl, kind, helm installes"),
        ("4",  "start-sre-copilot.ps1 termine sans erreur"),
        ("5",  "3 noeuds KinD Ready"),
        ("6",  "Tous les pods Running (monitoring, argocd, sre-agent, kyverno)"),
        ("7",  "6 port-forwards actifs"),
        ("8",  "SRE Agent /health et /status OK"),
        ("9",  "Prometheus scrape sre-agent (health=up)"),
        ("10", "Alertmanager receiver n8n-webhook present"),
        ("11", "Workflow n8n actif"),
        ("12", "Application ArgoCD demo-app enregistree"),
        ("13", "4 ClusterPolicies Kyverno Ready"),
        ("14", "Alerte test envoyee a Alertmanager (200)"),
        ("15", "Execution n8n visible"),
        ("16", "Logs SRE Agent : Webhook received + Phase 1/2/3"),
        ("17", "Reset + crash-app redeploye en CrashLoopBackOff"),
        ("18", "/test-crash-loop termine avec classification valide"),
        ("19", "3 phases enchainees dans les logs"),
        ("20", "Incident visible dans la dashboard /dashboard"),
        ("21", "crash-app a 0 replicas OU redemarre"),
        ("22", "Commit Git visible dans /gitops-repo/"),
        ("23", "ArgoCD operationState.phase=Succeeded"),
        ("24", "Compteurs Prometheus >= 1"),
        ("25", "4 pods de violation deployes"),
        ("26", "PolicyReports generes"),
        ("27", "Comptage par politique fonctionne"),
        ("28", "Detail violation explicite (kubectl describe)"),
        ("29", "Dashboard Grafana 'SRE Copilot — KPIs' present"),
        ("30", "9 panels remplis"),
        ("31", "PromQL manuelle renvoie des donnees"),
        ("32", "Garde-fou fallback keyword (Mistral down)"),
        ("33", "Garde-fou deduplication n8n"),
        ("34", "Garde-fou filtre namespace"),
        ("35", "Garde-fou anti-boucle (scale_to_zero tient 60s)"),
        ("36", "reset-demo.ps1 propre"),
        ("37", "reset-demo.ps1 -ClearHistory vide la dashboard"),
        ("38", "stop-sre-copilot.ps1 ferme les ports"),
    ]

    t = doc.add_table(rows=1 + len(items), cols=3)
    t.style = 'Table Grid'
    for i, h in enumerate(["#", "Validation", "Statut"]):
        cell_text(t.rows[0].cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(t.rows[0].cells[i], HEX_TABLE_HDR)
    for r, (num, desc) in enumerate(items):
        bg = "FFFFFF" if r % 2 == 0 else HEX_TABLE_ALT
        cell_text(t.rows[r+1].cells[0], num, bold=True, color=C_NAVY, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(t.rows[r+1].cells[1], desc, size=9)
        cell_text(t.rows[r+1].cells[2], "[ ]", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
        for ci in range(3):
            set_cell_bg(t.rows[r+1].cells[ci], bg)
    for row in t.rows:
        row.cells[0].width = Inches(0.4)
        row.cells[1].width = Inches(5.5)
        row.cells[2].width = Inches(0.6)
    set_cell_borders(t)

    doc.add_paragraph()
    callout(doc,
        "Si TOUTES les cases sont cochees : le projet est valide bout en bout. Tu peux le defendre les yeux fermes devant le jury.",
        HEX_GREEN_BG, label="VALIDATION COMPLETE")


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════

def main():
    out_path = os.path.join(os.path.dirname(__file__), "..", "SRE-Copilot-Test-Validation.docx")
    out_path = os.path.abspath(out_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.0); section.right_margin = Cm(2.0)
    section.top_margin  = Cm(2.0); section.bottom_margin = Cm(2.0)
    doc.styles['Normal'].font.name = "Calibri"
    doc.styles['Normal'].font.size = Pt(11)

    add_cover(doc)
    add_phase_1(doc)
    add_phase_2(doc)
    add_phase_3(doc)
    add_phase_4(doc)
    add_phase_5(doc)
    add_phase_6(doc)
    add_phase_7(doc)
    add_phase_8(doc)
    add_phase_9(doc)
    add_phase_10(doc)
    add_checklist(doc)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("PFE IFS 013/26 — PwC DigiTech Cloud Native Operations | SRE Copilot | Test Validation A à Z | 2026-06-11")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.italic = True; r.font.color.rgb = C_GRAY_LT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
