"""
Generates SRE-Copilot-Explication.docx
A simple, pedagogical document that explains:
  - What the SRE Copilot does (in plain language)
  - What happens step by step during an investigation
  - The difference between CrashLoopBackOff and ConfigError
  - The ArgoCD fix we just applied
  - Why Mistral is slow

Run: python scripts/generate_explication.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Color palette ──────────────────────────────────────────────────────────────
C_NAVY     = RGBColor(0x1F, 0x38, 0x64)
C_BLUE     = RGBColor(0x2E, 0x75, 0xB6)
C_TEAL     = RGBColor(0x00, 0x70, 0x7F)
C_GREEN    = RGBColor(0x37, 0x56, 0x23)
C_ORANGE   = RGBColor(0xC5, 0x50, 0x0C)
C_RED      = RGBColor(0x9C, 0x00, 0x00)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_TEXT     = RGBColor(0x2C, 0x2C, 0x2C)
C_MUTED    = RGBColor(0x6E, 0x7B, 0x8B)
C_GRAY_LT  = RGBColor(0xD6, 0xDC, 0xE4)

HEX_NAVY      = "1F3864"
HEX_TEAL      = "00707F"
HEX_BLUE      = "2E75B6"
HEX_TABLE_HDR = "1F3864"
HEX_TABLE_ALT = "DEEAF1"
HEX_CODE_BG   = "F2F2F2"
HEX_GREEN_BG  = "E2EFDA"
HEX_ORANGE_BG = "FFF2CC"
HEX_BLUE_BG   = "D9E1F2"
HEX_NEUTRAL_BG = "F5F2EC"


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_cell_borders(table, color="D6DCE4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                b = OxmlElement(f'w:{side}')
                b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), '4')
                b.set(qn('w:space'), '0'); b.set(qn('w:color'), color)
                tcBorders.append(b)
            tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, color=None, size=10, font="Calibri",
              align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = font; r.font.size = Pt(size); r.bold = bold
    if color: r.font.color.rgb = color
    return r


# ── Typography ─────────────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(18); r.bold = True; r.font.color.rgb = C_NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '8')
    bot.set(qn('w:space'), '2'); bot.set(qn('w:color'), HEX_NAVY)
    pBdr.append(bot); pPr.append(pBdr)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(13); r.bold = True; r.font.color.rgb = C_BLUE
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.bold = True; r.font.color.rgb = C_TEAL
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = C_TEXT
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.6)
    r = p.add_run(text)
    r.font.name = "Calibri"; r.font.size = Pt(11); r.font.color.rgb = C_TEXT
    return p


def callout(doc, text, bg_hex, label=None, label_color=None):
    """Boîte colorée pour mettre en avant une idée importante."""
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after  = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"; lr.font.size = Pt(9); lr.bold = True
        lr.font.color.rgb = label_color if label_color else C_WHITE
        set_para_bg(lp, HEX_TEAL)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"  {text}  ")
    r.font.name = "Calibri"; r.font.size = Pt(10); r.font.color.rgb = C_TEXT
    set_para_bg(p, bg_hex)
    return p


def code_block(doc, lines, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after  = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"; lr.font.size = Pt(9); lr.bold = True; lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.3)
        r = p.add_run(line if line else " ")
        r.font.name = "Courier New"; r.font.size = Pt(9); r.font.color.rgb = C_TEXT
        set_para_bg(p, HEX_CODE_BG)

    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(4)


def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell_text(t.rows[0].cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(t.rows[0].cells[i], HEX_TABLE_HDR)
    for r, row in enumerate(rows):
        bg = "FFFFFF" if r % 2 == 0 else HEX_TABLE_ALT
        for c, val in enumerate(row):
            cell_text(t.rows[r+1].cells[c], val, size=10)
            set_cell_bg(t.rows[r+1].cells[c], bg)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    set_cell_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ────────────────────────────────────────────────────────────────────────────────
# CONTENU
# ────────────────────────────────────────────────────────────────────────────────

def add_cover(doc):
    bar = doc.add_paragraph()
    set_para_bg(bar, HEX_NAVY)
    r = bar.add_run("  SRE Copilot — Explications")
    r.font.name = "Calibri"; r.font.size = Pt(26); r.bold = True; r.font.color.rgb = C_WHITE

    sub = doc.add_paragraph()
    set_para_bg(sub, HEX_TEAL)
    sr = sub.add_run("  Comment ça marche, vraiment, sans jargon")
    sr.font.name = "Calibri"; sr.font.size = Pt(13); sr.font.color.rgb = C_WHITE

    doc.add_paragraph()

    info = [
        ("Pour qui", "Toi-même quand tu reviens sur le projet — et le jury PFE."),
        ("Objectif", "Expliquer ce qui se passe en arrière-plan pendant une démo, et pourquoi."),
        ("Format",   "Analogies, pas de jargon inutile, exemples concrets."),
        ("Date",     "2026-06-09"),
    ]
    t = doc.add_table(rows=len(info), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(info):
        bg = HEX_TABLE_ALT if i % 2 == 0 else "FFFFFF"
        cell_text(t.rows[i].cells[0], k, bold=True, color=C_NAVY, size=11)
        cell_text(t.rows[i].cells[1], v, size=11)
        set_cell_bg(t.rows[i].cells[0], bg)
        set_cell_bg(t.rows[i].cells[1], bg)
    for row in t.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.2)
    set_cell_borders(t)


# ── Section 1 : Ce que fait le SRE Copilot ────────────────────────────────────

def add_section_overview(doc):
    heading1(doc, "1.  Le SRE Copilot en 60 secondes")

    body(doc, "Imagine un SRE (l'ingénieur qui surveille la prod) qui ne dort jamais. Quand une alerte tombe à 3h du matin, normalement c'est lui qui :")
    bullet(doc, "Reçoit la notification (PagerDuty, Slack, etc.)")
    bullet(doc, "Se connecte au cluster, lance kubectl pour comprendre le problème")
    bullet(doc, "Lit les logs, regarde les events, cherche la cause")
    bullet(doc, "Applique un fix temporaire pour stopper le saignement")
    bullet(doc, "Documente l'incident, prévient l'équipe, etc.")

    body(doc, "Le SRE Copilot fait tout ça à sa place, en 3 minutes au lieu de 30, sans le réveiller.")

    callout(doc,
        "Le SRE Copilot n'est PAS un chatbot. C'est un agent autonome qui : (1) reçoit des alertes Prometheus, (2) enquête tout seul avec kubectl, (3) demande à Mistral 7B de classifier le problème, (4) applique automatiquement un fix sûr, (5) documente le tout dans Git.",
        HEX_BLUE_BG, label="EN UNE PHRASE", label_color=C_WHITE)

    heading2(doc, "Pourquoi c'est utile pour PwC ?")
    body(doc, "Dans un cabinet d'audit IT, on intervient chez des clients qui n'ont pas forcément d'équipe SRE 24/7. Le SRE Copilot fait office de premier répondant automatique : il évite que les incidents banals (un container qui crash) traînent des heures avant qu'un humain les regarde.")


# ── Section 2 : Vie d'une alerte ───────────────────────────────────────────────

def add_section_lifecycle(doc):
    heading1(doc, "2.  Vie d'une alerte — du début à la fin")

    body(doc, "Quand un pod commence à crasher, voici tout ce qui se passe, dans l'ordre :")

    make_table(doc,
        ["#", "Qui", "Quoi", "Combien de temps"],
        [
            ("1", "Le pod crash-app",
             "Le container exécute 'exit 1'. Kubernetes le redémarre. Il re-crash. Boucle.",
             "≈ 1-2 min pour atteindre 3+ restarts"),
            ("2", "Prometheus",
             "Scrape les métriques. Détecte kube_pod_container_status_restarts_total > 3.",
             "Évalué toutes les 30s"),
            ("3", "Alertmanager",
             "Reçoit l'alerte 'PodRestartingTooMuch'. Vérifie la route : namespace=demo + severity=critical → receiver 'n8n-webhook'.",
             "< 1s"),
            ("4", "n8n (webhook)",
             "Reçoit le POST. Répond HTTP 200 immédiatement (responseMode: onReceived). Lance le workflow en arrière-plan.",
             "< 1s pour répondre"),
            ("5", "n8n (filter & dedup)",
             "Vérifie si la même alerte n'a pas été traitée dans les 10 dernières minutes. Si oui, ignore. Sinon, continue.",
             "< 100ms"),
            ("6", "n8n (call agent)",
             "POST vers http://host.docker.internal:8001/alert avec le payload de l'alerte.",
             "Démarre le timer"),
            ("7", "SRE Agent — Phase 1",
             "Lance plusieurs commandes kubectl : get pods, logs --previous, describe pod, get events. Récupère ~4 sections de données.",
             "≈ 2 secondes"),
            ("8", "SRE Agent — Phase 2",
             "Envoie ces données à Mistral 7B (Ollama). Mistral classifie l'incident en JSON : 'CrashLoopBackOff', root_cause, runbook 3 étapes, fix proposé.",
             "≈ 150-250 secondes (Mistral CPU)"),
            ("9", "SRE Agent — Phase 3 / action",
             "Exécute le fix recommandé. Pour CrashLoopBackOff : kubectl scale deployment crash-app --replicas=0 (stoppe le saignement).",
             "< 1s"),
            ("10", "SRE Agent — Phase 3 / git",
             "Crée un commit dans /gitops-repo/incidents/demo/crash-app-TIMESTAMP.yaml. Ce commit documente : qui, quoi, quand, pourquoi, quelle action.",
             "< 1s"),
            ("11", "SRE Agent — Phase 3 / ArgoCD",
             "Appelle POST /api/v1/applications/demo-app/sync. ArgoCD reçoit l'ordre, va lire le Git, ré-applique le manifest.",
             "≈ 5s"),
            ("12", "SRE Agent — réponse",
             "Renvoie un JSON au n8n avec tous les détails (classification, runbook, métriques). Persiste aussi dans /gitops-repo/incidents/responses.jsonl.",
             "< 1s"),
            ("13", "n8n (build report)",
             "Formate un rapport Markdown lisible avec tous les détails. Visible dans l'historique n8n.",
             "< 100ms"),
            ("14", "Dashboard",
             "L'incident apparaît automatiquement dans http://localhost:8001/dashboard (auto-refresh 30s).",
             "Au prochain refresh"),
        ],
        col_widths=[0.3, 1.5, 4.0, 1.4]
    )

    callout(doc,
        "Au final, environ 3 minutes après le crash, le pod est stoppé, l'incident est documenté dans Git, ArgoCD le sait, et tu peux voir le rapport complet dans une jolie dashboard. Pendant ces 3 minutes, tu n'as RIEN fait.",
        HEX_GREEN_BG, label="LE MAGIQUE", label_color=C_GREEN)


# ── Section 3 : Les 3 phases ───────────────────────────────────────────────────

def add_section_phases(doc):
    heading1(doc, "3.  Les 3 phases de l'agent, vues de l'intérieur")

    body(doc, "Le code de l'agent (sre-agent/agent.py) est divisé en 3 phases bien distinctes. C'est important parce qu'avant on avait essayé une approche 'tout-en-un' (LangChain ReAct) qui était instable. La séparation en 3 phases rend tout prévisible.")

    heading2(doc, "Phase 1 — Collecte (déterministe, rapide)")
    body(doc, "Le mot 'déterministe' veut dire : pas d'IA, pas de hasard. C'est juste du code Python qui appelle kubectl. Toujours le même comportement, toujours les mêmes données.")
    body(doc, "Pourquoi avant l'IA ? Parce que l'IA a besoin de données concrètes pour analyser. On les collecte d'abord, ensuite on les donne à Mistral.")
    code_block(doc, [
        "# Ce qu'on fait en Phase 1",
        "data = {}",
        "data['pod_status']        = kubectl get pods -n demo -o wide",
        "data['pod_logs']          = kubectl logs crash-app -n demo --tail=50 --previous",
        "data['pod_describe']      = kubectl describe pod crash-app -n demo",
        "data['namespace_events']  = kubectl get events -n demo --field-selector type!=Normal",
    ], label="PHASE 1 EN PYTHON")

    heading2(doc, "Phase 2 — Analyse IA (Mistral en mode JSON)")
    body(doc, "On envoie les 4 sections collectées à Mistral 7B. On lui demande de renvoyer EXACTEMENT ce format JSON :")
    code_block(doc, [
        "{",
        "  \"classification\":   \"CrashLoopBackOff\",",
        "  \"root_cause\":       \"Le container exit 1 immédiatement après le démarrage...\",",
        "  \"runbook\": [",
        "    {\"order\": 1, \"action\": \"...\", \"command\": \"kubectl ...\"},",
        "    ...",
        "  ],",
        "  \"proposed_fix\": {",
        "    \"type\":        \"scale_to_zero\",",
        "    \"description\": \"...\"",
        "  }",
        "}",
    ], label="CE QUE MISTRAL DOIT RENVOYER")
    body(doc, "Pourquoi imposer le format JSON ? Parce que sans contrainte, l'IA peut renvoyer du texte libre, des markdown, etc. Avec format='json' on force Mistral à respecter le schéma, donc on peut parser sans bug.")

    callout(doc,
        "Si Mistral plante ou répond du JSON invalide, l'agent passe en 'keyword fallback' : il cherche des mots-clés dans les logs (crashloopbackoff, oomkilled, imagepullbackoff). Pas aussi intelligent que Mistral, mais l'agent reste fonctionnel.",
        HEX_ORANGE_BG, label="ROBUSTESSE", label_color=C_ORANGE)

    heading2(doc, "Phase 3 — Remédiation (3 actions en chaîne)")
    body(doc, "Si Mistral a proposé un fix applicable ET que AUTO_REMEDIATE=true (variable d'environnement), l'agent enchaîne 3 actions :")

    make_table(doc,
        ["Action", "Quoi", "Pourquoi"],
        [
            ("scale ou restart",
             "kubectl scale crash-app --replicas=0 (ou rollout restart)",
             "Stoppe le 'saignement' immédiatement. Le client ne voit plus l'erreur."),
            ("git commit",
             "Crée un fichier YAML dans /gitops-repo/incidents/demo/ + git commit",
             "Trace écrite, auditable, qui peut être pushé sur GitHub."),
            ("ArgoCD sync",
             "POST /api/v1/applications/demo-app/sync à ArgoCD",
             "Notifie le système GitOps : 'Va relire Git et applique l'état souhaité'."),
        ],
        col_widths=[1.6, 3.4, 2.2]
    )


# ── Section 4 : CrashLoopBackOff vs ConfigError ────────────────────────────────

def add_section_classifications(doc):
    heading1(doc, "4.  CrashLoopBackOff vs ConfigError — c'est quoi la différence ?")

    body(doc, "Dans la dashboard tu as vu ces deux classifications. Voici comment Mistral fait la différence et pourquoi le fix appliqué change.")

    heading2(doc, "CrashLoopBackOff")
    body(doc, "Le pod démarre, puis CRASH (exit code != 0) immédiatement. Kubernetes le redémarre, il crash encore, Kubernetes attend de plus en plus longtemps (backoff exponentiel). D'où le nom 'crash + loop + back-off'.")
    body(doc, "Indices que Mistral cherche dans les données kubectl :")
    bullet(doc, "État du pod = 'CrashLoopBackOff' (visible dans 'kubectl get pods')")
    bullet(doc, "Restart count > 0 et qui augmente")
    bullet(doc, "Logs --previous montrent une erreur fatale (exit 1, segfault, exception Python)")
    bullet(doc, "Events: 'Back-off restarting failed container'")

    callout(doc,
        "Pour crash-app : c'est exactement le cas. Le container exécute 'exit 1' au démarrage, donc Kubernetes le relance en boucle. Le fix appliqué : scale_to_zero (replicas=0) pour stopper la boucle.",
        HEX_NEUTRAL_BG, label="EXEMPLE CONCRET", label_color=C_NAVY)

    heading2(doc, "ConfigError")
    body(doc, "Le pod ne crash pas forcément, mais il a un problème de configuration : variable d'environnement manquante, ConfigMap mal monté, secret introuvable, port mal défini, etc.")
    body(doc, "Indices que Mistral cherche :")
    bullet(doc, "Events: 'CreateContainerConfigError', 'InvalidImageName', 'CreateContainerError'")
    bullet(doc, "Logs: erreurs de type 'connection refused', 'env var not set', 'config not found'")
    bullet(doc, "Pod peut être en état 'CreateContainerConfigError' (jamais démarré)")

    callout(doc,
        "ConfigError = problème entre la config et le code. Le code est probablement OK, mais quelque chose ne lui est pas fourni correctement. Un simple restart peut résoudre (si la config a été corrigée entre-temps).",
        HEX_NEUTRAL_BG, label="EXEMPLE CONCRET", label_color=C_NAVY)

    heading2(doc, "Tableau récapitulatif")
    make_table(doc,
        ["Critère", "CrashLoopBackOff", "ConfigError"],
        [
            ("Symptôme visible",
             "Pod en boucle de restart, état CrashLoopBackOff",
             "Pod jamais démarré OU démarre puis erreur de config"),
            ("Cause typique",
             "Bug dans le code, exit immédiat, segfault",
             "Variable d'env manquante, secret introuvable, ConfigMap mal monté"),
            ("Visible dans les logs",
             "Erreurs fatales, stack traces, exit codes non-zéro",
             "'connection refused', 'env not set', erreurs de démarrage"),
            ("Fix appliqué par l'agent",
             "scale_to_zero (replicas=0) → stoppe la boucle",
             "restart (rollout restart) → redéploie proprement"),
            ("Pourquoi ce fix ?",
             "Le code est cassé. Inutile de relancer 1000 fois, ça refera la même chose. On stoppe, un humain corrige le code, ensuite on remet replicas=1.",
             "Souvent transitoire (DNS, network). Un restart peut suffire à récupérer."),
            ("Que fait l'humain après",
             "Lire les logs --previous, identifier le bug, corriger le code, redéployer.",
             "Vérifier ConfigMap/Secret/Env, corriger si nécessaire."),
        ],
        col_widths=[1.8, 2.6, 2.8]
    )

    heading2(doc, "Autres classifications possibles")
    body(doc, "Mistral peut aussi renvoyer :")
    bullet(doc, "OOMKilled — le pod a été tué par le kernel parce qu'il dépasse sa limite mémoire. Fix: restart + augmenter la limite.")
    bullet(doc, "ImagePullBackOff — Kubernetes n'arrive pas à télécharger l'image (tag inexistant, registry HS, secret manquant). Fix: aucun (humain requis).")
    bullet(doc, "ResourceExhausted — le cluster n'a plus assez de CPU/RAM. Fix: aucun (humain requis).")
    bullet(doc, "NetworkError — problème réseau (DNS, service mesh). Fix: aucun automatique.")
    bullet(doc, "Unknown — Mistral n'a pas su classifier. Fix: aucun, investigation manuelle requise.")


# ── Section 5 : Le fix ArgoCD ──────────────────────────────────────────────────

def add_section_argocd_fix(doc):
    heading1(doc, "5.  Le fix ArgoCD du jour — pourquoi 'non' avant, 'oui' maintenant")

    body(doc, "Dans la dashboard tu voyais 'ArgoCD : non' pour chaque incident. Voici exactement ce qui se passait et ce qu'on a corrigé.")

    heading2(doc, "Le bug, étape par étape")

    make_table(doc,
        ["#", "Ce que faisait l'agent", "Pourquoi ça plantait"],
        [
            ("1",
             "POST https://argocd-server.../api/v1/applications/demo-app/sync",
             "ArgoCD répond HTTP 404. Pourquoi ? Parce qu'AUCUNE Application nommée 'demo-app' n'était enregistrée dans ArgoCD. ArgoCD ne savait pas ce que c'est."),
            ("2",
             "Le code Python tente alors un fallback : appeler le binaire 'argocd' (CLI)",
             "Mais le binaire 'argocd' n'est PAS installé dans l'image Docker du SRE Agent. Seuls kubectl et git le sont. → Errno 2: No such file or directory."),
            ("3",
             "L'agent log un WARNING et continue",
             "Phase 3 ne fait que 2/3 de la boucle GitOps (scale OK, git commit OK, mais ArgoCD sync = false)."),
        ],
        col_widths=[0.3, 3.4, 3.5]
    )

    heading2(doc, "Le fix — 3 corrections")

    heading3(doc, "Correction 1 : Créer l'Application ArgoCD 'demo-app'")
    body(doc, "On a ajouté un fichier argocd/demo-app-application.yaml qui dit à ArgoCD :")
    bullet(doc, "Surveille le repo GitHub github.com/GhadaBenMansour/SRE-Copilot")
    bullet(doc, "Lis uniquement le fichier demo-app.yaml de la branche main")
    bullet(doc, "Quand on te demande 'sync', applique ce manifest dans le namespace demo")
    body(doc, "Maintenant, quand l'agent appelle POST /sync demo-app, ArgoCD existe pour de vrai et accepte l'ordre.")

    heading3(doc, "Correction 2 : Donner les permissions RBAC au compte sre-copilot")
    body(doc, "ArgoCD a son propre système de permissions (RBAC) indépendant de Kubernetes. Le compte sre-copilot existait avec la capacité 'apiKey' (peut générer un token permanent), mais n'avait AUCUNE permission sur les Applications.")
    code_block(doc, [
        "# Avant : pas de policy → 403 Forbidden",
        "",
        "# Après : on autorise get / sync / update",
        "p, sre-copilot, applications, get,    */*, allow",
        "p, sre-copilot, applications, sync,   */*, allow",
        "p, sre-copilot, applications, update, */*, allow",
    ], label="POLICY RBAC AJOUTÉE")

    heading3(doc, "Correction 3 : Redémarrer argocd-server")
    body(doc, "Le serveur ArgoCD met en cache la policy RBAC au démarrage. Sans redémarrage, le RBAC mis à jour n'est pas pris en compte → toujours 403.")
    code_block(doc, [
        "kubectl rollout restart deployment argocd-server -n argocd",
    ], label="REDÉMARRAGE")

    heading2(doc, "Résultat — la boucle GitOps complète")

    callout(doc,
        "Maintenant, quand l'agent finit une investigation : (1) il scale crash-app à 0 [STOP saignement], (2) il commit dans /gitops-repo [TRACE], (3) il dit à ArgoCD 'redéploie demo-app' [GitOps]. Les 3 étapes du loop GitOps sont alors complètes.",
        HEX_GREEN_BG, label="ÉTAT FINAL", label_color=C_GREEN)


# ── Section 6 : Pourquoi Mistral est lent ─────────────────────────────────────

def add_section_mistral_speed(doc):
    heading1(doc, "6.  Pourquoi Mistral met 3 minutes ?")

    body(doc, "C'est la question qui revient tout le temps. Réponse courte : Mistral 7B sur CPU, c'est lent. Réponse longue :")

    heading2(doc, "Ce que fait Mistral à chaque appel")
    bullet(doc, "Reçoit le prompt (notre alerte + les données kubectl) = ~1500 tokens")
    bullet(doc, "Génère la réponse JSON token par token = ~500 tokens")
    bullet(doc, "Chaque token demande un calcul matriciel sur ~7 milliards de paramètres")
    bullet(doc, "Sur CPU (sans GPU) : ~3-5 tokens par seconde")
    bullet(doc, "500 tokens × ~5 sec/token = 100-250 secondes au total")

    heading2(doc, "Comparaison")
    make_table(doc,
        ["Configuration", "Temps Mistral 7B", "Coût"],
        [
            ("Notre setup (CPU laptop)",       "150-250 secondes",  "0 € (local)"),
            ("GPU RTX 4090 (24 GB)",            "5-10 secondes",     "≈ 2000 € matériel"),
            ("Cloud API (Anthropic Claude)",    "1-2 secondes",      "≈ 0.01 € / appel"),
            ("Mistral via Mistral AI (cloud)",  "1-2 secondes",      "≈ 0.001 € / appel"),
        ],
        col_widths=[2.8, 2.2, 2.0]
    )

    heading2(doc, "Pourquoi on accepte la lenteur ?")
    body(doc, "Pour le PFE PwC, les choix techniques étaient :")
    bullet(doc, "100% local (pas de dépendance cloud, pas de fuite de logs clients)")
    bullet(doc, "Open-source (Mistral, pas Azure OpenAI)")
    bullet(doc, "Gratuit (pas de carte bancaire)")
    bullet(doc, "Reproductible (le même setup chez le client ou en pré-prod)")
    body(doc, "En prod, on remplacerait Mistral 7B par un modèle hébergé cloud (latence 1-2s). L'architecture du SRE Copilot ne change pas, juste l'URL Ollama vers une API distante.")

    callout(doc,
        "Pour la démo PFE : le timeout Invoke-WebRequest doit être réglé à 360 secondes (6 min). En option 'fire-and-forget' avec Start-Job, on ne bloque pas le terminal et on regarde les logs en direct — c'est plus pédagogique pour expliquer au jury ce qui se passe.",
        HEX_BLUE_BG, label="ASTUCE DÉMO", label_color=C_NAVY)


# ── Section 7 : Récapitulatif ─────────────────────────────────────────────────

def add_section_summary(doc):
    heading1(doc, "7.  Ce que tu peux montrer au jury")

    body(doc, "À la fin de la démo, voici ce que le jury aura vu :")

    make_table(doc,
        ["Ce qu'ils voient", "Ce que ça démontre"],
        [
            ("Cluster Kubernetes opérationnel (3 nœuds, 25+ pods)",
             "Maîtrise de l'infrastructure cloud-native"),
            ("Pipeline Alertmanager → n8n → SRE Agent",
             "Orchestration d'événements, intégration multi-outils"),
            ("Investigation IA avec Mistral 7B en JSON mode",
             "Application concrète d'un LLM open-source à un problème d'opérations"),
            ("Auto-remédiation kubectl scale --replicas=0",
             "Décision déterministe basée sur la classification IA"),
            ("Commit Git automatique dans /gitops-repo/",
             "Traçabilité, auditabilité, principe GitOps"),
            ("ArgoCD sync vers github.com/.../SRE-Copilot",
             "Boucle GitOps complète (alerte → fix → Git → sync)"),
            ("Politiques Kyverno qui auditent les workloads",
             "Sécurité as code, conformité automatique"),
            ("Dashboard /dashboard avec stats agrégées",
             "Observabilité, MTTR mesuré, runbook par incident"),
        ],
        col_widths=[3.4, 3.5]
    )

    callout(doc,
        "Le SRE Copilot répond exactement au sujet IFS 013/26 : 'remédiation GitOps pilotée par un agent IA orchestré avec n8n'. Chaque mot de l'intitulé est implémenté concrètement, avec des outils open-source, en local, sans dépendance cloud.",
        HEX_GREEN_BG, label="LIEN AVEC LE SUJET PFE", label_color=C_GREEN)


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    out_path = os.path.join(os.path.dirname(__file__), "..", "SRE-Copilot-Explication.docx")
    out_path = os.path.abspath(out_path)

    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    doc.styles['Normal'].font.name = "Calibri"
    doc.styles['Normal'].font.size = Pt(11)

    add_cover(doc)
    add_section_overview(doc)
    add_section_lifecycle(doc)
    add_section_phases(doc)
    add_section_classifications(doc)
    add_section_argocd_fix(doc)
    add_section_mistral_speed(doc)
    add_section_summary(doc)

    # Footer
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    r = p.add_run("PFE IFS 013/26 — PwC DigiTech Cloud Native Operations | SRE Copilot | Explications | 2026-06-09")
    r.font.name = "Calibri"; r.font.size = Pt(8); r.italic = True; r.font.color.rgb = C_GRAY_LT
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
