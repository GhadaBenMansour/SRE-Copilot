"""
Generates SRE-Copilot-Scenario-E2E.docx
End-to-end test scenario covering Missions 1 -> 4 in sequence.

Run: python scripts/generate_scenario_e2e.py
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ── Color palette ──────────────────────────────────────────────────────────────
C_NAVY     = RGBColor(0x1F, 0x38, 0x64)
C_BLUE     = RGBColor(0x2E, 0x75, 0xB6)
C_TEAL     = RGBColor(0x00, 0x70, 0x7F)
C_GREEN    = RGBColor(0x37, 0x56, 0x23)
C_ORANGE   = RGBColor(0xC5, 0x50, 0x0C)
C_PURPLE   = RGBColor(0x70, 0x30, 0xA0)
C_WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY_TXT = RGBColor(0x40, 0x40, 0x40)
C_GRAY_LT  = RGBColor(0xD6, 0xDC, 0xE4)
C_CODE_TXT = RGBColor(0x1F, 0x1F, 0x1F)

HEX_NAVY      = "1F3864"
HEX_TEAL      = "00707F"
HEX_TABLE_HDR = "1F3864"
HEX_TABLE_ALT = "DEEAF1"
HEX_CODE_BG   = "F2F2F2"
HEX_GREEN_BG  = "E2EFDA"
HEX_ORANGE_BG = "FFF2CC"
HEX_BLUE_BG   = "D9E1F2"
HEX_PURPLE_BG = "EBE4F2"

# Mission colors (cover banner per mission section)
MISSION_HEX = {
    1: "1F3864",   # navy
    2: "00707F",   # teal
    3: "375623",   # green
    4: "7030A0",   # purple
    5: "C5500C",   # orange — Mission 5 (Métriques)
}


# ── XML helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_para_bg(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_cell_borders(table, color="D6DCE4"):
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), color)
                tcBorders.append(border)
            tcPr.append(tcBorders)


def cell_text(cell, text, bold=False, color=None, size=10, font="Calibri",
              align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return run


# ── Typography helpers ─────────────────────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text.upper())
    run.font.name = "Calibri"
    run.font.size = Pt(15)
    run.bold = True
    run.font.color.rgb = C_NAVY
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), HEX_NAVY)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def body(doc, text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = color if color else C_GRAY_TXT
    return p


def mission_banner(doc, mission_num, name, description):
    """Big colored banner for the start of each mission section."""
    color_hex = MISSION_HEX[mission_num]
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run(f"  MISSION {mission_num} — {name.upper()}")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.bold = True
    run.font.color.rgb = C_WHITE
    set_para_bg(p, color_hex)

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(10)
    run2 = p2.add_run(f"  {description}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.italic = True
    run2.font.color.rgb = C_WHITE
    set_para_bg(p2, color_hex)


def step_title(doc, step_num, title):
    """Numbered step title with badge."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    # Step number badge
    r_num = p.add_run(f" {step_num:02d} ")
    r_num.font.name = "Calibri"
    r_num.font.size = Pt(11)
    r_num.bold = True
    r_num.font.color.rgb = C_WHITE
    # We need shading on the run — use a hack: write the badge as a separate paragraph
    # Actually, since we can't shade a run easily, let's just style the title

    # Clear and rewrite as styled title with number prefix in color
    p._p.clear()
    pPr = p._p.get_or_add_pPr()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r_num = p.add_run(f"#{step_num:02d}  ")
    r_num.font.name = "Calibri"
    r_num.font.size = Pt(13)
    r_num.bold = True
    r_num.font.color.rgb = C_TEAL
    r_title = p.add_run(title)
    r_title.font.name = "Calibri"
    r_title.font.size = Pt(13)
    r_title.bold = True
    r_title.font.color.rgb = C_NAVY
    return p


def why_line(doc, why):
    """Italic 'why' line under the step title."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.5)
    r = p.add_run(f"Pourquoi : {why}")
    r.font.name = "Calibri"
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = C_GRAY_TXT
    return p


def code_block(doc, lines, label=None):
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.space_before = Pt(6)
        lp.paragraph_format.space_after  = Pt(0)
        lr = lp.add_run(f"  {label}")
        lr.font.name = "Calibri"
        lr.font.size = Pt(9)
        lr.bold = True
        lr.font.color.rgb = C_WHITE
        set_para_bg(lp, HEX_TEAL)

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.3)
        run = p.add_run(line if line else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = C_CODE_TXT
        set_para_bg(p, HEX_CODE_BG)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after  = Pt(4)


def expected_result(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(f"  ✓  Résultat attendu : {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.bold = True
    run.font.color.rgb = C_GREEN
    set_para_bg(p, HEX_GREEN_BG)


def info_box(doc, text, bg_hex=HEX_BLUE_BG, icon="ℹ"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    run = p.add_run(f"  {icon}  {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = C_GRAY_TXT
    set_para_bg(p, bg_hex)
    return p


def make_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell_text(hdr.cells[i], h, bold=True, color=C_WHITE, size=10)
        set_cell_bg(hdr.cells[i], HEX_TABLE_HDR)
    for r_idx, row in enumerate(rows):
        tr = t.rows[r_idx + 1]
        bg = "FFFFFF" if r_idx % 2 == 0 else HEX_TABLE_ALT
        for c_idx, val in enumerate(row):
            cell_text(tr.cells[c_idx], val, size=9)
            set_cell_bg(tr.cells[c_idx], bg)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    set_cell_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ── Cover ──────────────────────────────────────────────────────────────────────

def add_cover(doc):
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(0)
    bar.paragraph_format.space_after  = Pt(0)
    set_para_bg(bar, HEX_NAVY)
    r = bar.add_run("  SRE COPILOT")
    r.font.name = "Calibri"; r.font.size = Pt(28); r.bold = True; r.font.color.rgb = C_WHITE

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after  = Pt(0)
    set_para_bg(sub, HEX_TEAL)
    sr = sub.add_run("  SCÉNARIO DE TEST END-TO-END | MISSIONS 1 → 4")
    sr.font.name = "Calibri"; sr.font.size = Pt(13); sr.bold = True; sr.font.color.rgb = C_WHITE

    doc.add_paragraph()

    info_rows = [
        ("Projet",       "PFE IFS 013/26 — PwC DigiTech Cloud Native Operations"),
        ("Objectif",     "Démontrer le fonctionnement complet de l'agent SRE Copilot"),
        ("Périmètre",    "Infrastructure → Alertes → Agent IA → Sécurité Kyverno"),
        ("Durée totale", "~20 minutes pour le scénario complet (cluster déjà démarré)"),
        ("Date",         "2026-06-03"),
    ]
    t = doc.add_table(rows=len(info_rows), cols=2)
    t.style = 'Table Grid'
    for i, (k, v) in enumerate(info_rows):
        bg = HEX_TABLE_ALT if i % 2 == 0 else "FFFFFF"
        cell_text(t.rows[i].cells[0], k, bold=True, color=C_NAVY, size=10)
        cell_text(t.rows[i].cells[1], v, size=10)
        set_cell_bg(t.rows[i].cells[0], bg)
        set_cell_bg(t.rows[i].cells[1], bg)
    for row in t.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.2)
    set_cell_borders(t)

    doc.add_paragraph()


def add_overview(doc):
    heading1(doc, "Vue d'ensemble du scénario")

    body(doc, "Ce scénario teste les 4 missions complétées dans l'ordre logique : démarrage de l'infrastructure, déclenchement d'une alerte, investigation et auto-remédiation par l'agent IA, puis vérification que les politiques Kyverno auditent correctement les workloads.")

    make_table(doc,
        ["#", "Mission", "Ce qu'on teste", "Étapes"],
        [
            ("1", "Infrastructure",       "Cluster KinD + stack monitoring + ArgoCD opérationnels", "01-04"),
            ("2", "Pipeline d'alertes",    "Alertmanager → n8n route correctement les alertes",       "05-08"),
            ("3", "Agent IA + Remédiation","Pipeline 3 phases + dashboard centralisée",                "09-14"),
            ("4", "Sécurité Kyverno",      "Détection automatique des violations",                     "15-19"),
            ("5", "Métriques + Grafana",   "/metrics + ServiceMonitor + dashboard auto",               "20-22"),
            ("✓", "Validation finale",     "Tous les composants fonctionnent ensemble",                "23-24"),
        ],
        col_widths=[0.3, 1.7, 4.0, 0.7]
    )

    info_box(doc, "Avant de commencer : Docker Desktop doit être démarré et Ollama doit tourner avec le modèle Mistral. Les commandes sont en PowerShell.")


# ════════════════════════════════════════════════════════════════════════════
# MISSION 1 — Infrastructure
# ════════════════════════════════════════════════════════════════════════════

def add_mission_1(doc):
    mission_banner(doc, 1, "Infrastructure KinD", "Vérifier que tous les composants de base sont opérationnels.")

    # Step 1
    step_title(doc, 1, "Lancer le script de démarrage")
    why_line(doc, "Le script automatise cluster + Prometheus + ArgoCD + n8n + port-forwards.")
    code_block(doc, [
        "cd c:\\Users\\PC\\Desktop\\SRE-Copilot",
        ".\\scripts\\start-sre-copilot.ps1",
    ], label="COMMANDE")
    expected_result(doc, "5 lignes vertes 'OK' à la fin (Grafana, Prometheus, Alertmanager, ArgoCD, SRE Agent)")

    # Step 2
    step_title(doc, 2, "Vérifier les nœuds Kubernetes")
    why_line(doc, "Le cluster KinD doit avoir 1 control-plane + 2 workers en Ready.")
    code_block(doc, [
        "kubectl get nodes",
    ], label="COMMANDE")
    expected_result(doc, "3 lignes 'Ready' : sre-cluster-control-plane, sre-cluster-worker, sre-cluster-worker2")

    # Step 3
    step_title(doc, 3, "Vérifier les pods système")
    why_line(doc, "Tous les composants critiques doivent être Running.")
    code_block(doc, [
        "kubectl get pods -n monitoring   # Prometheus + Grafana + Alertmanager",
        "kubectl get pods -n argocd       # ArgoCD",
        "kubectl get pods -n sre-agent    # SRE Agent",
        "kubectl get pods -n kyverno      # 4 controllers Kyverno",
    ], label="COMMANDE")
    expected_result(doc, "Tous les pods en Running. (argocd-applicationset-controller est desactive volontairement par le start script — CRD manquante, inutilise dans le projet)")

    # Step 4
    step_title(doc, 4, "Vérifier Ollama et le modèle Mistral")
    why_line(doc, "L'agent IA dépend d'Ollama pour l'inférence. Sans Mistral, fallback keyword.")
    code_block(doc, [
        "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
        "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
        "",
        "kubectl exec -n sre-agent $pod -- python3 -c \"",
        "import urllib.request, json",
        "req = urllib.request.Request('http://host.docker.internal:11434/api/tags')",
        "with urllib.request.urlopen(req, timeout=5) as r:",
        "    print('Models:', [m['name'] for m in json.loads(r.read()).get('models',[])])",
        "\"",
    ], label="COMMANDE")
    expected_result(doc, "Models: ['mistral:latest'] (ou similaire avec un autre tag)")


# ════════════════════════════════════════════════════════════════════════════
# MISSION 2 — Pipeline d'alertes
# ════════════════════════════════════════════════════════════════════════════

def add_mission_2(doc):
    mission_banner(doc, 2, "Pipeline d'alertes", "Vérifier qu'une alerte Alertmanager arrive bien jusqu'au SRE Agent via n8n.")

    # Step 5
    step_title(doc, 5, "Vérifier la configuration Alertmanager (receivers)")
    why_line(doc, "Le receiver 'n8n-webhook' doit pointer vers http://host.docker.internal:5678/webhook/sre-alert.")
    code_block(doc, [
        "# Via API ou UI Alertmanager",
        "Invoke-WebRequest 'http://localhost:9093/api/v2/status' | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty config",
        "",
        "# Ou en navigateur : http://localhost:9093 → Status → Config",
    ], label="COMMANDE")
    expected_result(doc, "Bloc 'receivers' montre 'n8n-webhook' avec l'URL host.docker.internal:5678")

    # Step 6
    step_title(doc, 6, "Vérifier le workflow n8n")
    why_line(doc, "Le workflow doit être actif avec 5 nœuds et responseMode=onReceived.")
    code_block(doc, [
        "Start-Process 'http://localhost:5678/workflows'",
        "",
        "# Dans l'UI n8n :",
        "# 1. Onglet 'Workflows' → ouvrir le workflow qui contient le nœud 'Alertmanager Webhook'",
        "# 2. Vérifier le toggle 'Active' en haut à droite (vert)",
        "# 3. Cliquer sur le nœud webhook-trigger :",
        "#    - HTTP Method : POST",
        "#    - Path        : sre-alert",
        "#    - Respond     : Immediately (= onReceived)",
    ], label="COMMANDE")
    expected_result(doc, "Workflow ACTIVE avec 5 nœuds en série : Webhook → Filter & Dedup → Call SRE Agent → Parse → Report")

    # Step 7
    step_title(doc, 7, "Envoyer une alerte test via Alertmanager (API v2)")
    why_line(doc, "Simule une alerte Prometheus qui passe par tout le pipeline.")
    info_box(doc, "Important : Alertmanager v1 supprimée depuis v0.27.0. On utilise /api/v2/alerts avec un tableau JSON. Le here-string PowerShell (@\"...\"@) est la forme la plus lisible — pas de backslash hell.", HEX_ORANGE_BG, icon="⚠")
    code_block(doc, [
        "$now = (Get-Date -Format 'o')",
        "$end = (Get-Date).AddHours(1).ToString('o')",
        "",
        "$body = @\"",
        "[{",
        "  \"labels\": {",
        "    \"alertname\": \"PodRestartingTooMuch\",",
        "    \"namespace\": \"demo\",",
        "    \"pod\":       \"crash-app\",",
        "    \"severity\":  \"critical\"",
        "  },",
        "  \"annotations\": {",
        "    \"summary\":     \"E2E test\",",
        "    \"description\": \"Scenario via API v2\"",
        "  },",
        "  \"startsAt\": \"$now\",",
        "  \"endsAt\":   \"$end\"",
        "}]",
        "\"@",
        "",
        "Invoke-WebRequest -Uri 'http://localhost:9093/api/v2/alerts' `",
        "    -Method POST -Body $body -ContentType 'application/json'",
    ], label="COMMANDE")
    expected_result(doc, "Status : 200 OK | Body vide (Alertmanager accepte sans corps de réponse)")
    info_box(doc, "Note : le `\"@` de fermeture du here-string DOIT être en colonne 0 (sans indentation), sinon PowerShell renvoie une erreur de parsing.", HEX_BLUE_BG)

    # Step 8
    step_title(doc, 8, "Vérifier l'exécution dans n8n")
    why_line(doc, "Une nouvelle exécution doit apparaître dans Executions → Running, puis Succeeded.")
    code_block(doc, [
        "# Ouvrir l'onglet Executions du workflow",
        "Start-Process 'http://localhost:5678/executions'",
        "",
        "# Ou depuis l'UI :",
        "# Workflows → ouvrir le workflow → bouton 'Executions' (icône horloge)",
        "# La nouvelle exécution doit s'afficher en haut :",
        "#   - État 'Running' pendant ~3-4 min (Mistral en Phase 2)",
        "#   - Puis 'Succeeded' avec tous les 5 nœuds verts",
    ], label="COMMANDE")
    expected_result(doc, "1 nouvelle exécution 'Succeeded' (~3-4 min, dont 2-3 min sur le nœud Call SRE Agent)")


# ════════════════════════════════════════════════════════════════════════════
# MISSION 3 — Agent IA + Remédiation
# ════════════════════════════════════════════════════════════════════════════

def add_mission_3(doc):
    mission_banner(doc, 3, "Agent IA + Auto-remédiation", "Pipeline 3 phases : collecte → LLM → GitOps. Validation visuelle via dashboard.")

    # Step 9
    step_title(doc, 9, "Déployer crash-app pour générer un cas réel")
    why_line(doc, "crash-app (busybox exit 1) déclenche immédiatement un CrashLoopBackOff.")
    code_block(doc, [
        "kubectl apply -f demo-app.yaml",
        "Start-Sleep -Seconds 60",
        "",
        "# Vérifier l'état CrashLoopBackOff",
        "kubectl get pods -n demo",
    ], label="COMMANDE")
    expected_result(doc, "crash-app-XXX  0/1  CrashLoopBackOff  3+  RESTARTS")

    # Step 10
    step_title(doc, 10, "Déclencher une investigation complète via /test-crash-loop")
    why_line(doc, "Endpoint dédié qui résout dynamiquement le pod crash-app et lance les 3 phases.")
    info_box(doc, "Mistral 7B sur CPU prend 150-250 secondes par investigation (Phase 2 = LLM = goulot d'étranglement). On met TimeoutSec=360 pour avoir de la marge. Si tu vois 'Le délai de l'opération a expiré', ce n'est PAS un échec : l'agent continue dans le pod, le résultat est dispo dans /api/incidents et la dashboard quelques secondes plus tard.", HEX_ORANGE_BG, icon="⚠")
    code_block(doc, [
        "Write-Host 'Investigation en cours (peut prendre 150-250 secondes)...' -ForegroundColor Yellow",
        "$start = Get-Date",
        "",
        "$resp = Invoke-WebRequest -Uri 'http://localhost:8001/test-crash-loop' `",
        "    -Method POST -TimeoutSec 360",
        "",
        "$dur = (Get-Date) - $start",
        "Write-Host \"Termine en $([math]::Round($dur.TotalSeconds,1))s\"",
        "",
        "$result = $resp.Content | ConvertFrom-Json",
        "$result | Select-Object classification, auto_remediation_applied, `",
        "    duration_seconds, git_pr_url, argocd_sync_triggered",
    ], label="OPTION A — ATTENDRE LA RÉPONSE")
    expected_result(doc, "classification=CrashLoopBackOff | auto_remediation_applied=True | duration_seconds ~150-250s")

    body(doc, "Alternative : lancer l'investigation en arrière-plan et suivre les logs directement (recommandé si tu veux observer chaque phase en temps réel).")
    code_block(doc, [
        "# Déclenchement non-bloquant (PowerShell job)",
        "Start-Job -ScriptBlock {",
        "    Invoke-WebRequest 'http://localhost:8001/test-crash-loop' -Method POST -TimeoutSec 600",
        "} | Out-Null",
        "Write-Host 'Investigation lancée en background.' -ForegroundColor Green",
        "",
        "# Suivre l'enchaînement Phase 1 → 2 → 3 en temps réel",
        "kubectl logs -f deployment/sre-agent -n sre-agent --tail=10",
        "# Ctrl+C quand tu vois : [Phase 3] Complete",
    ], label="OPTION B — FIRE-AND-FORGET + LOGS LIVE")
    expected_result(doc, "Les 3 phases défilent : Phase 1 (rapide) → Phase 2 (long, Mistral) → Phase 3 (rapide)")

    # Step 11
    step_title(doc, 11, "Vérifier les logs des 3 phases du pipeline")
    why_line(doc, "Les logs montrent l'enchaînement Phase 1 (kubectl, rapide) → Phase 2 (Mistral, ~3 min) → Phase 3 (GitOps, rapide).")
    code_block(doc, [
        "kubectl logs deployment/sre-agent -n sre-agent --tail=40",
        "",
        "# Lignes attendues, dans l'ordre :",
        "# [Phase 1] kubectl collection — demo/crash-app-XXX        ← T+0s",
        "# [Phase 1] Collected 4 sections                            ← T+2s",
        "# [Phase 2] LLM JSON-mode analysis                          ← T+2s",
        "# [Phase 2] Classification: CrashLoopBackOff | Fix: scale_to_zero  ← T+150-250s",
        "# [Phase 3] AUTO_REMEDIATE=true — applying fix_type=scale_to_zero",
        "# [Phase 3] Git result: SUCCESS: fix committed locally on branch 'fix/...'",
        "# [Phase 3] Complete: {'auto_remediation_applied': True, ...}    ← T+~200s",
    ], label="COMMANDE")
    expected_result(doc, "Les 3 phases défilent dans l'ordre, sans ERROR (la Phase 2 est l'étape lente)")

    # Step 12
    step_title(doc, 12, "Vérifier la remédiation sur crash-app")
    why_line(doc, "Le déploiement crash-app doit être scaled à 0 replicas.")
    code_block(doc, [
        "kubectl get deployment crash-app -n demo",
        "# Attendu : READY 0/0",
        "",
        "kubectl get pods -n demo -l app=crash-app",
        "# Attendu : 'No resources found' (le pod a été terminé)",
    ], label="COMMANDE")
    expected_result(doc, "crash-app  READY 0/0  UP-TO-DATE 0  AVAILABLE 0")

    # Step 13
    step_title(doc, 13, "Vérifier le commit Git automatique dans gitops-repo")
    why_line(doc, "L'agent commit l'incident dans /gitops-repo/incidents/ pour traçabilité.")
    code_block(doc, [
        "$pod = kubectl get pod -n sre-agent -l app=sre-agent --no-headers `",
        "    -o custom-columns='NAME:.metadata.name' | Select-Object -First 1",
        "",
        "# Historique des commits",
        "kubectl exec -n sre-agent $pod -- git -C /gitops-repo log --oneline",
        "",
        "# Lister les fichiers d'incidents",
        "kubectl exec -n sre-agent $pod -- ls /gitops-repo/incidents/demo/",
    ], label="COMMANDE")
    expected_result(doc, "Commits 'fix(demo): auto-remediate CrashLoopBackOff on crash-app' + fichier YAML d'incident")

    # Step 14
    step_title(doc, 14, "Ouvrir la dashboard et vérifier les stats")
    why_line(doc, "La dashboard centralise tous les payloads webhook avec stats agrégées.")
    code_block(doc, [
        "Start-Process 'http://localhost:8001/dashboard'",
        "",
        "# Vérifier les cartes en haut :",
        "# - Incidents       : >= 1",
        "# - Auto-remédiés  : >= 1",
        "# - Taux de succès : 100%",
        "# - MTTR moyen     : ~150-250s (Mistral 7B CPU)",
        "",
        "# Cliquer sur un incident → voir le payload JSON complet et le runbook",
    ], label="COMMANDE")
    expected_result(doc, "Dashboard affiche l'incident avec classification CrashLoopBackOff + runbook 3 étapes")


# ════════════════════════════════════════════════════════════════════════════
# MISSION 4 — Sécurité Kyverno
# ════════════════════════════════════════════════════════════════════════════

def add_mission_4(doc):
    mission_banner(doc, 4, "Sécurité Kyverno", "Auditer les workloads avec 4 politiques de sécurité.")

    # Step 15
    step_title(doc, 15, "Vérifier que les 4 politiques sont actives")
    why_line(doc, "Sans politique Ready=True, aucun audit ne se fait.")
    code_block(doc, [
        "kubectl get clusterpolicies",
        "",
        "# Attendu : 4 lignes avec ADMISSION=true, BACKGROUND=true, READY=True",
    ], label="COMMANDE")
    expected_result(doc, "4 ClusterPolicies en Ready : disallow-latest-tag, disallow-privileged-containers, require-labels, require-resource-limits")

    # Step 16
    step_title(doc, 16, "Déployer les 4 pods de test (violations volontaires)")
    why_line(doc, "Chaque pod viole une politique différente. Le mode Audit ne bloque pas.")
    code_block(doc, [
        "# Avant : remettre crash-app si elle a été scaled à 0",
        "kubectl scale deployment crash-app --replicas=1 -n demo",
        "",
        "# Appliquer les 4 violations",
        "kubectl apply -f kyverno/test-violations/",
        "",
        "# Lister les pods",
        "kubectl get pods -n demo",
    ], label="COMMANDE")
    expected_result(doc, "4 pods 'violation-*' créés (mode Audit = pas de blocage)")

    # Step 17
    step_title(doc, 17, "Lire les PolicyReports générés")
    why_line(doc, "Le reports-controller crée un PolicyReport par ressource auditée.")
    code_block(doc, [
        "Start-Sleep -Seconds 5",
        "",
        "kubectl get policyreport -n demo",
        "",
        "# Attendu : 6 PolicyReports",
        "# - 4 pour les pods 'violation-*' (Fail: 1, Pass: 4)",
        "# - 1 pour crash-app Deployment (Fail: 3, Pass: 2)",
        "# - 1 pour crash-app ReplicaSet (Fail: 3, Pass: 2)",
    ], label="COMMANDE")
    expected_result(doc, "6 PolicyReports avec des violations détectées")

    # Step 18
    step_title(doc, 18, "Compter les violations par politique")
    why_line(doc, "Vue agrégée du niveau de conformité du cluster.")
    code_block(doc, [
        "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty items | `",
        "    Select-Object -ExpandProperty results -ErrorAction SilentlyContinue | `",
        "    Where-Object { $_.result -eq 'fail' } | `",
        "    Group-Object -Property policy | `",
        "    Select-Object Name, Count | Format-Table -AutoSize",
        "",
        "# Attendu :",
        "# disallow-latest-tag                  3",
        "# disallow-privileged-containers       1",
        "# require-labels                       3",
        "# require-resource-limits              3",
    ], label="COMMANDE")
    expected_result(doc, "Total ~10 violations réparties sur les 4 politiques")

    # Step 19
    step_title(doc, 19, "Détail des violations sur crash-app (cas pédagogique)")
    why_line(doc, "Sans modification, crash-app viole déjà 3 best practices.")
    code_block(doc, [
        "kubectl get policyreport -n demo -o json | ConvertFrom-Json | `",
        "    Select-Object -ExpandProperty items | `",
        "    Where-Object { $_.scope.name -eq 'crash-app' -and $_.scope.kind -eq 'Deployment' } | `",
        "    Select-Object -ExpandProperty results | `",
        "    Where-Object { $_.result -eq 'fail' } | `",
        "    Select-Object policy, message | Format-List",
    ], label="COMMANDE")
    expected_result(doc, "3 violations sur crash-app : tag manquant, label 'owner' absent, limits.cpu manquant")


# ════════════════════════════════════════════════════════════════════════════
# MISSION 5 — Métriques + Grafana
# ════════════════════════════════════════════════════════════════════════════

def add_mission_5(doc):
    mission_banner(doc, 5, "Métriques + Dashboard Grafana", "Vérifier l'observabilité bout-en-bout : /metrics, ServiceMonitor, Grafana auto-import.")

    # Step 20
    step_title(doc, 20, "Tester l'endpoint /metrics du SRE Agent")
    why_line(doc, "Vérifie que prometheus_client expose correctement les compteurs.")
    code_block(doc, [
        "$r = Invoke-WebRequest 'http://localhost:8001/metrics' -TimeoutSec 5",
        "Write-Host \"Status: $($r.StatusCode) | Content-Type: $($r.Headers.'Content-Type')\"",
        "",
        "# Voir les premières métriques sre_*",
        "$r.Content -split \"`n\" | Select-String '^sre_' | Select-Object -First 10",
    ], label="COMMANDE")
    expected_result(doc, "HTTP 200 + Content-Type: text/plain; version=0.0.4 + lignes sre_alerts_total, sre_investigation_duration_seconds_bucket, etc.")

    # Step 21
    step_title(doc, 21, "Vérifier que Prometheus scrape l'agent (ServiceMonitor)")
    why_line(doc, "Le ServiceMonitor 'sre-agent' permet à Prometheus de découvrir le target automatiquement.")
    code_block(doc, [
        "$t = (Invoke-WebRequest 'http://localhost:9090/api/v1/targets' -TimeoutSec 5).Content | `",
        "    ConvertFrom-Json",
        "$t.data.activeTargets | Where-Object { $_.labels.job -eq 'sre-agent-service' } | `",
        "    Select-Object scrapeUrl, health, lastScrape | Format-List",
    ], label="COMMANDE")
    expected_result(doc, "health=up | scrapeUrl=http://<pod-ip>:8000/metrics | lastScrape récent")

    # Step 22
    step_title(doc, 22, "Ouvrir le dashboard Grafana 'SRE Copilot — KPIs'")
    why_line(doc, "Le dashboard est auto-importé par le sidecar Grafana via la ConfigMap labellée grafana_dashboard=1.")
    code_block(doc, [
        "Start-Process 'http://localhost:3000/d/sre-copilot/sre-copilot-kpis'",
        "",
        "# Login : admin / prom-operator",
        "# Le dashboard a 9 panels :",
        "#   1. Alertes totales (stat)",
        "#   2. Taux de remédiation auto (gauge %)",
        "#   3. MTTR médian (stat s)",
        "#   4. ArgoCD sync OK (stat)",
        "#   5. Fallbacks LLM (stat)",
        "#   6. Taux d'alertes par classification (time series)",
        "#   7. Répartition classifications (donut)",
        "#   8. Durée investigation p50/p95 (time series)",
        "#   9. Remédiations par type+résultat (bar chart)",
    ], label="COMMANDE")
    expected_result(doc, "Dashboard visible avec 9 panels, valeurs remplies après quelques investigations")


# ════════════════════════════════════════════════════════════════════════════
# Validation finale
# ════════════════════════════════════════════════════════════════════════════

def add_validation(doc):
    heading1(doc, "Validation Finale du Scénario")

    # Step 23
    step_title(doc, 23, "Récapitulatif des données dans le dashboard SRE Copilot")
    why_line(doc, "Confirmer la persistance des incidents et les stats agrégées.")
    code_block(doc, [
        "Invoke-WebRequest 'http://localhost:8001/api/stats' | `",
        "    Select-Object -ExpandProperty Content | ConvertFrom-Json | `",
        "    ConvertTo-Json -Depth 4",
    ], label="COMMANDE")
    expected_result(doc, "Total >= 1, auto_remediated >= 1, success_rate ~= 100%")

    # Step 24
    step_title(doc, 24, "Cleanup (optionnel)")
    why_line(doc, "Préparer l'environnement pour un nouveau scénario.")
    code_block(doc, [
        "# Supprimer les pods de test Kyverno",
        "kubectl delete -f kyverno/test-violations/",
        "",
        "# Remettre crash-app à 1 replica pour relancer un test",
        "kubectl scale deployment crash-app --replicas=1 -n demo",
        "",
        "# (Optionnel) Arrêter les port-forwards sans supprimer le cluster",
        ".\\scripts\\stop-sre-copilot.ps1",
    ], label="COMMANDE")
    expected_result(doc, "Pods de test supprimés, crash-app à 1 replica, port-forwards fermés")


# ════════════════════════════════════════════════════════════════════════════
# Summary
# ════════════════════════════════════════════════════════════════════════════

def add_summary(doc):
    doc.add_page_break()
    heading1(doc, "Récapitulatif du Scénario")

    make_table(doc,
        ["Mission", "Composant testé", "Résultat attendu", "Statut"],
        [
            ("1", "Cluster KinD + monitoring stack",
             "3 nœuds Ready, 6+ pods monitoring Running",
             "[ ]"),
            ("1", "Ollama + Mistral disponible",
             "Models: ['mistral:latest']",
             "[ ]"),
            ("2", "Alertmanager → n8n routing",
             "POST /api/v2/alerts → exec n8n Succeeded",
             "[ ]"),
            ("2", "Workflow n8n actif (5 nœuds)",
             "responseMode: onReceived, maxConcurrent: 1",
             "[ ]"),
            ("3", "Pipeline 3 phases SRE Agent",
             "Phase 1 + 2 + 3 dans les logs",
             "[ ]"),
            ("3", "Remédiation automatique",
             "crash-app scaled à 0 + git commit",
             "[ ]"),
            ("3", "Dashboard centralisée",
             "/dashboard affiche l'incident + payload JSON",
             "[ ]"),
            ("4", "4 ClusterPolicies Kyverno actives",
             "READY=True sur les 4",
             "[ ]"),
            ("4", "Détection des 4 types de violation",
             "4 PolicyReports avec FAIL",
             "[ ]"),
            ("4", "Audit automatique crash-app",
             "3 violations détectées sans intervention",
             "[ ]"),
        ],
        col_widths=[0.6, 2.4, 3.2, 0.5]
    )

    doc.add_paragraph()
    heading1(doc, "Si tout est vert : le projet PFE est techniquement opérationnel")
    body(doc, "Ce scénario valide bout en bout :")
    body(doc, "• La stack d'observabilité (Prometheus + Alertmanager + Grafana)")
    body(doc, "• L'orchestration de workflows (n8n)")
    body(doc, "• L'agent IA local avec Mistral et le pipeline 3 phases")
    body(doc, "• La remédiation GitOps automatique (scale + git commit + ArgoCD sync)")
    body(doc, "• La dashboard centralisée pour les payloads webhook")
    body(doc, "• Le scan de sécurité Kyverno (Policy-as-Code)")
    doc.add_paragraph()
    info_box(doc, "Toutes les 5 missions sont livrées et testées. Le projet technique est complet — prochaine étape : la rédaction du rapport.", HEX_GREEN_BG, icon="✓")


# ── Main ───────────────────────────────────────────────────────────────────────

def main():
    out_path = os.path.join(os.path.dirname(__file__), "..", "SRE-Copilot-Scenario-E2E.docx")
    out_path = os.path.abspath(out_path)

    doc = Document()

    section = doc.sections[0]
    section.page_width    = Cm(21)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    doc.styles['Normal'].font.name = "Calibri"
    doc.styles['Normal'].font.size = Pt(10)

    add_cover(doc)
    add_overview(doc)
    add_mission_1(doc)
    add_mission_2(doc)
    add_mission_3(doc)
    add_mission_4(doc)
    add_mission_5(doc)
    add_validation(doc)
    add_summary(doc)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("PFE IFS 013/26 — PwC DigiTech Cloud Native Operations | SRE Copilot | Scénario E2E M1→M4 | 2026-06-03")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = C_GRAY_LT
    run.italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    main()
