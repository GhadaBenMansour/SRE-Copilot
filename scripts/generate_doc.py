from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Styles helpers
def set_heading(paragraph, text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
    paragraph.clear()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)

def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    return p

def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x23, 0x85, 0x56)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        p.add_run(text).font.size = Pt(11)
    else:
        p.add_run(text).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # Gray background shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    return p

def add_note(doc, text, color_fill='FFF3CD', color_text=RGBColor(0x85, 0x64, 0x04)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_fill)
    pPr.append(shd)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = color_text
    return p

def add_separator(doc):
    p = doc.add_paragraph()
    run = p.add_run('─' * 90)
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ══════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
add_title(doc, "SRE Copilot")
add_subtitle(doc, "Guide de test et d'utilisation — Workflow complet")
doc.add_paragraph()
add_subtitle(doc, "Stage PFE · PwC DigiTech / Cloud Native Operations")
add_subtitle(doc, "Remédiation GitOps pilotée par un agent IA orchestré avec n8n")
doc.add_paragraph()
add_subtitle(doc, "Auteur : Stagiaire PFE  |  Date : Avril 2026")
doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES (manuelle)
# ══════════════════════════════════════════════════════════════
add_h1(doc, "Table des matières")
toc_items = [
    ("1.", "Architecture et vue d'ensemble du projet"),
    ("2.", "Description du workflow n8n"),
    ("3.", "Description du SRE Agent (FastAPI + LangChain + Mistral)"),
    ("4.", "L'application de test : crash-app"),
    ("5.", "Étapes de lancement et de test"),
    ("   5.1", "Étape 1  — Démarrer l'environnement"),
    ("   5.2", "Étape 2  — Déployer le SRE Agent"),
    ("   5.3", "Étape 3  — Vérifier le SRE Agent"),
    ("   5.4", "Étape 4  — Importer le workflow dans n8n"),
    ("   5.5", "Étape 5  — Appliquer la config Alertmanager"),
    ("   5.6", "Étape 6  — Vérifier l'application de test"),
    ("   5.7", "Étape 7  — Vérifier l'alerte dans Prometheus"),
    ("   5.8", "Étape 8  — Tester le flux complet"),
    ("   5.9", "Étape 9  — Tester l'agent directement"),
    ("   5.10", "Étape 10 — Lire et interpréter le résultat"),
    ("6.", "Récapitulatif des URLs et credentials"),
    ("7.", "Commandes de diagnostic utiles"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    r1 = p.add_run(num + "  ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 1. ARCHITECTURE
# ══════════════════════════════════════════════════════════════
add_h1(doc, "1. Architecture et vue d'ensemble du projet")
add_separator(doc)

add_body(doc,
    "Le projet SRE Copilot est un agent IA de remédiation GitOps qui s'exécute entièrement en local "
    "sur un PC personnel, sans dépendre d'une infrastructure PwC. Il repose sur un cluster Kubernetes "
    "local (KinD), un LLM local (Ollama + Mistral 7B) et un orchestrateur de workflows (n8n)."
)

add_h2(doc, "1.1  Stack technique locale")

stack = [
    ("KinD (sre-cluster)",        "1 control-plane + 2 workers — cluster Kubernetes local via Docker Desktop"),
    ("Prometheus + Alertmanager", "kube-prometheus-stack (Helm) — namespace monitoring — détection et routage des alertes"),
    ("Grafana",                   "Dashboards Kubernetes — namespace monitoring — port 3000"),
    ("ArgoCD",                    "Déploiement GitOps — namespace argocd — port 8080"),
    ("n8n",                       "Orchestrateur de workflow — namespace n8n — port 5678 (NodePort 30567)"),
    ("Ollama + Mistral 7B",       "LLM local sur Windows — accessible via host.docker.internal:11434"),
    ("SRE Agent (FastAPI)",       "API Python LangChain — namespace sre-agent — port 8000"),
]

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = "Composant"
hdr[1].text = "Description"
for cell in hdr:
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1F497D')
    cell.paragraphs[0]._p.get_or_add_pPr().append(shd)

for comp, desc in stack:
    row = table.add_row().cells
    row[0].text = comp
    row[1].text = desc
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

add_h2(doc, "1.2  Flux de traitement bout en bout")
add_body(doc, "Le flux complet d'une alerte, de sa détection à la remédiation :")

flow_steps = [
    ("crash-app (namespace demo)",   "Pod busybox avec commande 'exit 1' → CrashLoopBackOff permanent"),
    ("Prometheus",                   "Détecte via la règle PodRestartingTooMuch (> 5 redémarrages en 10 min)"),
    ("Alertmanager",                 "Route les alertes critical/warning → webhook POST vers n8n"),
    ("n8n Workflow",                 "Réceptionne le webhook → vérifie status=firing → appelle le SRE Agent"),
    ("SRE Agent /alert",             "LangChain ReAct : investigate kubectl → envoie contexte à Mistral"),
    ("Ollama / Mistral 7B",          "Génère : classification + root cause + runbook structuré"),
    ("n8n (suite)",                  "Parse la réponse → construit rapport incident → Audit Log → répond au webhook"),
    ("ArgoCD (optionnel)",           "Si PR GitOps créée → sync automatique du déploiement corrigé"),
]

for i, (actor, action) in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"  {i}.  {actor}  →  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    r2 = p.add_run(action)
    r2.font.size = Pt(10)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 2. WORKFLOW N8N
# ══════════════════════════════════════════════════════════════
add_h1(doc, "2. Description du workflow n8n")
add_separator(doc)

add_body(doc,
    "Le fichier n8n-workflows/sre-alert-workflow.json contient le workflow "
    "'SRE Copilot - Alert Remediation'. Il comporte 8 noeuds connectés qui orchestrent "
    "le traitement complet d'une alerte Prometheus."
)

nodes = [
    ("1", "Alertmanager Webhook",  "Noeud déclencheur",
     "Écoute les POST sur /webhook/sre-alert. C'est le point d'entrée du workflow. "
     "Alertmanager envoie ici les alertes via l'URL interne : "
     "http://n8n-service.n8n.svc.cluster.local:5678/webhook/sre-alert"),

    ("2", "Is Firing?",            "Noeud de condition (IF)",
     "Vérifie que le champ body.status == 'firing'. Si l'alerte est résolue (resolved) "
     "ou en attente (pending), elle est ignorée. Cela évite de traiter des alertes non actives."),

    ("3", "Call SRE Agent",        "Noeud HTTP Request",
     "Envoie un POST vers http://sre-agent-service.sre-agent.svc.cluster.local:8000/alert "
     "avec le payload complet de l'alerte. Timeout : 120 secondes (Mistral peut être lent)."),

    ("4", "Parse Agent Response",  "Noeud Code (JavaScript)",
     "Extrait et structure la réponse JSON du SRE Agent : alert, namespace, severity, "
     "classification, rootCause, runbook, prUrl, autoRemediated, argoCDSync."),

    ("5", "Has PR?",               "Noeud de condition (IF)",
     "Vérifie si une Pull Request GitOps a été créée (prUrl != 'none'). "
     "Si oui → passe par Audit Log. Sinon → va directement à Webhook Response."),

    ("6", "Build Incident Report", "Noeud Code (JavaScript)",
     "Construit un rapport d'incident lisible (format Slack/Teams) avec emojis "
     "de sévérité, résumé de la cause, statut de remédiation et runbook formaté."),

    ("7", "Audit Log",             "Noeud Code (JavaScript)",
     "Enregistre une entrée d'audit dans les logs n8n : timestamp, alerte, namespace, "
     "sévérité, classification, statut de remédiation, URL PR. Traçabilité complète."),

    ("8", "Webhook Response",      "Noeud Respond to Webhook",
     "Renvoie une réponse JSON à Alertmanager avec le statut 'processed' et le "
     "détail de tous les résultats traités. Alertmanager confirme ainsi la livraison."),
]

for num, name, ntype, desc in nodes:
    add_h3(doc, f"Noeud {num} — {name}  ({ntype})")
    add_body(doc, desc)

add_note(doc,
    "IMPORTANT : Le workflow doit être activé dans n8n (toggle Inactive → Active) "
    "pour que le webhook soit opérationnel. Sans activation, les alertes ne sont pas reçues."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 3. SRE AGENT
# ══════════════════════════════════════════════════════════════
add_h1(doc, "3. Description du SRE Agent (FastAPI + LangChain + Mistral)")
add_separator(doc)

add_body(doc,
    "Le SRE Agent est une API Python (FastAPI) qui expose plusieurs endpoints REST. "
    "Il utilise LangChain en mode ReAct (Reasoning + Acting) avec Ollama/Mistral 7B "
    "comme modèle de langage local."
)

add_h2(doc, "3.1  Endpoints de l'API")

endpoints = [
    ("GET  /health",           "Vérification que le service est démarré. Retourne { status: ok }."),
    ("GET  /status",           "Vérifie la connexion à Ollama et la disponibilité du modèle Mistral."),
    ("POST /alert",            "Endpoint principal : reçoit le payload Alertmanager complet (liste d'alertes)."),
    ("POST /analyze",          "Analyse une seule alerte. Utile pour les tests directs sans Alertmanager."),
    ("POST /test-crash-loop",  "Génère en interne une alerte CrashLoopBackOff fictive et la traite immédiatement."),
]

for ep, desc in endpoints:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(ep.ljust(28))
    r1.bold = True
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x23, 0x85, 0x56)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

add_h2(doc, "3.2  Fonctionnement de l'agent ReAct (agent.py)")

add_body(doc,
    "L'agent suit le pattern ReAct (Reasoning + Acting) défini par LangChain. "
    "À chaque étape, il alterne entre :"
)
add_bullet(doc, "Thought  — raisonnement interne sur l'alerte reçue", "")
add_bullet(doc, "Action   — appel d'un outil kubectl (logs, events, describe...)", "")
add_bullet(doc, "Observation — résultat de l'outil renvoyé à Mistral", "")
add_bullet(doc, "Final Answer — réponse structurée JSON avec classification + runbook", "")

add_body(doc, "Paramètres de l'agent :")
add_bullet(doc, "Modèle : Mistral 7B (Q4_K_M, 4.3 GB) via Ollama sur Windows", "")
add_bullet(doc, "Température : 0.1 (réponses déterministes et précises)", "")
add_bullet(doc, "Max iterations : 15 (évite les boucles infinies)", "")
add_bullet(doc, "Timeout Ollama : 120 secondes par requête", "")

add_h2(doc, "3.3  Outils kubectl disponibles (kubectl_tools.py)")
add_body(doc, "L'agent dispose des outils suivants pour investiguer le cluster :")

kubectl_tools = [
    "get_pod_logs         — récupère les logs d'un pod (dernières 100 lignes)",
    "describe_pod         — kubectl describe pod (events, état des containers)",
    "get_pod_events       — events Kubernetes du namespace",
    "get_deployment_info  — état du déploiement (replicas, stratégie)",
    "get_pod_status       — liste des pods avec statut et restarts",
]
for t in kubectl_tools:
    add_bullet(doc, t)

add_h2(doc, "3.4  Format de réponse JSON")
add_body(doc, "L'agent retourne toujours un JSON structuré :")
add_code(doc,
    '{\n'
    '  "classification":          "CrashLoopBackOff | OOMKilled | ImagePullBackOff | ConfigError | ...",\n'
    '  "root_cause":              "Explication concise de la cause racine",\n'
    '  "runbook": [\n'
    '    { "order": 1, "action": "Investigate", "command": "kubectl logs ...", "description": "..." },\n'
    '    { "order": 2, "action": "Fix",         "command": "kubectl set image ...", "description": "..." }\n'
    '  ],\n'
    '  "auto_remediation_applied": false,\n'
    '  "remediation_details":     "Raison de la décision de remédiation",\n'
    '  "git_pr_url":              null,\n'
    '  "argocd_sync_triggered":   false\n'
    '}'
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 4. APPLICATION DE TEST
# ══════════════════════════════════════════════════════════════
add_h1(doc, "4. L'application de test : crash-app")
add_separator(doc)

add_body(doc,
    "L'application de test est un pod Kubernetes volontairement défectueux, "
    "situé dans le namespace demo. Elle est définie dans :"
)
add_bullet(doc, "demo-app.yaml  (à la racine du projet)")
add_bullet(doc, "gitops-repo/apps/demo/crash-app.yaml  (version gérée par ArgoCD)")

add_h2(doc, "4.1  Ce qu'elle fait")
add_body(doc,
    "C'est un conteneur busybox:latest avec une seule commande : exit 1. "
    "Dès qu'il démarre, il se termine immédiatement avec le code d'erreur 1. "
    "Kubernetes tente de le redémarrer (politique RestartPolicy: Always par défaut), "
    "ce qui provoque une boucle infinie de crashs : le CrashLoopBackOff."
)
add_code(doc, 'command: ["sh", "-c", "exit 1"]   # crash immédiat intentionnel')

add_h2(doc, "4.2  Ce que cela déclenche")
steps_crash = [
    "Kubernetes marque le pod CrashLoopBackOff après quelques redémarrages",
    "Le compteur de restarts monte continuellement (actuellement > 225 restarts)",
    "Prometheus évalue la règle PodRestartingTooMuch toutes les 30 secondes",
    "Après 5 redémarrages en 10 minutes → alerte passe en état Firing",
    "Alertmanager route l'alerte vers n8n via webhook",
    "Le workflow SRE Copilot s'exécute automatiquement",
]
for s in steps_crash:
    add_bullet(doc, s)

add_h2(doc, "4.3  Ce que l'agent diagnostique")
add_body(doc,
    "L'agent Mistral identifie que la commande du container est 'exit 1', "
    "ce qui explique le crash immédiat. Le runbook généré propose typiquement :"
)
add_bullet(doc, "Investiguer les logs : kubectl logs crash-app -n demo")
add_bullet(doc, "Vérifier les events : kubectl describe pod crash-app -n demo")
add_bullet(doc, "Corriger la commande : kubectl set image ou modifier le manifest YAML")
add_bullet(doc, "Créer une PR GitOps avec le manifest corrigé (si AUTO_REMEDIATE=false)")

add_note(doc,
    "INFO : Le pod crash-app est en cours d'exécution depuis 34 jours dans le cluster. "
    "Il n'est pas nécessaire de le redéployer pour les tests — il génère déjà des alertes."
)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 5. ÉTAPES DE TEST
# ══════════════════════════════════════════════════════════════
add_h1(doc, "5. Étapes de lancement et de test")
add_separator(doc)

add_note(doc,
    "PRÉ-REQUIS : Docker Desktop doit être démarré. "
    "Ollama doit tourner sur Windows (ollama serve). "
    "Le cluster KinD sre-cluster doit être actif (kind get clusters)."
)

# ── 5.1
add_h2(doc, "Étape 1 — Démarrer l'environnement")
add_body(doc, "Lance le script de démarrage depuis PowerShell :")
add_code(doc, "cd C:\\Users\\PC\\Desktop\\SRE-Copilot\n.\\scripts\\start-sre-copilot.ps1")
add_body(doc, "Ce script :")
add_bullet(doc, "Tue les anciens processus kubectl port-forward (évite les conflits de ports)")
add_bullet(doc, "Vérifie Docker, le cluster KinD, les nodes et les pods")
add_bullet(doc, "Lance 5 fenêtres PowerShell avec port-forward persistant (auto-relance si coupure)")
add_bullet(doc, "Ouvre automatiquement Grafana, Prometheus, n8n et ArgoCD dans le navigateur")
add_body(doc, "Résultat attendu : 5 fenêtres PowerShell affichant 'Forwarding from 127.0.0.1:XXXX'")

# ── 5.2
add_h2(doc, "Étape 2 — Déployer le SRE Agent")
add_body(doc, "Lance le script de déploiement :")
add_code(doc, ".\\scripts\\deploy-sre-agent.ps1")
add_body(doc, "Ce script effectue automatiquement 6 opérations :")
ops = [
    ("docker build",       "Construit l'image Docker de l'agent FastAPI depuis sre-agent/Dockerfile"),
    ("kind load",          "Charge l'image dans le cluster KinD (sans passer par Docker Hub)"),
    ("kubectl apply",      "Crée le namespace, le ServiceAccount, les RBAC, le Deployment et le Service"),
    ("kubectl create secret", "Crée le secret kubeconfig pour que l'agent puisse exécuter des kubectl"),
    ("kubectl rollout status", "Attend que le pod soit Ready (timeout 120s)"),
    ("port-forward",       "Ouvre le port 8000 en local pour accéder à l'API"),
]
for action, desc in ops:
    add_bullet(doc, f"{desc}", f"{action}  —  ")

add_body(doc, "Vérification :")
add_code(doc, "kubectl get pods -n sre-agent\n# Attendu : sre-agent-xxx   1/1   Running   0")

# ── 5.3
add_h2(doc, "Étape 3 — Vérifier le SRE Agent")
add_body(doc, "Teste que l'API répond et que Mistral est accessible :")
add_code(doc,
    "# Santé de l'API\n"
    "curl http://localhost:8000/health\n\n"
    "# Connexion Ollama + modèle disponible\n"
    "curl http://localhost:8000/status"
)
add_body(doc, "Réponse attendue sur /status :")
add_code(doc, '{ "ollama": "reachable", "model": "mistral", "model_available": true }')
add_note(doc,
    "Si ollama est 'unreachable' : vérifier que Ollama tourne sur Windows "
    "(commande : ollama serve dans un terminal séparé)."
)

# ── 5.4
add_h2(doc, "Étape 4 — Importer le workflow dans n8n")
add_body(doc, "1. Ouvre http://localhost:5678 dans le navigateur")
add_body(doc, "2. Menu gauche → Workflows → bouton + → Import from file")
add_body(doc, "3. Sélectionne le fichier : n8n-workflows/sre-alert-workflow.json")
add_body(doc, "4. Une fois importé, ouvre le workflow")
add_body(doc, "5. Clique sur le toggle Inactive → Active (en haut à droite)")
add_note(doc,
    "L'URL du webhook sera : http://n8n-service.n8n.svc.cluster.local:5678/webhook/sre-alert\n"
    "(URL interne cluster, utilisée par Alertmanager)\n"
    "Pour les tests manuels depuis Windows : http://localhost:5678/webhook/sre-alert"
)

# ── 5.5
add_h2(doc, "Étape 5 — Appliquer la configuration Alertmanager")
add_body(doc, "Cette config route les alertes critical/warning vers n8n :")
add_code(doc, "kubectl apply -f C:\\Users\\PC\\Desktop\\SRE-Copilot\\alertmanager-config.yaml")
add_body(doc, "Vérifie que la config est bien prise en compte sur http://localhost:9093 (onglet Status → Config).")

# ── 5.6
add_h2(doc, "Étape 6 — Vérifier l'application de test crash-app")
add_code(doc, "kubectl get pods -n demo")
add_body(doc, "Résultat attendu :")
add_code(doc, "NAME                         READY   STATUS             RESTARTS\ncrash-app-xxx                0/1     CrashLoopBackOff   225+")
add_body(doc, "Si le namespace demo n'existe pas :")
add_code(doc, "kubectl apply -f C:\\Users\\PC\\Desktop\\SRE-Copilot\\demo-app.yaml")

# ── 5.7
add_h2(doc, "Étape 7 — Vérifier l'alerte dans Prometheus")
add_body(doc, "Ouvre http://localhost:9090/alerts")
add_body(doc, "L'alerte PodRestartingTooMuch doit apparaître en rouge (état Firing).")
add_body(doc, "Si elle n'apparaît pas, applique la règle Prometheus :")
add_code(doc,
    "kubectl apply -f C:\\Users\\PC\\Desktop\\SRE-Copilot\\pod-restart-alert.yaml\n"
    "# Attendre environ 2 minutes pour que l'alerte passe en Firing"
)

# ── 5.8
add_h2(doc, "Étape 8 — Tester le flux complet (webhook → n8n → Agent → Mistral)")
add_body(doc,
    "Cette commande simule exactement ce qu'Alertmanager envoie à n8n. "
    "Lance-la depuis PowerShell :"
)
add_code(doc,
    "Invoke-RestMethod -Uri 'http://localhost:5678/webhook/sre-alert' `\n"
    "  -Method POST `\n"
    "  -ContentType 'application/json' `\n"
    "  -Body '{\n"
    "    \"receiver\": \"n8n-webhook\",\n"
    "    \"status\": \"firing\",\n"
    "    \"alerts\": [{\n"
    "      \"status\": \"firing\",\n"
    "      \"labels\": {\n"
    "        \"alertname\": \"PodRestartingTooMuch\",\n"
    "        \"namespace\": \"demo\",\n"
    "        \"pod\": \"crash-app\",\n"
    "        \"severity\": \"critical\"\n"
    "      },\n"
    "      \"annotations\": {\n"
    "        \"summary\": \"Pod crash-app redémarre trop souvent\"\n"
    "      }\n"
    "    }]\n"
    "  }'"
)
add_body(doc,
    "Pendant l'exécution (30 à 60 secondes), surveille dans n8n : "
    "Workflows → Executions. Tu verras les noeuds s'allumer un par un."
)

# ── 5.9
add_h2(doc, "Étape 9 — Tester l'agent directement (sans n8n)")
add_body(doc,
    "L'endpoint /test-crash-loop génère une alerte fictive en interne "
    "et la traite sans passer par n8n. Utile pour déboguer uniquement l'agent :"
)
add_code(doc,
    "Invoke-RestMethod -Uri 'http://localhost:8000/test-crash-loop' -Method POST"
)
add_body(doc, "Ou depuis le navigateur via la doc Swagger interactive :")
add_code(doc, "http://localhost:8000/docs")

# ── 5.10
add_h2(doc, "Étape 10 — Lire et interpréter le résultat")
add_body(doc, "La réponse JSON typique ressemble à :")
add_code(doc,
    '{\n'
    '  "alert_name":               "PodRestartingTooMuch",\n'
    '  "namespace":                "demo",\n'
    '  "severity":                 "critical",\n'
    '  "classification":           "CrashLoopBackOff",\n'
    '  "root_cause":               "Container exits immediately with code 1 (command: exit 1)",\n'
    '  "runbook": [\n'
    '    { "order":1, "action":"Investigate", "command":"kubectl logs crash-app -n demo" },\n'
    '    { "order":2, "action":"Investigate", "command":"kubectl describe pod crash-app -n demo" },\n'
    '    { "order":3, "action":"Fix",         "command":"kubectl set image deployment/crash-app crash-container=busybox -- sleep 3600 -n demo" }\n'
    '  ],\n'
    '  "auto_remediation_applied": false,\n'
    '  "remediation_details":      "AUTO_REMEDIATE is disabled. PR would be required.",\n'
    '  "git_pr_url":               null,\n'
    '  "argocd_sync_triggered":    false,\n'
    '  "investigation_log":        ["[get_pod_logs] ...", "[describe_pod] ..."]\n'
    '}'
)

add_body(doc, "Interprétation des champs :")
fields = [
    ("classification",           "Type de panne identifié par Mistral"),
    ("root_cause",               "Explication en langage naturel de la cause racine"),
    ("runbook",                  "Liste d'actions correctives ordonnées avec commandes kubectl"),
    ("auto_remediation_applied", "true si l'agent a appliqué le fix directement"),
    ("git_pr_url",               "URL de la PR GitOps créée (si remédiation via ArgoCD)"),
    ("argocd_sync_triggered",    "true si ArgoCD a été déclenché pour synchroniser"),
    ("investigation_log",        "Trace de toutes les actions kubectl exécutées par l'agent"),
]
for field, desc in fields:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"{field:<35}")
    r1.font.name = 'Courier New'
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x23, 0x85, 0x56)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 6. RÉCAPITULATIF URLS
# ══════════════════════════════════════════════════════════════
add_h1(doc, "6. Récapitulatif des URLs et credentials")
add_separator(doc)

urls = [
    ("Grafana",      "http://localhost:3000",      "admin",  "prom-operator"),
    ("Prometheus",   "http://localhost:9090",      "—",      "—"),
    ("Alertmanager", "http://localhost:9093",      "—",      "—"),
    ("n8n",          "http://localhost:5678",      "—",      "—"),
    ("ArgoCD",       "http://localhost:8080",      "admin",  "admin123"),
    ("SRE Agent API","http://localhost:8000",      "—",      "—"),
    ("SRE Agent Docs","http://localhost:8000/docs","—",      "—"),
    ("Ollama",       "http://localhost:11434",     "—",      "—"),
]

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Table Grid'
hdrs = ["Service", "URL", "Utilisateur", "Mot de passe"]
for i, h in enumerate(hdrs):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '2E74B5')
    cell.paragraphs[0]._p.get_or_add_pPr().append(shd)

for svc, url, user, pwd in urls:
    row = table2.add_row().cells
    row[0].text = svc
    row[1].text = url
    row[2].text = user
    row[3].text = pwd
    for c in row:
        for p in c.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════
# 7. COMMANDES DE DIAGNOSTIC
# ══════════════════════════════════════════════════════════════
add_h1(doc, "7. Commandes de diagnostic utiles")
add_separator(doc)

add_h2(doc, "7.1  État général du cluster")
add_code(doc,
    "kubectl get nodes\n"
    "kubectl get pods -A --field-selector=status.phase!=Running\n"
    "kubectl get pods -n monitoring\n"
    "kubectl get pods -n n8n\n"
    "kubectl get pods -n argocd\n"
    "kubectl get pods -n sre-agent\n"
    "kubectl get pods -n demo"
)

add_h2(doc, "7.2  Logs des composants")
add_code(doc,
    "# Logs SRE Agent\n"
    "kubectl logs -n sre-agent -l app=sre-agent -f\n\n"
    "# Logs n8n\n"
    "kubectl logs -n n8n -l app=n8n -f\n\n"
    "# Logs Alertmanager\n"
    "kubectl logs -n monitoring -l app.kubernetes.io/name=alertmanager -f\n\n"
    "# Logs crash-app (pour voir les erreurs)\n"
    "kubectl logs -n demo -l app=crash-app"
)

add_h2(doc, "7.3  Vérifier Ollama et Mistral")
add_code(doc,
    "# Modèles disponibles\n"
    "curl http://localhost:11434/api/tags\n\n"
    "# Tester Mistral directement\n"
    "curl http://localhost:11434/api/generate -d '{\"model\":\"mistral\",\"prompt\":\"Hello\",\"stream\":false}'"
)

add_h2(doc, "7.4  Arrêter et redémarrer les port-forwards")
add_code(doc,
    "# Arrêter tous les port-forwards\n"
    "Get-Process kubectl | Stop-Process -Force\n\n"
    "# Relancer l'environnement complet\n"
    ".\\scripts\\start-sre-copilot.ps1"
)

add_h2(doc, "7.5  Vérifier que l'alerte est bien routée")
add_code(doc,
    "# Voir la config active d'Alertmanager\n"
    "kubectl get secret -n monitoring alertmanager-monitoring-kube-prometheus-alertmanager ^\n"
    "  -o jsonpath='{.data.alertmanager\\.yaml}' | ^\n"
    "  powershell -c \"[System.Text.Encoding]::UTF8.GetString([Convert]::FromBase64String((cat -)))\""
)

add_note(doc,
    "En cas de problème avec n8n : vérifier que le workflow est bien ACTIVE "
    "(pas juste sauvegardé). Sans activation, le webhook n'est pas enregistré et "
    "les alertes d'Alertmanager retourneront une erreur 404.",
    color_fill='FFE0E0',
    color_text=RGBColor(0x8B, 0x00, 0x00)
)

doc.add_paragraph()
add_separator(doc)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("SRE Copilot — Guide de test  |  Stage PFE PwC DigiTech  |  Avril 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.italic = True

# ── Save
output_path = r"C:\Users\PC\Desktop\SRE-Copilot\SRE-Copilot-Guide-Test.docx"
doc.save(output_path)
print(f"Document généré : {output_path}")
